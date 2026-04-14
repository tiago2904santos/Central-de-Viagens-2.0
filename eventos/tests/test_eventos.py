import json
import os
import sys
from io import BytesIO
from pathlib import Path
from datetime import date, datetime, timedelta
from decimal import Decimal
from zipfile import ZipFile
import builtins
import importlib
import unittest
from urllib.parse import quote
from unittest.mock import patch, MagicMock

from django.conf import settings
from django.core.exceptions import ValidationError
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, Client, RequestFactory
from django.urls import reverse
from django.contrib.auth import get_user_model
from django.utils import timezone as tz

from cadastros.models import AssinaturaConfiguracao, Cargo, ConfiguracaoSistema, Estado, Cidade, UnidadeLotacao, Viajante, Veiculo, CombustivelVeiculo
from eventos.models import (
    Evento,
    EventoDestino,
    EventoFinalizacao,
    EventoParticipante,
    EventoTermoParticipante,
    Justificativa,
    ModeloJustificativa,
    ModeloMotivoViagem,
    Oficio,
    OficioTrecho,
    OrdemServico,
    PlanoTrabalho,
    RoteiroEvento,
    RoteiroEventoDestino,
    RoteiroEventoTrecho,
    TermoAutorizacao,
    TipoDemandaEvento,
)
from eventos.services.justificativa import get_prazo_justificativa_dias, oficio_exige_justificativa
from eventos.services.documentos import (
    DocumentoOficioTipo,
    build_document_filename,
    build_justificativa_document_context,
    build_oficio_document_context,
    build_ordem_servico_document_context,
    build_plano_trabalho_document_context,
    build_termo_autorizacao_document_context,
    get_assinaturas_documento,
    get_document_backend_capabilities,
    get_docx_backend_availability,
    get_pdf_backend_availability,
    reset_document_backend_capabilities_cache,
    validate_oficio_for_document_generation,
)
from eventos.services.documentos.context import format_document_display
from eventos.services.documentos.oficio import build_oficio_template_context

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

User = get_user_model()


class PtBrEncodingTest(TestCase):
    """Valida strings crÃ­ticas em PT-BR sem caracteres corrompidos."""

    def test_modelos_e_choices_criticos_estao_em_pt_br_correto(self):
        textos = [
            Evento._meta.get_field('titulo').verbose_name,
            Evento._meta.get_field('descricao').verbose_name,
            Oficio._meta.verbose_name,
            dict(Oficio.CUSTEIO_CHOICES)[Oficio.CUSTEIO_UNIDADE],
            dict(Oficio.CUSTEIO_CHOICES)[Oficio.CUSTEIO_ONUS_LIMITADOS],
            dict(Evento.TIPO_CHOICES)[Evento.TIPO_OPERACAO],
            dict(Evento.TIPO_CHOICES)[Evento.TIPO_PARANA],
            ModeloMotivoViagem._meta.verbose_name_plural,
        ]
        for texto in textos:
            self.assertNotIn('\ufffd', texto)
        self.assertEqual(Evento._meta.get_field('titulo').verbose_name, 'TÃ­tulo')
        self.assertEqual(Oficio._meta.verbose_name, 'OfÃ­cio')
        self.assertEqual(
            dict(Oficio.CUSTEIO_CHOICES)[Oficio.CUSTEIO_UNIDADE],
            'UNIDADE - DPC (diÃ¡rias e combustÃ­vel custeados pela DPC).',
        )

    def test_format_document_display_corrige_nome_e_conectores_em_pt_br(self):
        raw = 'JOÃƒO Mario DE GOES, Assessor DE ComunicaÃ§Ã£o Social - PCPR'
        self.assertEqual(
            format_document_display(raw),
            'JoÃ£o Mario de Goes, Assessor de ComunicaÃ§Ã£o Social - PCPR',
        )
        self.assertEqual(
            dict(Oficio.CUSTEIO_CHOICES)[Oficio.CUSTEIO_ONUS_LIMITADOS],
            'Ã”NUS LIMITADOS AOS PRÃ“PRIOS VENCIMENTOS',
        )


class EventoListaAuthTest(TestCase):
    """Lista de eventos exige login."""

    def setUp(self):
        self.client = Client()

    def test_lista_redireciona_sem_login(self):
        response = self.client.get(reverse('eventos:lista'))
        self.assertEqual(response.status_code, 302)
        self.assertIn('login', response.url)

    def test_lista_ok_com_login(self):
        User.objects.create_user(username='u', password='p')
        self.client.login(username='u', password='p')
        response = self.client.get(reverse('eventos:lista'))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Eventos')

    def test_lista_exibe_titulo_destinos_editar_etapa_1(self):
        """Lista usa tÃ­tulo gerado, destinos e link Editar Etapa 1."""
        User.objects.create_user(username='u', password='p')
        self.client.login(username='u', password='p')
        tipo = TipoDemandaEvento.objects.filter(ativo=True).first()
        estado = Estado.objects.create(nome='ParanÃ¡', sigla='PR', codigo_ibge='41')
        cidade = Cidade.objects.create(nome='Curitiba', estado=estado, codigo_ibge='4106902')
        ev = Evento.objects.create(
            titulo='PCPR - CURITIBA - 01/01/2025',
            data_inicio=date(2025, 1, 1),
            data_fim=date(2025, 1, 1),
            status=Evento.STATUS_RASCUNHO,
        )
        ev.tipos_demanda.add(tipo)
        EventoDestino.objects.create(evento=ev, estado=estado, cidade=cidade, ordem=0)
        response = self.client.get(reverse('eventos:lista'))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'PCPR')
        self.assertContains(response, 'Curitiba')
        self.assertContains(response, 'Editar Etapa 1')
        self.assertContains(response, reverse('eventos:guiado-etapa-1', kwargs={'pk': ev.pk}))
        self.assertNotContains(response, '/guiado/painel/')


class EventoCRUDTest(TestCase):
    """CriaÃ§Ã£o e ediÃ§Ã£o unificadas no fluxo guiado."""

    def setUp(self):
        self.user = User.objects.create_user(username='u', password='p')
        self.client = Client()
        self.client.login(username='u', password='p')

    def test_cadastrar_redireciona_para_fluxo_guiado(self):
        """Cadastrar evento redireciona para guiado-novo (fonte Ãºnica de criaÃ§Ã£o)."""
        response = self.client.get(reverse('eventos:cadastrar'))
        self.assertEqual(response.status_code, 302)
        self.assertIn('guiado/novo', response.url)

    def test_editar_redireciona_para_etapa_1(self):
        """Editar evento redireciona para a Etapa 1 do fluxo guiado (mesma tela/lÃ³gica)."""
        ev = Evento.objects.create(
            titulo='Original',
            data_inicio=date(2025, 1, 1),
            data_fim=date(2025, 1, 5),
            status=Evento.STATUS_RASCUNHO,
        )
        response = self.client.get(reverse('eventos:editar', kwargs={'pk': ev.pk}))
        self.assertEqual(response.status_code, 302)
        self.assertIn(str(ev.pk), response.url)
        self.assertIn('guiado/etapa-1', response.url)


class EventoValidacaoTest(TestCase):
    """ValidaÃ§Ãµes na Etapa 1 (data e destinos)."""

    def setUp(self):
        self.user = User.objects.create_user(username='u', password='p')
        self.client = Client()
        self.client.login(username='u', password='p')
        self.tipo = TipoDemandaEvento.objects.filter(ativo=True).first()
        self.estado = Estado.objects.create(nome='ParanÃ¡', sigla='PR', codigo_ibge='41')
        self.cidade = Cidade.objects.create(nome='Curitiba', estado=self.estado, codigo_ibge='4106902')

    def test_etapa1_data_fim_menor_que_data_inicio_rejeita(self):
        """Na Etapa 1, data_fim < data_inicio Ã© rejeitado. NÃ£o enviar 'data_unica' para que seja False (checkbox desmarcado)."""
        self.assertIsNotNone(self.tipo)
        ev = Evento.objects.create(titulo='', data_inicio=date(2025, 1, 10), data_fim=date(2025, 1, 10), status=Evento.STATUS_RASCUNHO)
        data = {
            'tipos_demanda': [self.tipo.pk],
            'data_inicio': '2025-01-15',
            'data_fim': '2025-01-10',
            'descricao': '',
            'tem_convite_ou_oficio_evento': False,
            'destino_estado_0': self.estado.pk,
            'destino_cidade_0': self.cidade.pk,
        }
        response = self.client.post(reverse('eventos:guiado-etapa-1', kwargs={'pk': ev.pk}), data)
        self.assertEqual(response.status_code, 200)
        self.assertIn('data_fim', response.context['form'].errors)


class EventoDetalheTest(TestCase):
    """PÃ¡gina de detalhe do evento (modelo unificado)."""

    def setUp(self):
        self.user = User.objects.create_user(username='u', password='p')
        self.client = Client()

    def test_detalhe_redireciona_sem_login(self):
        ev = Evento.objects.create(
            titulo='E',
            data_inicio=date(2025, 1, 1),
            data_fim=date(2025, 1, 2),
        )
        response = self.client.get(reverse('eventos:detalhe', kwargs={'pk': ev.pk}))
        self.assertEqual(response.status_code, 302)
        self.assertIn('login', response.url)

    def test_detalhe_ok_mostra_dados_do_modelo_novo(self):
        """Detalhe redireciona para ediÃ§Ã£o do evento (Etapa 1)."""
        self.client.login(username='u', password='p')
        tipo = TipoDemandaEvento.objects.filter(ativo=True).first()
        estado = Estado.objects.create(nome='ParanÃ¡', sigla='PR', codigo_ibge='41')
        cidade = Cidade.objects.create(nome='Curitiba', estado=estado, codigo_ibge='4106902')
        ev = Evento.objects.create(
            titulo='PCPR - CURITIBA - 01/01/2025',
            data_inicio=date(2025, 1, 1),
            data_fim=date(2025, 1, 1),
            data_unica=True,
            status=Evento.STATUS_EM_ANDAMENTO,
        )
        ev.tipos_demanda.add(tipo)
        EventoDestino.objects.create(evento=ev, estado=estado, cidade=cidade, ordem=0)
        response = self.client.get(reverse('eventos:detalhe', kwargs={'pk': ev.pk}))
        self.assertEqual(response.status_code, 302)
        self.assertEqual(response.url, reverse('eventos:editar', kwargs={'pk': ev.pk}))


class EventoExcluirTest(TestCase):
    """ExclusÃ£o de evento: exige login, POST, bloqueia quando hÃ¡ roteiros."""

    def setUp(self):
        self.user = User.objects.create_user(username='u', password='p')
        self.client = Client()

    def test_excluir_exige_login(self):
        """ExclusÃ£o por POST redireciona para login se nÃ£o autenticado."""
        ev = Evento.objects.create(titulo='E', data_inicio=date(2025, 1, 1), data_fim=date(2025, 1, 1), status=Evento.STATUS_RASCUNHO)
        url = reverse('eventos:excluir', kwargs={'pk': ev.pk})
        response = self.client.post(url, {})
        self.assertEqual(response.status_code, 302)
        self.assertIn('login', response.url)
        self.assertTrue(Evento.objects.filter(pk=ev.pk).exists())

    def test_evento_sem_vinculos_pode_ser_excluido(self):
        """Evento sem roteiros pode ser excluÃ­do; redireciona para lista com mensagem de sucesso."""
        self.client.login(username='u', password='p')
        tipo = TipoDemandaEvento.objects.filter(ativo=True).first()
        estado = Estado.objects.create(nome='ParanÃ¡', sigla='PR', codigo_ibge='41')
        cidade = Cidade.objects.create(nome='Curitiba', estado=estado, codigo_ibge='4106902')
        ev = Evento.objects.create(titulo='Evento X', data_inicio=date(2025, 1, 1), data_fim=date(2025, 1, 1), status=Evento.STATUS_RASCUNHO)
        if tipo:
            ev.tipos_demanda.add(tipo)
        EventoDestino.objects.create(evento=ev, estado=estado, cidade=cidade, ordem=0)
        url = reverse('eventos:excluir', kwargs={'pk': ev.pk})
        response = self.client.post(url, {})
        self.assertEqual(response.status_code, 302)
        self.assertEqual(response.url, reverse('eventos:lista'))
        self.assertFalse(Evento.objects.filter(pk=ev.pk).exists())
        response_lista = self.client.get(reverse('eventos:lista'))
        self.assertContains(response_lista, 'excluÃ­do com sucesso')

    def test_evento_com_roteiros_nao_pode_ser_excluido(self):
        """Evento com roteiros vinculados nÃ£o pode ser excluÃ­do; mensagem de erro."""
        self.client.login(username='u', password='p')
        ev = Evento.objects.create(titulo='Evento Y', data_inicio=date(2025, 1, 1), data_fim=date(2025, 1, 1), status=Evento.STATUS_RASCUNHO)
        RoteiroEvento.objects.create(evento=ev)
        url = reverse('eventos:excluir', kwargs={'pk': ev.pk})
        response = self.client.post(url, {})
        self.assertEqual(response.status_code, 302)
        self.assertEqual(response.url, reverse('eventos:lista'))
        self.assertTrue(Evento.objects.filter(pk=ev.pk).exists())
        response_lista = self.client.get(reverse('eventos:lista'))
        self.assertContains(response_lista, 'nÃ£o pode ser excluÃ­do')
        self.assertContains(response_lista, 'roteiros')

    def test_excluir_redireciona_para_lista_com_sucesso(self):
        """ApÃ³s exclusÃ£o bem-sucedida, redireciona para lista e evento some."""
        self.client.login(username='u', password='p')
        ev = Evento.objects.create(titulo='Z', data_inicio=date(2025, 1, 1), data_fim=date(2025, 1, 1), status=Evento.STATUS_RASCUNHO)
        pk = ev.pk
        response = self.client.post(reverse('eventos:excluir', kwargs={'pk': pk}), {})
        self.assertEqual(response.status_code, 302)
        self.assertEqual(response.url, reverse('eventos:lista'))
        self.assertFalse(Evento.objects.filter(pk=pk).exists())

    def test_excluir_so_aceita_post(self):
        """ExclusÃ£o sÃ³ aceita POST; GET retorna 405."""
        self.client.login(username='u', password='p')
        ev = Evento.objects.create(titulo='W', data_inicio=date(2025, 1, 1), data_fim=date(2025, 1, 1), status=Evento.STATUS_RASCUNHO)
        response = self.client.get(reverse('eventos:excluir', kwargs={'pk': ev.pk}))
        self.assertEqual(response.status_code, 405)
        self.assertTrue(Evento.objects.filter(pk=ev.pk).exists())


class EventoGuiadoTest(TestCase):
    """Fluxo guiado: novo, etapa 1, painel."""

    def setUp(self):
        self.user = User.objects.create_user(username='u', password='p')
        self.client = Client()
        self.client.login(username='u', password='p')

    def test_guiado_novo_exige_login_e_redireciona(self):
        self.client.logout()
        response = self.client.post(reverse('eventos:guiado-novo'))
        self.assertEqual(response.status_code, 302)
        self.assertIn('login', response.url)

    def test_guiado_novo_cria_evento_e_redireciona_para_etapa_1(self):
        response = self.client.post(reverse('eventos:guiado-novo'))
        self.assertEqual(response.status_code, 302)
        self.assertEqual(Evento.objects.count(), 1)
        ev = Evento.objects.get()
        self.assertEqual(ev.status, Evento.STATUS_RASCUNHO)
        self.assertIn(str(ev.pk), response.url)
        self.assertIn('etapa-1', response.url)

    def test_etapa_1_salva_corretamente(self):
        tipo = TipoDemandaEvento.objects.filter(ativo=True).first()
        self.assertIsNotNone(tipo, 'Precisa existir pelo menos um TipoDemandaEvento (migraÃ§Ã£o 0004).')
        estado = Estado.objects.create(nome='ParanÃ¡', sigla='PR', codigo_ibge='41')
        cidade = Cidade.objects.create(nome='Curitiba', estado=estado, codigo_ibge='4106902')
        ev = Evento.objects.create(
            titulo='',
            data_inicio=date(2025, 2, 1),
            data_fim=date(2025, 2, 5),
            status=Evento.STATUS_RASCUNHO,
        )
        data = {
            'tipos_demanda': [tipo.pk],
            'data_inicio': '2025-02-01',
            'data_fim': '2025-02-05',
            'descricao': 'Desc',
            'tem_convite_ou_oficio_evento': False,
            'destino_estado_0': estado.pk,
            'destino_cidade_0': cidade.pk,
        }
        response = self.client.post(reverse('eventos:guiado-etapa-1', kwargs={'pk': ev.pk}), data)
        self.assertEqual(response.status_code, 302, response.context.get('form', {}).errors if response.context else None)
        ev.refresh_from_db()
        self.assertIn(tipo.nome, ev.titulo)
        self.assertIn('CURITIBA', ev.titulo)
        self.assertEqual(ev.data_inicio, date(2025, 2, 1))
        self.assertEqual(ev.data_fim, date(2025, 2, 5))
        self.assertFalse(ev.data_unica)
        if not tipo.is_outros:
            self.assertEqual(ev.descricao, '')  # sem OUTROS a descriÃ§Ã£o fica vazia
        self.assertEqual(ev.destinos.count(), 1)
        self.assertEqual(ev.tipos_demanda.count(), 1)

    def test_etapa_1_destino_cidade_fora_do_estado_gera_erro(self):
        tipo = TipoDemandaEvento.objects.filter(ativo=True).first()
        self.assertIsNotNone(tipo)
        sp = Estado.objects.create(nome='SÃ£o Paulo', sigla='SP', codigo_ibge='35')
        rj = Estado.objects.create(nome='Rio de Janeiro', sigla='RJ', codigo_ibge='33')
        cidade_rj = Cidade.objects.create(nome='Rio', estado=rj, codigo_ibge='3304557')
        ev = Evento.objects.create(
            titulo='E',
            data_inicio=date(2025, 1, 1),
            data_fim=date(2025, 1, 5),
            status=Evento.STATUS_RASCUNHO,
        )
        data = {
            'tipos_demanda': [tipo.pk],
            'data_inicio': '2025-01-01',
            'data_fim': '2025-01-05',
            'destino_estado_0': sp.pk,
            'destino_cidade_0': cidade_rj.pk,
            'descricao': 'Desc',
            'tem_convite_ou_oficio_evento': False,
        }
        response = self.client.post(reverse('eventos:guiado-etapa-1', kwargs={'pk': ev.pk}), data)
        self.assertEqual(response.status_code, 200)
        self.assertIn('form', response.context)
        form = response.context['form']
        self.assertTrue(form.errors.get('__all__') or 'cidade' in str(form.errors).lower() or 'estado' in str(form.errors).lower())

class EventoEtapa2RoteirosTest(TestCase):
    """Etapa 2: Roteiros do evento."""

    def setUp(self):
        self.user = User.objects.create_user(username='u', password='p')
        self.client = Client()
        self.client.login(username='u', password='p')
        self.estado = Estado.objects.create(nome='ParanÃ¡', sigla='PR', codigo_ibge='41')
        self.cidade_a = Cidade.objects.create(nome='Curitiba', estado=self.estado, codigo_ibge='4106902')
        self.cidade_b = Cidade.objects.create(nome='Londrina', estado=self.estado, codigo_ibge='4113700')
        self.evento = Evento.objects.create(
            titulo='Evento Roteiros',
            tipo_demanda=Evento.TIPO_OUTRO,
            data_inicio=date(2025, 1, 1),
            data_fim=date(2025, 1, 5),
            status=Evento.STATUS_EM_ANDAMENTO,
            cidade_base=self.cidade_a,
        )

    def test_etapa_2_lista_exige_login(self):
        self.client.logout()
        response = self.client.get(reverse('eventos:guiado-etapa-2', kwargs={'evento_id': self.evento.pk}))
        self.assertEqual(response.status_code, 302)
        self.assertIn('login', response.url)

    def test_etapa_2_lista_filtra_apenas_roteiros_do_evento_atual(self):
        outro_evento = Evento.objects.create(
            titulo='Outro evento',
            tipo_demanda=Evento.TIPO_OUTRO,
            data_inicio=date(2025, 2, 1),
            data_fim=date(2025, 2, 3),
            status=Evento.STATUS_EM_ANDAMENTO,
            cidade_base=self.cidade_a,
        )
        roteiro_evento = RoteiroEvento.objects.create(
            evento=self.evento,
            origem_estado=self.estado,
            origem_cidade=self.cidade_a,
            status=RoteiroEvento.STATUS_FINALIZADO,
            tipo=RoteiroEvento.TIPO_EVENTO,
        )
        RoteiroEventoDestino.objects.create(roteiro=roteiro_evento, estado=self.estado, cidade=self.cidade_b, ordem=0)
        roteiro_outro_evento = RoteiroEvento.objects.create(
            evento=outro_evento,
            origem_estado=self.estado,
            origem_cidade=self.cidade_a,
            status=RoteiroEvento.STATUS_FINALIZADO,
            tipo=RoteiroEvento.TIPO_EVENTO,
        )
        RoteiroEventoDestino.objects.create(roteiro=roteiro_outro_evento, estado=self.estado, cidade=self.cidade_b, ordem=0)

        response = self.client.get(reverse('eventos:guiado-etapa-2', kwargs={'evento_id': self.evento.pk}))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, f'roteiro-card-{roteiro_evento.pk}')
        self.assertNotContains(response, f'roteiro-card-{roteiro_outro_evento.pk}')
        self.assertEqual([item.pk for item in response.context['object_list']], [roteiro_evento.pk])

    def test_etapa_2_lista_usa_shell_moderno_com_toggle_persistido_e_fab(self):
        response = self.client.get(reverse('eventos:guiado-etapa-2', kwargs={'evento_id': self.evento.pk}))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'data-view-storage-key="central-viagens.guiado.etapa2.roteiros.view-mode"')
        self.assertContains(response, 'oficios-view-pane--basic')
        self.assertContains(response, 'oficios-view-pane--rich')
        self.assertContains(response, 'aria-label="Alternar visualizacao dos roteiros do evento"')
        self.assertContains(response, 'Subir lista')
        self.assertContains(response, 'Novo roteiro')

    def test_criar_roteiro_dentro_de_evento(self):
        saida = datetime(2025, 1, 2, 8, 0)
        data = {
            'origem_estado': self.estado.pk,
            'origem_cidade': self.cidade_a.pk,
            'destino_estado_0': self.estado.pk,
            'destino_cidade_0': self.cidade_b.pk,
            'saida_dt': saida.strftime('%Y-%m-%dT%H:%M'),
            'retorno_saida_dt': '',
            'observacoes': '',
            'trecho_0_saida_dt': saida.strftime('%Y-%m-%dT%H:%M'),
            'trecho_0_chegada_dt': (saida + timedelta(minutes=120)).strftime('%Y-%m-%dT%H:%M'),
            'trecho_1_saida_dt': '',
            'trecho_1_chegada_dt': '',
        }
        response = self.client.post(
            reverse('eventos:guiado-etapa-2-cadastrar', kwargs={'evento_id': self.evento.pk}),
            data,
        )
        self.assertEqual(response.status_code, 302)
        self.assertEqual(RoteiroEvento.objects.filter(evento=self.evento).count(), 1)
        r = RoteiroEvento.objects.get(evento=self.evento)
        self.assertEqual(r.origem_cidade_id, self.cidade_a.pk)
        self.assertEqual(r.destinos.count(), 1)
        self.assertEqual(r.destinos.first().cidade_id, self.cidade_b.pk)
        self.assertIsNotNone(r.saida_dt)
        self.assertIsNotNone(r.chegada_dt, 'chegada_dt preenchido a partir dos trechos')
        self.assertEqual((r.chegada_dt - r.saida_dt).total_seconds(), 120 * 60)

    def test_origem_padrao_vem_da_configuracao_sede(self):
        """Sede prÃ©-preenchida vem de ConfiguracaoSistema.cidade_sede_padrao."""
        from cadastros.models import ConfiguracaoSistema
        config = ConfiguracaoSistema.get_singleton()
        config.cidade_sede_padrao = self.cidade_a
        config.save(update_fields=['cidade_sede_padrao'])
        response = self.client.get(
            reverse('eventos:guiado-etapa-2-cadastrar', kwargs={'evento_id': self.evento.pk})
        )
        self.assertEqual(response.status_code, 200)
        form = response.context['form']
        self.assertEqual(form.initial.get('origem_cidade'), self.cidade_a.pk)
        self.assertEqual(form.initial.get('origem_estado'), self.estado.pk)

    def test_origem_padrao_cai_para_cidade_base_do_evento_quando_nao_houver_config(self):
        """Sem cidade_sede_padrao, a etapa 2 deve usar a cidade base do evento."""
        from cadastros.models import ConfiguracaoSistema

        config = ConfiguracaoSistema.get_singleton()
        config.cidade_sede_padrao = None
        config.save(update_fields=['cidade_sede_padrao'])
        self.evento.cidade_base = self.cidade_a
        self.evento.save(update_fields=['cidade_base'])
        EventoDestino.objects.create(evento=self.evento, estado=self.estado, cidade=self.cidade_b, ordem=0)

        response = self.client.get(
            reverse('eventos:guiado-etapa-2-cadastrar', kwargs={'evento_id': self.evento.pk})
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.context['sede_cidade_id'], self.cidade_a.pk)
        self.assertEqual(response.context['sede_estado_id'], self.estado.pk)
        self.assertGreaterEqual(len(response.context.get('trechos') or []), 2)
        self.assertContains(response, f'value="{self.cidade_a.pk}" selected')

    def test_cidade_destino_deve_pertencer_ao_estado(self):
        outro_estado = Estado.objects.create(nome='SÃ£o Paulo', sigla='SP', codigo_ibge='35')
        cidade_sp = Cidade.objects.create(nome='SÃ£o Paulo', estado=outro_estado, codigo_ibge='3550308')
        data = {
            'origem_estado': self.estado.pk,
            'origem_cidade': self.cidade_a.pk,
            'destino_estado_0': self.estado.pk,
            'destino_cidade_0': cidade_sp.pk,
            'saida_dt': '2025-01-02T08:00',
            'retorno_saida_dt': '',
            'observacoes': '',
        }
        response = self.client.post(
            reverse('eventos:guiado-etapa-2-cadastrar', kwargs={'evento_id': self.evento.pk}),
            data,
        )
        self.assertEqual(response.status_code, 200)
        self.assertTrue(response.context['form'].errors.get('__all__') or 'pertencer' in str(response.context['form'].errors).lower())
        self.assertEqual(RoteiroEvento.objects.filter(evento=self.evento).count(), 0)

    def test_calculo_chegada_ao_salvar(self):
        """chegada_dt do roteiro Ã© preenchida a partir do Ãºltimo trecho de ida."""
        data = {
            'origem_estado': self.estado.pk,
            'origem_cidade': self.cidade_a.pk,
            'destino_estado_0': self.estado.pk,
            'destino_cidade_0': self.cidade_b.pk,
            'saida_dt': '2025-01-02T14:00',
            'retorno_saida_dt': '',
            'observacoes': '',
            'trecho_0_saida_dt': '2025-01-02T14:00',
            'trecho_0_chegada_dt': '2025-01-02T15:30',
            'trecho_1_saida_dt': '',
            'trecho_1_chegada_dt': '',
        }
        self.client.post(
            reverse('eventos:guiado-etapa-2-cadastrar', kwargs={'evento_id': self.evento.pk}),
            data,
        )
        r = RoteiroEvento.objects.get(evento=self.evento)
        self.assertIsNotNone(r.chegada_dt)
        self.assertEqual((r.chegada_dt - r.saida_dt).total_seconds(), 90 * 60)

    def test_salvar_incompleto_rascunho(self):
        """Com um destino mas sem saida/duracao o roteiro fica RASCUNHO."""
        data = {
            'origem_estado': self.estado.pk,
            'origem_cidade': self.cidade_a.pk,
            'destino_estado_0': self.estado.pk,
            'destino_cidade_0': self.cidade_b.pk,
            'saida_dt': '',
            'retorno_saida_dt': '',
            'observacoes': 'PARCIAL',
        }
        response = self.client.post(
            reverse('eventos:guiado-etapa-2-cadastrar', kwargs={'evento_id': self.evento.pk}),
            data,
        )
        self.assertEqual(response.status_code, 302)
        r = RoteiroEvento.objects.get(evento=self.evento)
        self.assertEqual(r.status, RoteiroEvento.STATUS_RASCUNHO)
        self.assertEqual(r.observacoes, 'PARCIAL')

    def test_autosave_cria_roteiro_em_rascunho_e_retorna_url_edicao(self):
        response = self.client.post(
            reverse('eventos:guiado-etapa-2-cadastrar', kwargs={'evento_id': self.evento.pk}),
            {
                'autosave': '1',
                'autosave_obj_id': '',
                'origem_estado': self.estado.pk,
                'origem_cidade': self.cidade_a.pk,
                'destino_estado_0': self.estado.pk,
                'destino_cidade_0': self.cidade_b.pk,
                'observacoes': 'rascunho autosave',
            },
            HTTP_X_REQUESTED_WITH='XMLHttpRequest',
        )

        self.assertEqual(response.status_code, 200)
        body = response.json()
        self.assertTrue(body.get('ok'))
        self.assertTrue(body.get('edit_url'))

        roteiro = RoteiroEvento.objects.get(evento=self.evento)
        self.assertEqual(roteiro.status, RoteiroEvento.STATUS_RASCUNHO)
        self.assertEqual(roteiro.observacoes, 'RASCUNHO AUTOSAVE')
        self.assertIn(str(roteiro.pk), body.get('edit_url'))

    def test_salvar_completo_finalizado(self):
        """Roteiro fica FINALIZADO quando tem saida_dt e chegada_dt (preenchidos a partir dos trechos)."""
        data = {
            'origem_estado': self.estado.pk,
            'origem_cidade': self.cidade_a.pk,
            'destino_estado_0': self.estado.pk,
            'destino_cidade_0': self.cidade_b.pk,
            'saida_dt': '2025-01-02T08:00',
            'retorno_saida_dt': '',
            'observacoes': '',
            'trecho_0_saida_dt': '2025-01-02T08:00',
            'trecho_0_chegada_dt': '2025-01-02T10:00',
            'trecho_1_saida_dt': '',
            'trecho_1_chegada_dt': '',
        }
        self.client.post(
            reverse('eventos:guiado-etapa-2-cadastrar', kwargs={'evento_id': self.evento.pk}),
            data,
        )
        r = RoteiroEvento.objects.get(evento=self.evento)
        self.assertEqual(r.status, RoteiroEvento.STATUS_FINALIZADO)

    def test_excluir_remove_do_banco(self):
        r = RoteiroEvento.objects.create(
            evento=self.evento,
            origem_estado=self.estado,
            origem_cidade=self.cidade_a,
            saida_dt=datetime(2025, 1, 2, 8, 0),
            duracao_min=60,
            chegada_dt=datetime(2025, 1, 2, 9, 0),
            status=RoteiroEvento.STATUS_FINALIZADO,
        )
        RoteiroEventoDestino.objects.create(roteiro=r, estado=self.estado, cidade=self.cidade_b, ordem=0)
        response = self.client.post(
            reverse('eventos:guiado-etapa-2-excluir', kwargs={'evento_id': self.evento.pk, 'pk': r.pk}),
            {},
        )
        self.assertEqual(response.status_code, 302)
        self.assertFalse(RoteiroEvento.objects.filter(pk=r.pk).exists())

    def test_cadastro_roteiro_herda_sede_da_configuracao(self):
        """Cadastro novo de roteiro prÃ©-preenche sede com ConfiguracaoSistema.cidade_sede_padrao."""
        from cadastros.models import ConfiguracaoSistema
        config = ConfiguracaoSistema.get_singleton()
        config.cidade_sede_padrao = self.cidade_a
        config.save(update_fields=['cidade_sede_padrao'])
        response = self.client.get(reverse('eventos:guiado-etapa-2-cadastrar', kwargs={'evento_id': self.evento.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.context['form'].initial.get('origem_estado'), self.estado.pk)
        self.assertEqual(response.context['form'].initial.get('origem_cidade'), self.cidade_a.pk)

    def test_cadastro_roteiro_herda_destinos_do_evento(self):
        """Cadastro novo de roteiro prÃ©-preenche destinos da Etapa 1 do evento."""
        EventoDestino.objects.create(evento=self.evento, estado=self.estado, cidade=self.cidade_a, ordem=0)
        EventoDestino.objects.create(evento=self.evento, estado=self.estado, cidade=self.cidade_b, ordem=1)
        response = self.client.get(reverse('eventos:guiado-etapa-2-cadastrar', kwargs={'evento_id': self.evento.pk}))
        self.assertEqual(response.status_code, 200)
        destinos_atuais = response.context['destinos_atuais']
        self.assertEqual(len(destinos_atuais), 2)
        self.assertEqual(destinos_atuais[0]['cidade_id'], self.cidade_a.pk)
        self.assertEqual(destinos_atuais[1]['cidade_id'], self.cidade_b.pk)

    def test_edicao_roteiro_mostra_dados_salvos_nao_do_evento(self):
        """EdiÃ§Ã£o do roteiro exibe os destinos salvos do roteiro, nÃ£o os atuais do evento."""
        r = RoteiroEvento.objects.create(
            evento=self.evento,
            origem_estado=self.estado,
            origem_cidade=self.cidade_a,
            saida_dt=datetime(2025, 1, 2, 8, 0),
            duracao_min=60,
        )
        RoteiroEventoDestino.objects.create(roteiro=r, estado=self.estado, cidade=self.cidade_b, ordem=0)
        EventoDestino.objects.create(evento=self.evento, estado=self.estado, cidade=self.cidade_a, ordem=0)
        response = self.client.get(reverse('eventos:guiado-etapa-2-editar', kwargs={'evento_id': self.evento.pk, 'pk': r.pk}))
        self.assertEqual(response.status_code, 200)
        destinos_atuais = response.context['destinos_atuais']
        self.assertEqual(len(destinos_atuais), 1)
        self.assertEqual(destinos_atuais[0]['cidade_id'], self.cidade_b.pk)

    def test_cadastro_roteiro_abre_sem_quebrar_sem_sede_ni_destinos(self):
        """FormulÃ¡rio de cadastro abre mesmo sem cidade sede na config e sem destinos no evento."""
        response = self.client.get(reverse('eventos:guiado-etapa-2-cadastrar', kwargs={'evento_id': self.evento.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertIn('destinos_atuais', response.context)
        self.assertIn('form', response.context)

    def test_cadastro_roteiro_multiplos_destinos_evento(self):
        """MÃºltiplos destinos do evento aparecem no formulÃ¡rio de novo roteiro."""
        EventoDestino.objects.create(evento=self.evento, estado=self.estado, cidade=self.cidade_a, ordem=0)
        EventoDestino.objects.create(evento=self.evento, estado=self.estado, cidade=self.cidade_b, ordem=1)
        response = self.client.get(reverse('eventos:guiado-etapa-2-cadastrar', kwargs={'evento_id': self.evento.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertEqual(len(response.context['destinos_atuais']), 2)
        self.assertContains(response, 'destino_estado_0')
        self.assertContains(response, 'destino_estado_1')

    def test_bloco_duracao_apoio_removido(self):
        """O bloco '3) DuraÃ§Ã£o (apoio)' e o campo DuraÃ§Ã£o (HH:MM) foram removidos da Etapa 2."""
        response = self.client.get(
            reverse('eventos:guiado-etapa-2-cadastrar', kwargs={'evento_id': self.evento.pk})
        )
        self.assertEqual(response.status_code, 200)
        content = response.content.decode('utf-8')
        self.assertNotIn('DuraÃ§Ã£o (apoio)', content)
        self.assertNotIn('id_duracao_hhmm', content)

    def test_bloco_global_ida_retorno_nao_aparece(self):
        """FormulÃ¡rio nÃ£o exibe mais bloco global de SaÃ­da ida / SaÃ­da retorno / Chegada calculada."""
        r = RoteiroEvento.objects.create(
            evento=self.evento,
            origem_estado=self.estado,
            origem_cidade=self.cidade_a,
        )
        RoteiroEventoDestino.objects.create(roteiro=r, estado=self.estado, cidade=self.cidade_b, ordem=0)
        response = self.client.get(reverse('eventos:guiado-etapa-2-editar', kwargs={'evento_id': self.evento.pk, 'pk': r.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, 'SaÃ­da ida â€” data')
        self.assertNotContains(response, 'SaÃ­da retorno â€” data')
        self.assertNotContains(response, 'Chegada ida (calculada)')
        self.assertNotContains(response, 'Chegada retorno (calculada)')

    def test_trechos_gerados_container_presente(self):
        """PÃ¡gina de roteiro exibe o bloco de trechos (cada trecho com campos prÃ³prios, preenchido via JS)."""
        r = RoteiroEvento.objects.create(
            evento=self.evento,
            origem_estado=self.estado,
            origem_cidade=self.cidade_a,
        )
        RoteiroEventoDestino.objects.create(roteiro=r, estado=self.estado, cidade=self.cidade_b, ordem=0)
        response = self.client.get(reverse('eventos:guiado-etapa-2-editar', kwargs={'evento_id': self.evento.pk, 'pk': r.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, '3) Trechos')
        self.assertContains(response, 'trechos-gerados-container')

    def test_script_trechos_tem_campos_por_trecho(self):
        """JS gera inputs de saÃ­da/chegada por trecho (saÃ­da data/hora, chegada data/hora)."""
        r = RoteiroEvento.objects.create(
            evento=self.evento,
            origem_estado=self.estado,
            origem_cidade=self.cidade_a,
        )
        RoteiroEventoDestino.objects.create(roteiro=r, estado=self.estado, cidade=self.cidade_b, ordem=0)
        response = self.client.get(reverse('eventos:guiado-etapa-2-editar', kwargs={'evento_id': self.evento.pk, 'pk': r.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertIn(b'_saida_date', response.content)
        self.assertIn(b'_saida_time', response.content)
        self.assertIn(b'_chegada_date', response.content)
        self.assertIn(b'_chegada_time', response.content)

    def test_multiplos_destinos_geram_trechos_no_contexto(self):
        """MÃºltiplos destinos geram mÃºltiplos trechos (ida + retorno) no contexto para o JS."""
        r = RoteiroEvento.objects.create(
            evento=self.evento,
            origem_estado=self.estado,
            origem_cidade=self.cidade_a,
        )
        RoteiroEventoDestino.objects.create(roteiro=r, estado=self.estado, cidade=self.cidade_b, ordem=0)
        RoteiroEventoDestino.objects.create(roteiro=r, estado=self.estado, cidade=self.cidade_a, ordem=1)
        response = self.client.get(reverse('eventos:guiado-etapa-2-editar', kwargs={'evento_id': self.evento.pk, 'pk': r.pk}))
        self.assertEqual(response.status_code, 200)
        trechos = response.context.get('trechos') or []
        self.assertGreaterEqual(len(trechos), 3)

    def test_salvar_reabrir_mantem_horarios_por_trecho(self):
        """Salvar ediÃ§Ã£o e reabrir: horÃ¡rios de cada trecho permanecem (persistÃªncia por trecho)."""
        from eventos.models import RoteiroEventoTrecho
        r = RoteiroEvento.objects.create(
            evento=self.evento,
            origem_estado=self.estado,
            origem_cidade=self.cidade_a,
        )
        RoteiroEventoDestino.objects.create(roteiro=r, estado=self.estado, cidade=self.cidade_b, ordem=0)
        data = {
            'origem_estado': self.estado.pk,
            'origem_cidade': self.cidade_a.pk,
            'destino_estado_0': self.estado.pk,
            'destino_cidade_0': self.cidade_b.pk,
            'observacoes': '',
            'trecho_0_saida_dt': '2025-02-10T09:00',
            'trecho_0_chegada_dt': '2025-02-10T10:30',
            'trecho_1_saida_dt': '2025-02-12T14:00',
            'trecho_1_chegada_dt': '2025-02-12T15:30',
        }
        self.client.post(
            reverse('eventos:guiado-etapa-2-editar', kwargs={'evento_id': self.evento.pk, 'pk': r.pk}),
            data,
        )
        r.refresh_from_db()
        trechos = list(RoteiroEventoTrecho.objects.filter(roteiro=r).order_by('ordem'))
        self.assertEqual(len(trechos), 2)
        self.assertIsNotNone(trechos[0].saida_dt)
        self.assertIsNotNone(trechos[0].chegada_dt)
        self.assertEqual((trechos[0].chegada_dt - trechos[0].saida_dt).total_seconds(), 90 * 60)
        self.assertIsNotNone(trechos[1].chegada_dt)
        response2 = self.client.get(reverse('eventos:guiado-etapa-2-editar', kwargs={'evento_id': self.evento.pk, 'pk': r.pk}))
        self.assertEqual(response2.status_code, 200)
        trechos_json = response2.context.get('trechos_json', '[]')
        import json
        initial = json.loads(trechos_json) if isinstance(trechos_json, str) else trechos_json
        self.assertEqual(len(initial), 2)
        self.assertIn('2025-02-10', initial[0].get('saida_dt', ''))
        self.assertIn('2025-02-12', initial[1].get('chegada_dt', ''))

    def test_multiplos_trechos_horarios_no_trecho_certo_salvar_reabrir(self):
        """Salvar com mÃºltiplos trechos (sede->d1->d2->sede) mantÃ©m cada saÃ­da/chegada no trecho correto; reabrir preserva."""
        from eventos.models import RoteiroEventoTrecho
        cidade_c = Cidade.objects.create(nome='MaringÃ¡', estado=self.estado, codigo_ibge='4115200')
        EventoDestino.objects.create(evento=self.evento, estado=self.estado, cidade=self.cidade_b, ordem=0)
        EventoDestino.objects.create(evento=self.evento, estado=self.estado, cidade=cidade_c, ordem=1)
        r = RoteiroEvento.objects.create(
            evento=self.evento,
            origem_estado=self.estado,
            origem_cidade=self.cidade_a,
        )
        RoteiroEventoDestino.objects.create(roteiro=r, estado=self.estado, cidade=self.cidade_b, ordem=0)
        RoteiroEventoDestino.objects.create(roteiro=r, estado=self.estado, cidade=cidade_c, ordem=1)
        data = {
            'origem_estado': self.estado.pk,
            'origem_cidade': self.cidade_a.pk,
            'saida_dt': '2025-02-10T08:00',
            'retorno_saida_dt': '2025-02-11T09:00',
            'destino_estado_0': self.estado.pk,
            'destino_cidade_0': self.cidade_b.pk,
            'destino_estado_1': self.estado.pk,
            'destino_cidade_1': cidade_c.pk,
            'observacoes': '',
            'trecho_0_saida_dt': '2025-02-10T08:00',
            'trecho_0_chegada_dt': '2025-02-10T10:00',
            'trecho_1_saida_dt': '2025-02-10T11:00',
            'trecho_1_chegada_dt': '2025-02-10T13:00',
            'trecho_2_saida_dt': '2025-02-11T09:00',
            'trecho_2_chegada_dt': '2025-02-11T12:00',
        }
        response = self.client.post(
            reverse('eventos:guiado-etapa-2-editar', kwargs={'evento_id': self.evento.pk, 'pk': r.pk}),
            data,
        )
        self.assertEqual(response.status_code, 302)
        trechos = list(RoteiroEventoTrecho.objects.filter(roteiro=r).order_by('ordem'))
        self.assertEqual(len(trechos), 3)
        self.assertEqual(trechos[0].tipo, RoteiroEventoTrecho.TIPO_IDA)
        self.assertEqual(tz.localtime(trechos[0].saida_dt).replace(tzinfo=None), datetime(2025, 2, 10, 8, 0))
        self.assertEqual(tz.localtime(trechos[0].chegada_dt).replace(tzinfo=None), datetime(2025, 2, 10, 10, 0))
        self.assertEqual(trechos[1].tipo, RoteiroEventoTrecho.TIPO_IDA)
        self.assertEqual(tz.localtime(trechos[1].saida_dt).replace(tzinfo=None), datetime(2025, 2, 10, 11, 0))
        self.assertEqual(tz.localtime(trechos[1].chegada_dt).replace(tzinfo=None), datetime(2025, 2, 10, 13, 0))
        self.assertEqual(trechos[2].tipo, RoteiroEventoTrecho.TIPO_RETORNO)
        self.assertEqual(tz.localtime(trechos[2].saida_dt).replace(tzinfo=None), datetime(2025, 2, 11, 9, 0))
        self.assertEqual(tz.localtime(trechos[2].chegada_dt).replace(tzinfo=None), datetime(2025, 2, 11, 12, 0))
        response2 = self.client.get(
            reverse('eventos:guiado-etapa-2-editar', kwargs={'evento_id': self.evento.pk, 'pk': r.pk})
        )
        self.assertEqual(response2.status_code, 200)
        trechos_json = response2.context.get('trechos_json', '[]')
        initial = json.loads(trechos_json) if isinstance(trechos_json, str) else trechos_json
        self.assertEqual(len(initial), 3)
        self.assertIn('2025-02-10T08:00', initial[0].get('saida_dt', ''))
        self.assertIn('2025-02-10T10:00', initial[0].get('chegada_dt', ''))
        self.assertIn('2025-02-10T11:00', initial[1].get('saida_dt', ''))
        self.assertIn('2025-02-10T13:00', initial[1].get('chegada_dt', ''))
        self.assertIn('2025-02-11T09:00', initial[2].get('saida_dt', ''))
        self.assertIn('2025-02-11T12:00', initial[2].get('chegada_dt', ''))

    def test_parse_trechos_times_post_ordem_correta(self):
        """_parse_trechos_times_post retorna lista na ordem trecho_0, trecho_1, ... para associaÃ§Ã£o correta."""
        from eventos.views import _parse_trechos_times_post
        rf = RequestFactory()
        request = rf.post('/x/', {
            'trecho_0_saida_dt': '2025-02-10T08:00',
            'trecho_0_chegada_dt': '2025-02-10T10:00',
            'trecho_1_saida_dt': '2025-02-10T11:00',
            'trecho_1_chegada_dt': '2025-02-10T13:00',
            'trecho_2_saida_dt': '2025-02-11T09:00',
            'trecho_2_chegada_dt': '2025-02-11T12:00',
        })
        result = _parse_trechos_times_post(request, 3)
        self.assertEqual(len(result), 3)
        self.assertIsNotNone(result[0].get('saida_dt'))
        self.assertEqual(result[0]['saida_dt'].hour, 8)
        self.assertEqual(result[0]['chegada_dt'].hour, 10)
        self.assertEqual(result[1]['saida_dt'].hour, 11)
        self.assertEqual(result[1]['chegada_dt'].hour, 13)
        self.assertEqual(result[2]['saida_dt'].hour, 9)
        self.assertEqual(result[2]['chegada_dt'].hour, 12)

    def test_salvar_trechos_roteiro_associa_por_ordem(self):
        """_salvar_trechos_roteiro associa trechos_data[0] ao primeiro trecho, etc."""
        from eventos.views import _salvar_trechos_roteiro
        cidade_c = Cidade.objects.create(nome='MaringÃ¡', estado=self.estado, codigo_ibge='4115200')
        r = RoteiroEvento.objects.create(
            evento=self.evento,
            origem_estado=self.estado,
            origem_cidade=self.cidade_a,
        )
        RoteiroEventoDestino.objects.create(roteiro=r, estado=self.estado, cidade=self.cidade_b, ordem=0)
        RoteiroEventoDestino.objects.create(roteiro=r, estado=self.estado, cidade=cidade_c, ordem=1)
        destinos_list = [(self.estado.pk, self.cidade_b.pk), (self.estado.pk, cidade_c.pk)]
        d0 = {'saida_dt': datetime(2025, 2, 10, 8, 0), 'chegada_dt': datetime(2025, 2, 10, 10, 0)}
        d1 = {'saida_dt': datetime(2025, 2, 10, 11, 0), 'chegada_dt': datetime(2025, 2, 10, 13, 0)}
        d2 = {'saida_dt': datetime(2025, 2, 11, 9, 0), 'chegada_dt': datetime(2025, 2, 11, 12, 0)}
        trechos_data = [d0, d1, d2]
        _salvar_trechos_roteiro(r, destinos_list, trechos_data)
        t0 = RoteiroEventoTrecho.objects.get(roteiro=r, ordem=0)
        t1 = RoteiroEventoTrecho.objects.get(roteiro=r, ordem=1)
        t2 = RoteiroEventoTrecho.objects.get(roteiro=r, ordem=2)
        local0 = tz.localtime(t0.saida_dt).replace(tzinfo=None) if t0.saida_dt else None
        local1 = tz.localtime(t1.saida_dt).replace(tzinfo=None) if t1.saida_dt else None
        local2 = tz.localtime(t2.saida_dt).replace(tzinfo=None) if t2.saida_dt else None
        self.assertEqual(local0, datetime(2025, 2, 10, 8, 0), 'trecho ordem=0 deve ter saida 08:00')
        self.assertEqual(local1, datetime(2025, 2, 10, 11, 0), 'trecho ordem=1 deve ter saida 11:00')
        self.assertEqual(local2, datetime(2025, 2, 11, 9, 0), 'trecho ordem=2 deve ter saida 09:00')

    def test_trecho_create_persiste_saida_dt(self):
        """Model persiste saida_dt corretamente ao criar trecho direto."""
        r = RoteiroEvento.objects.create(
            evento=self.evento,
            origem_estado=self.estado,
            origem_cidade=self.cidade_a,
        )
        RoteiroEventoTrecho.objects.create(
            roteiro=r, ordem=0, tipo=RoteiroEventoTrecho.TIPO_IDA,
            origem_estado=self.estado, origem_cidade=self.cidade_a,
            destino_estado=self.estado, destino_cidade=self.cidade_b,
            saida_dt=datetime(2025, 2, 10, 8, 0),
            chegada_dt=datetime(2025, 2, 10, 10, 0),
        )
        t0 = RoteiroEventoTrecho.objects.get(roteiro=r, ordem=0)
        # USE_TZ=True: comparar em hora local
        local_saida = tz.localtime(t0.saida_dt).replace(tzinfo=None) if t0.saida_dt else None
        self.assertEqual(local_saida, datetime(2025, 2, 10, 8, 0))

    def test_salvar_trechos_um_destino_dois_trechos(self):
        """Com 1 destino: trecho 0 = ida (trechos_data[0]), trecho 1 = retorno (trechos_data[1])."""
        from eventos.views import _salvar_trechos_roteiro
        r = RoteiroEvento.objects.create(
            evento=self.evento,
            origem_estado=self.estado,
            origem_cidade=self.cidade_a,
        )
        destinos_list = [(self.estado.pk, self.cidade_b.pk)]
        d0 = {'saida_dt': datetime(2025, 2, 10, 8, 0), 'chegada_dt': datetime(2025, 2, 10, 10, 0)}
        d1 = {'saida_dt': datetime(2025, 2, 11, 9, 0), 'chegada_dt': datetime(2025, 2, 11, 12, 0)}
        self.assertEqual(d0['saida_dt'].hour, 8, 'd0 deve ter hora 8')
        trechos_data = [d0, d1]
        _salvar_trechos_roteiro(r, destinos_list, trechos_data)
        self.assertEqual(RoteiroEventoTrecho.objects.filter(roteiro=r).count(), 2)
        t0 = RoteiroEventoTrecho.objects.get(roteiro=r, ordem=0)
        t1 = RoteiroEventoTrecho.objects.get(roteiro=r, ordem=1)
        local0 = tz.localtime(t0.saida_dt).replace(tzinfo=None) if t0.saida_dt else None
        local1 = tz.localtime(t1.saida_dt).replace(tzinfo=None) if t1.saida_dt else None
        self.assertEqual(local0, datetime(2025, 2, 10, 8, 0))
        self.assertEqual(local1, datetime(2025, 2, 11, 9, 0))

    def test_calculo_automatico_ida_persistido(self):
        """Chegada da ida = saÃ­da + duraÃ§Ã£o (3h30); valor persistido no banco."""
        from eventos.models import RoteiroEventoTrecho
        r = RoteiroEvento.objects.create(
            evento=self.evento,
            origem_estado=self.estado,
            origem_cidade=self.cidade_a,
            saida_dt=datetime(2025, 3, 13, 14, 0),
            duracao_min=210,
        )
        RoteiroEventoDestino.objects.create(roteiro=r, estado=self.estado, cidade=self.cidade_b, ordem=0)
        data = {
            'origem_estado': self.estado.pk,
            'origem_cidade': self.cidade_a.pk,
            'destino_estado_0': self.estado.pk,
            'destino_cidade_0': self.cidade_b.pk,
            'saida_dt': '2025-03-13T14:00',
            'retorno_saida_dt': '2025-03-16T08:00',
            'observacoes': '',
            'trecho_0_saida_dt': '2025-03-13T14:00',
            'trecho_0_chegada_dt': '2025-03-13T17:30',
            'trecho_1_saida_dt': '2025-03-16T08:00',
            'trecho_1_chegada_dt': '2025-03-16T11:30',
        }
        response = self.client.post(
            reverse('eventos:guiado-etapa-2-editar', kwargs={'evento_id': self.evento.pk, 'pk': r.pk}),
            data,
        )
        self.assertEqual(response.status_code, 302)
        trechos = list(RoteiroEventoTrecho.objects.filter(roteiro=r).order_by('ordem'))
        self.assertEqual(len(trechos), 2)
        delta = (trechos[0].chegada_dt - trechos[0].saida_dt).total_seconds()
        self.assertEqual(delta, 210 * 60)

    def test_calculo_automatico_retorno_persistido(self):
        """Chegada do retorno = saÃ­da retorno + duraÃ§Ã£o (3h30); valor persistido no banco."""
        from eventos.models import RoteiroEventoTrecho
        r = RoteiroEvento.objects.create(
            evento=self.evento,
            origem_estado=self.estado,
            origem_cidade=self.cidade_a,
            saida_dt=datetime(2025, 3, 13, 14, 0),
            duracao_min=210,
        )
        RoteiroEventoDestino.objects.create(roteiro=r, estado=self.estado, cidade=self.cidade_b, ordem=0)
        data = {
            'origem_estado': self.estado.pk,
            'origem_cidade': self.cidade_a.pk,
            'destino_estado_0': self.estado.pk,
            'destino_cidade_0': self.cidade_b.pk,
            'saida_dt': '2025-03-13T14:00',
            'retorno_saida_dt': '2025-03-16T08:00',
            'observacoes': '',
            'trecho_0_saida_dt': '2025-03-13T14:00',
            'trecho_0_chegada_dt': '2025-03-13T17:30',
            'trecho_1_saida_dt': '2025-03-16T08:00',
            'trecho_1_chegada_dt': '2025-03-16T11:30',
        }
        response = self.client.post(
            reverse('eventos:guiado-etapa-2-editar', kwargs={'evento_id': self.evento.pk, 'pk': r.pk}),
            data,
        )
        self.assertEqual(response.status_code, 302)
        trechos = list(RoteiroEventoTrecho.objects.filter(roteiro=r).order_by('ordem'))
        self.assertEqual(len(trechos), 2)
        self.assertEqual(trechos[1].tipo, 'RETORNO')
        delta = (trechos[1].chegada_dt - trechos[1].saida_dt).total_seconds()
        self.assertEqual(delta, 210 * 60)

    def test_chegada_retorno_calculada(self):
        """Chegada do retorno Ã© preenchida a partir do trecho de retorno salvo."""
        data = {
            'origem_estado': self.estado.pk,
            'origem_cidade': self.cidade_a.pk,
            'destino_estado_0': self.estado.pk,
            'destino_cidade_0': self.cidade_b.pk,
            'saida_dt': '2025-01-02T08:00',
            'retorno_saida_dt': '2025-01-02T18:00',
            'observacoes': '',
            'trecho_0_saida_dt': '2025-01-02T08:00',
            'trecho_0_chegada_dt': '2025-01-02T09:00',
            'trecho_1_saida_dt': '2025-01-02T18:00',
            'trecho_1_chegada_dt': '2025-01-02T19:00',
        }
        self.client.post(
            reverse('eventos:guiado-etapa-2-cadastrar', kwargs={'evento_id': self.evento.pk}),
            data,
        )
        r = RoteiroEvento.objects.get(evento=self.evento)
        self.assertIsNotNone(r.retorno_chegada_dt)
        self.assertEqual((r.retorno_chegada_dt - r.retorno_saida_dt).total_seconds(), 60 * 60)

    def test_trechos_multiplos_destinos(self):
        """MÃºltiplos trechos na exibiÃ§Ã£o: ida (sede -> destino) e retorno (Ãºltimo destino -> sede)."""
        r = RoteiroEvento.objects.create(
            evento=self.evento,
            origem_estado=self.estado,
            origem_cidade=self.cidade_a,
            saida_dt=datetime(2025, 1, 2, 8, 0),
            duracao_min=60,
            retorno_saida_dt=datetime(2025, 1, 2, 18, 0),
        )
        RoteiroEventoDestino.objects.create(roteiro=r, estado=self.estado, cidade=self.cidade_b, ordem=0)
        response = self.client.get(reverse('eventos:guiado-etapa-2-editar', kwargs={'evento_id': self.evento.pk, 'pk': r.pk}))
        self.assertEqual(response.status_code, 200)
        trechos = response.context.get('trechos') or []
        self.assertGreaterEqual(len(trechos), 2)
        self.assertIn('Curitiba', str(trechos))
        self.assertIn('Londrina', str(trechos))
        self.assertTrue(any(t.get('tipo') == 'RETORNO' for t in trechos))

    def test_edicao_nao_sobrescreve_sede_com_config(self):
        """Na ediÃ§Ã£o, a sede exibida Ã© a salva no roteiro, nÃ£o a da configuraÃ§Ã£o atual."""
        from cadastros.models import ConfiguracaoSistema
        config = ConfiguracaoSistema.get_singleton()
        config.cidade_sede_padrao = self.cidade_a
        config.save(update_fields=['cidade_sede_padrao'])
        r = RoteiroEvento.objects.create(
            evento=self.evento,
            origem_estado=self.estado,
            origem_cidade=self.cidade_b,
            saida_dt=datetime(2025, 1, 2, 8, 0),
            duracao_min=60,
        )
        RoteiroEventoDestino.objects.create(roteiro=r, estado=self.estado, cidade=self.cidade_a, ordem=0)
        response = self.client.get(reverse('eventos:guiado-etapa-2-editar', kwargs={'evento_id': self.evento.pk, 'pk': r.pk}))
        self.assertEqual(response.status_code, 200)
        form = response.context['form']
        self.assertEqual(form.instance.origem_cidade_id, self.cidade_b.pk)

    def test_cidade_sede_selecionada_ao_abrir_cadastro(self):
        """Ao abrir cadastro novo, Cidade (Sede) deve vir prÃ©-preenchida e selecionada da configuraÃ§Ã£o."""
        from cadastros.models import ConfiguracaoSistema
        config = ConfiguracaoSistema.get_singleton()
        config.cidade_sede_padrao = self.cidade_a
        config.save(update_fields=['cidade_sede_padrao'])
        response = self.client.get(reverse('eventos:guiado-etapa-2-cadastrar', kwargs={'evento_id': self.evento.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, f'value="{self.cidade_a.pk}"')
        html = response.content.decode()
        self.assertIn(f'value="{self.cidade_a.pk}" selected', html)

    def test_cadastro_e_edicao_usam_mesmo_template_roteiro_form(self):
        """Cadastro novo e ediÃ§Ã£o de roteiro renderizam o mesmo template principal."""
        # Cadastro novo
        response_cad = self.client.get(
            reverse('eventos:guiado-etapa-2-cadastrar', kwargs={'evento_id': self.evento.pk})
        )
        self.assertEqual(response_cad.status_code, 200)
        template_names_cad = {t.name for t in response_cad.templates if t.name}
        self.assertIn('eventos/guiado/roteiro_form.html', template_names_cad)
        # EdiÃ§Ã£o
        r = RoteiroEvento.objects.create(
            evento=self.evento,
            origem_estado=self.estado,
            origem_cidade=self.cidade_a,
        )
        RoteiroEventoDestino.objects.create(roteiro=r, estado=self.estado, cidade=self.cidade_b, ordem=0)
        response_edit = self.client.get(
            reverse('eventos:guiado-etapa-2-editar', kwargs={'evento_id': self.evento.pk, 'pk': r.pk})
        )
        self.assertEqual(response_edit.status_code, 200)
        template_names_edit = {t.name for t in response_edit.templates if t.name}
        self.assertIn('eventos/guiado/roteiro_form.html', template_names_edit)

    def test_cadastro_novo_ja_exibe_bloco_trechos_sem_salvar(self):
        """Cadastro novo deve renderizar o bloco de trechos para o JS preencher imediatamente."""
        EventoDestino.objects.create(evento=self.evento, estado=self.estado, cidade=self.cidade_a, ordem=0)
        EventoDestino.objects.create(evento=self.evento, estado=self.estado, cidade=self.cidade_b, ordem=1)
        response = self.client.get(
            reverse('eventos:guiado-etapa-2-cadastrar', kwargs={'evento_id': self.evento.pk})
        )
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, '3) Trechos')
        self.assertContains(response, 'trechos-gerados-container')

    def test_cadastro_novo_persistir_trechos_conforme_formulario(self):
        """Cadastro novo deve salvar imediatamente os trechos de acordo com o que foi montado no formulÃ¡rio."""
        from eventos.models import RoteiroEventoTrecho
        data = {
            'origem_estado': self.estado.pk,
            'origem_cidade': self.cidade_a.pk,
            'destino_estado_0': self.estado.pk,
            'destino_cidade_0': self.cidade_b.pk,
            'observacoes': '',
            'trecho_0_saida_dt': '2025-02-10T09:00',
            'trecho_0_chegada_dt': '2025-02-10T10:00',
            'trecho_1_saida_dt': '2025-02-12T14:00',
            'trecho_1_chegada_dt': '2025-02-12T15:00',
        }
        response = self.client.post(
            reverse('eventos:guiado-etapa-2-cadastrar', kwargs={'evento_id': self.evento.pk}),
            data,
        )
        self.assertEqual(response.status_code, 302)
        # Salvar roteiro novo redireciona para a LISTA da Etapa 2, nÃ£o para editar.
        expected_url = reverse('eventos:guiado-etapa-2', kwargs={'evento_id': self.evento.pk})
        self.assertEqual(response.url, expected_url, 'Cadastro novo deve redirecionar para lista da Etapa 2')
        r = RoteiroEvento.objects.get(evento=self.evento)
        trechos = list(RoteiroEventoTrecho.objects.filter(roteiro=r).order_by('ordem'))
        self.assertEqual(len(trechos), 2)
        self.assertIsNotNone(trechos[0].saida_dt)
        self.assertIsNotNone(trechos[0].chegada_dt)
        self.assertIsNotNone(trechos[1].saida_dt)
        self.assertIsNotNone(trechos[1].chegada_dt)

    def test_cadastro_novo_trechos_json_preenchido_com_sede_e_destinos(self):
        """Cadastro novo com sede da config e 2 destinos do evento jÃ¡ envia trechos_list e trechos_json (3 trechos)."""
        from cadastros.models import ConfiguracaoSistema
        import json
        config = ConfiguracaoSistema.get_singleton()
        config.cidade_sede_padrao = self.cidade_a
        config.save(update_fields=['cidade_sede_padrao'])
        EventoDestino.objects.create(evento=self.evento, estado=self.estado, cidade=self.cidade_a, ordem=0)
        EventoDestino.objects.create(evento=self.evento, estado=self.estado, cidade=self.cidade_b, ordem=1)
        response = self.client.get(
            reverse('eventos:guiado-etapa-2-cadastrar', kwargs={'evento_id': self.evento.pk})
        )
        self.assertEqual(response.status_code, 200)
        trechos_list = response.context.get('trechos') or []
        trechos_json_raw = response.context.get('trechos_json') or '[]'
        trechos_initial = json.loads(trechos_json_raw) if isinstance(trechos_json_raw, str) else trechos_json_raw
        self.assertEqual(len(trechos_list), 3, '2 idas + 1 retorno')
        self.assertEqual(len(trechos_initial), 3)
        self.assertIn('initialTrechosData', response.content.decode())

    def test_script_adicionar_remover_destino_regenera_trechos(self):
        """PÃ¡gina de roteiro contÃ©m script que chama renderTrechos ao adicionar/remover destino."""
        response = self.client.get(
            reverse('eventos:guiado-etapa-2-cadastrar', kwargs={'evento_id': self.evento.pk})
        )
        self.assertEqual(response.status_code, 200)
        content = response.content.decode()
        self.assertIn('btn-adicionar-destino', content)
        self.assertIn('renderTrechos', content)
        self.assertIn('removerDestino', content)
        self.assertIn('trechos-gerados-container', content)

    def test_tempo_cru_adicional_persistidos_ao_salvar(self):
        """Salvar com tempo_cru e tempo_adicional persiste e tempo_total = cru + adicional."""
        from eventos.models import RoteiroEventoTrecho
        r = RoteiroEvento.objects.create(
            evento=self.evento,
            origem_estado=self.estado,
            origem_cidade=self.cidade_a,
        )
        RoteiroEventoDestino.objects.create(roteiro=r, estado=self.estado, cidade=self.cidade_b, ordem=0)
        data = {
            'origem_estado': self.estado.pk,
            'origem_cidade': self.cidade_a.pk,
            'destino_estado_0': self.estado.pk,
            'destino_cidade_0': self.cidade_b.pk,
            'observacoes': '',
            'trecho_0_saida_dt': '2025-02-10T09:00',
            'trecho_0_chegada_dt': '2025-02-10T10:45',
            'trecho_0_tempo_cru_estimado_min': '120',
            'trecho_0_tempo_adicional_min': '45',
            'trecho_1_saida_dt': '2025-02-12T14:00',
            'trecho_1_chegada_dt': '2025-02-12T15:30',
            'trecho_1_tempo_cru_estimado_min': '75',
            'trecho_1_tempo_adicional_min': '15',
        }
        response = self.client.post(
            reverse('eventos:guiado-etapa-2-editar', kwargs={'evento_id': self.evento.pk, 'pk': r.pk}),
            data,
        )
        self.assertEqual(response.status_code, 302)
        trechos = list(RoteiroEventoTrecho.objects.filter(roteiro=r).order_by('ordem'))
        self.assertEqual(len(trechos), 2)
        self.assertEqual(trechos[0].tempo_cru_estimado_min, 120)
        self.assertEqual(trechos[0].tempo_adicional_min, 45)
        self.assertEqual(trechos[0].duracao_estimada_min, 165)
        self.assertEqual(trechos[1].tempo_cru_estimado_min, 75)
        self.assertEqual(trechos[1].tempo_adicional_min, 15)
        self.assertEqual(trechos[1].duracao_estimada_min, 90)

    def test_tempo_total_igual_cru_mais_adicional(self):
        """Model: tempo_total_final_min = tempo_cru + tempo_adicional."""
        from eventos.models import RoteiroEventoTrecho
        r = RoteiroEvento.objects.create(
            evento=self.evento,
            origem_estado=self.estado,
            origem_cidade=self.cidade_a,
        )
        RoteiroEventoDestino.objects.create(roteiro=r, estado=self.estado, cidade=self.cidade_b, ordem=0)
        t = RoteiroEventoTrecho.objects.create(
            roteiro=r, ordem=0, tipo=RoteiroEventoTrecho.TIPO_IDA,
            origem_estado=self.estado, origem_cidade=self.cidade_a,
            destino_estado=self.estado, destino_cidade=self.cidade_b,
            tempo_cru_estimado_min=100, tempo_adicional_min=20,
        )
        self.assertEqual(t.tempo_total_final_min, 120)

    def test_tempo_adicional_aceita_zero(self):
        """Backend aceita e persiste tempo_adicional_min=0; total = cru + 0."""
        from eventos.models import RoteiroEventoTrecho
        r = RoteiroEvento.objects.create(
            evento=self.evento,
            origem_estado=self.estado,
            origem_cidade=self.cidade_a,
        )
        RoteiroEventoDestino.objects.create(roteiro=r, estado=self.estado, cidade=self.cidade_b, ordem=0)
        data = {
            'origem_estado': self.estado.pk,
            'origem_cidade': self.cidade_a.pk,
            'destino_estado_0': self.estado.pk,
            'destino_cidade_0': self.cidade_b.pk,
            'observacoes': '',
            'trecho_0_saida_dt': '2025-02-10T09:00',
            'trecho_0_chegada_dt': '2025-02-10T10:00',
            'trecho_0_tempo_cru_estimado_min': '90',
            'trecho_0_tempo_adicional_min': '0',
            'trecho_1_saida_dt': '2025-02-12T14:00',
            'trecho_1_chegada_dt': '2025-02-12T15:00',
            'trecho_1_tempo_cru_estimado_min': '60',
            'trecho_1_tempo_adicional_min': '0',
        }
        response = self.client.post(
            reverse('eventos:guiado-etapa-2-editar', kwargs={'evento_id': self.evento.pk, 'pk': r.pk}),
            data,
        )
        self.assertEqual(response.status_code, 302)
        trechos = list(RoteiroEventoTrecho.objects.filter(roteiro=r).order_by('ordem'))
        self.assertEqual(trechos[0].tempo_adicional_min, 0)
        self.assertEqual(trechos[0].duracao_estimada_min, 90)
        self.assertEqual(trechos[1].tempo_adicional_min, 0)
        self.assertEqual(trechos[1].duracao_estimada_min, 60)

    def test_tempo_adicional_negativo_clamped_para_zero(self):
        """Backend rejeita adicional negativo; valor Ã© clampeado para 0."""
        from eventos.models import RoteiroEventoTrecho
        r = RoteiroEvento.objects.create(
            evento=self.evento,
            origem_estado=self.estado,
            origem_cidade=self.cidade_a,
        )
        RoteiroEventoDestino.objects.create(roteiro=r, estado=self.estado, cidade=self.cidade_b, ordem=0)
        data = {
            'origem_estado': self.estado.pk,
            'origem_cidade': self.cidade_a.pk,
            'destino_estado_0': self.estado.pk,
            'destino_cidade_0': self.cidade_b.pk,
            'observacoes': '',
            'trecho_0_saida_dt': '2025-02-10T09:00',
            'trecho_0_chegada_dt': '2025-02-10T10:15',
            'trecho_0_tempo_cru_estimado_min': '60',
            'trecho_0_tempo_adicional_min': '-15',
            'trecho_1_saida_dt': '',
            'trecho_1_chegada_dt': '',
        }
        response = self.client.post(
            reverse('eventos:guiado-etapa-2-editar', kwargs={'evento_id': self.evento.pk, 'pk': r.pk}),
            data,
        )
        self.assertEqual(response.status_code, 302)
        trechos = list(RoteiroEventoTrecho.objects.filter(roteiro=r).order_by('ordem'))
        self.assertIsNotNone(trechos[0].tempo_adicional_min)
        self.assertGreaterEqual(trechos[0].tempo_adicional_min, 0, 'Adicional nunca fica negativo')

    def test_tempo_adicional_sem_restricao_absurda(self):
        """Adicional aceita valores fora da faixa 6..21 (ex.: 0, 30, 60, 120)."""
        from eventos.models import RoteiroEventoTrecho
        r = RoteiroEvento.objects.create(
            evento=self.evento,
            origem_estado=self.estado,
            origem_cidade=self.cidade_a,
        )
        RoteiroEventoDestino.objects.create(roteiro=r, estado=self.estado, cidade=self.cidade_b, ordem=0)
        data = {
            'origem_estado': self.estado.pk,
            'origem_cidade': self.cidade_a.pk,
            'destino_estado_0': self.estado.pk,
            'destino_cidade_0': self.cidade_b.pk,
            'observacoes': '',
            'trecho_0_saida_dt': '2025-02-10T09:00',
            'trecho_0_chegada_dt': '2025-02-10T11:30',
            'trecho_0_tempo_cru_estimado_min': '90',
            'trecho_0_tempo_adicional_min': '60',
            'trecho_1_saida_dt': '',
            'trecho_1_chegada_dt': '',
        }
        response = self.client.post(
            reverse('eventos:guiado-etapa-2-editar', kwargs={'evento_id': self.evento.pk, 'pk': r.pk}),
            data,
        )
        self.assertEqual(response.status_code, 302)
        trechos = list(RoteiroEventoTrecho.objects.filter(roteiro=r).order_by('ordem'))
        self.assertEqual(trechos[0].tempo_adicional_min, 60)
        self.assertEqual(trechos[0].duracao_estimada_min, 150)  # 90 + 60

    def test_script_renderiza_inputs_visiveis_por_trecho(self):
        """Script de roteiro gera inputs visÃ­veis (date/time) por trecho, nÃ£o sÃ³ hidden."""
        response = self.client.get(
            reverse('eventos:guiado-etapa-2-cadastrar', kwargs={'evento_id': self.evento.pk})
        )
        self.assertEqual(response.status_code, 200)
        content = response.content.decode()
        self.assertIn('type="date"', content, 'Script deve gerar input type=date')
        self.assertIn('type="time"', content, 'Script deve gerar input type=time')
        self.assertIn('_saida_date', content, 'Script deve gerar campo saÃ­da data por trecho')
        self.assertIn('_saida_time', content, 'Script deve gerar campo saÃ­da hora por trecho')
        self.assertIn('_chegada_date', content, 'Script deve gerar campo chegada data por trecho')
        self.assertIn('_chegada_time', content, 'Script deve gerar campo chegada hora por trecho')
        self.assertIn('data-trecho-ordem', content, 'Script deve marcar cada card com ordem do trecho')

    def test_script_tem_botoes_tempo_mais_menos(self):
        """Script de roteiro contÃ©m botÃµes +15 e -15 para tempo adicional."""
        response = self.client.get(
            reverse('eventos:guiado-etapa-2-cadastrar', kwargs={'evento_id': self.evento.pk})
        )
        self.assertEqual(response.status_code, 200)
        content = response.content.decode()
        self.assertIn('btn-tempo-menos', content)
        self.assertIn('btn-tempo-mais', content)
        self.assertIn('trecho-tempo-adicional', content)

    def test_cadastro_mostra_botao_estimar_km_tempo(self):
        """Cadastro de roteiro exibe o botÃ£o 'Estimar km/tempo' (mesma base funcional da ediÃ§Ã£o)."""
        EventoDestino.objects.create(evento=self.evento, estado=self.estado, cidade=self.cidade_a, ordem=0)
        response = self.client.get(
            reverse('eventos:guiado-etapa-2-cadastrar', kwargs={'evento_id': self.evento.pk})
        )
        self.assertEqual(response.status_code, 200)
        content = response.content.decode()
        self.assertIn('Estimar km/tempo', content)
        self.assertIn('btn-calcular-km', content)
        self.assertIn('urlTrechosEstimar', content)

    def test_cadastro_inclui_url_trechos_estimar_para_trecho_novo(self):
        """PÃ¡gina de cadastro/etapa-2 inclui urlTrechosEstimar para estimar trecho novo sem pk."""
        EventoDestino.objects.create(evento=self.evento, estado=self.estado, cidade=self.cidade_a, ordem=0)
        response = self.client.get(
            reverse('eventos:guiado-etapa-2-cadastrar', kwargs={'evento_id': self.evento.pk})
        )
        self.assertEqual(response.status_code, 200)
        content = response.content.decode()
        self.assertIn('urlTrechosEstimar', content)
        self.assertIn('trechos/estimar/', content)

    def test_cadastro_trechos_json_tem_origem_destino_cidade_id(self):
        """trechos_json no cadastro inclui origem_cidade_id e destino_cidade_id por trecho."""
        from cadastros.models import ConfiguracaoSistema
        import json
        config = ConfiguracaoSistema.get_singleton()
        config.cidade_sede_padrao = self.cidade_a
        config.save(update_fields=['cidade_sede_padrao'])
        EventoDestino.objects.create(evento=self.evento, estado=self.estado, cidade=self.cidade_a, ordem=0)
        EventoDestino.objects.create(evento=self.evento, estado=self.estado, cidade=self.cidade_b, ordem=1)
        response = self.client.get(
            reverse('eventos:guiado-etapa-2-cadastrar', kwargs={'evento_id': self.evento.pk})
        )
        self.assertEqual(response.status_code, 200)
        trechos_json_raw = response.context.get('trechos_json') or '[]'
        trechos_initial = json.loads(trechos_json_raw) if isinstance(trechos_json_raw, str) else trechos_json_raw
        self.assertGreaterEqual(len(trechos_initial), 2)
        first = trechos_initial[0]
        self.assertIn('origem_cidade_id', first)
        self.assertIn('destino_cidade_id', first)

    def test_edicao_abre_com_trechos_persistidos(self):
        """EdiÃ§Ã£o de roteiro abre com trechos e horÃ¡rios jÃ¡ salvos no contexto."""
        r = RoteiroEvento.objects.create(
            evento=self.evento,
            origem_estado=self.estado,
            origem_cidade=self.cidade_a,
            saida_dt=datetime(2025, 1, 2, 8, 0),
            duracao_min=60,
        )
        RoteiroEventoDestino.objects.create(roteiro=r, estado=self.estado, cidade=self.cidade_b, ordem=0)
        from eventos.models import RoteiroEventoTrecho
        RoteiroEventoTrecho.objects.create(
            roteiro=r, ordem=0, tipo=RoteiroEventoTrecho.TIPO_IDA,
            origem_estado=self.estado, origem_cidade=self.cidade_a,
            destino_estado=self.estado, destino_cidade=self.cidade_b,
            saida_dt=datetime(2025, 1, 2, 8, 0), chegada_dt=datetime(2025, 1, 2, 9, 0),
        )
        RoteiroEventoTrecho.objects.create(
            roteiro=r, ordem=1, tipo=RoteiroEventoTrecho.TIPO_RETORNO,
            origem_estado=self.estado, origem_cidade=self.cidade_b,
            destino_estado=self.estado, destino_cidade=self.cidade_a,
            saida_dt=datetime(2025, 1, 2, 18, 0), chegada_dt=datetime(2025, 1, 2, 19, 0),
        )
        response = self.client.get(
            reverse('eventos:guiado-etapa-2-editar', kwargs={'evento_id': self.evento.pk, 'pk': r.pk})
        )
        self.assertEqual(response.status_code, 200)
        trechos = response.context.get('trechos') or []
        self.assertEqual(len(trechos), 2)
        self.assertIn('2025-01-02', response.context.get('trechos_json', ''))

    def test_trechos_persistidos_ao_salvar_edicao(self):
        """Ao salvar a ediÃ§Ã£o com horÃ¡rios por trecho, os trechos devem ser persistidos no banco."""
        from eventos.models import RoteiroEventoTrecho
        r = RoteiroEvento.objects.create(
            evento=self.evento,
            origem_estado=self.estado,
            origem_cidade=self.cidade_a,
            saida_dt=datetime(2025, 1, 2, 8, 0),
            duracao_min=60,
        )
        RoteiroEventoDestino.objects.create(roteiro=r, estado=self.estado, cidade=self.cidade_b, ordem=0)
        self.assertEqual(RoteiroEventoTrecho.objects.filter(roteiro=r).count(), 0)
        data = {
            'origem_estado': self.estado.pk,
            'origem_cidade': self.cidade_a.pk,
            'destino_estado_0': self.estado.pk,
            'destino_cidade_0': self.cidade_b.pk,
            'saida_dt': '2025-01-02T08:00',
            'retorno_saida_dt': '',
            'observacoes': '',
            'trecho_0_saida_dt': '2025-01-02T08:00',
            'trecho_0_chegada_dt': '2025-01-02T09:00',
            'trecho_1_saida_dt': '2025-01-02T18:00',
            'trecho_1_chegada_dt': '2025-01-02T19:00',
        }
        response = self.client.post(
            reverse('eventos:guiado-etapa-2-editar', kwargs={'evento_id': self.evento.pk, 'pk': r.pk}),
            data,
        )
        self.assertEqual(response.status_code, 302)
        r.refresh_from_db()
        trechos = list(RoteiroEventoTrecho.objects.filter(roteiro=r).order_by('ordem'))
        self.assertEqual(len(trechos), 2)
        self.assertIsNotNone(trechos[0].saida_dt)
        self.assertIsNotNone(trechos[0].chegada_dt)
        self.assertIsNotNone(trechos[1].saida_dt)
        self.assertIsNotNone(trechos[1].chegada_dt)


class EventoEtapa1RefatoradoTest(TestCase):
    """Testes da Etapa 1 refatorada: mÃºltiplos tipos, tÃ­tulo automÃ¡tico, data Ãºnica, destinos, descriÃ§Ã£o."""

    def setUp(self):
        self.user = User.objects.create_user(username='u', password='p')
        self.client = Client()
        self.client.login(username='u', password='p')
        self.estado = Estado.objects.create(nome='ParanÃ¡', sigla='PR', codigo_ibge='41')
        self.cidade_a = Cidade.objects.create(nome='Curitiba', estado=self.estado, codigo_ibge='4106902')
        self.cidade_b = Cidade.objects.create(nome='Londrina', estado=self.estado, codigo_ibge='4113700')
        self.tipos = list(TipoDemandaEvento.objects.filter(ativo=True).order_by('ordem')[:3])

    def test_etapa_1_multiplos_tipos_demanda(self):
        self.assertGreaterEqual(len(self.tipos), 2, 'Precisa de ao menos 2 tipos (migraÃ§Ã£o 0004).')
        ev = Evento.objects.create(titulo='', data_inicio=date(2025, 3, 12), data_fim=date(2025, 3, 12), status=Evento.STATUS_RASCUNHO)
        data = {
            'tipos_demanda': [self.tipos[0].pk, self.tipos[1].pk],
            'data_unica': True,
            'data_inicio': '2025-03-12',
            'data_fim': '2025-03-12',
            'descricao': '',
            'tem_convite_ou_oficio_evento': False,
            'destino_estado_0': self.estado.pk,
            'destino_cidade_0': self.cidade_a.pk,
        }
        response = self.client.post(reverse('eventos:guiado-etapa-1', kwargs={'pk': ev.pk}), data)
        self.assertEqual(response.status_code, 302)
        ev.refresh_from_db()
        self.assertEqual(ev.tipos_demanda.count(), 2)

    def test_etapa_1_autosave_persiste_dados_sem_redirect(self):
        tipo = self.tipos[0] if self.tipos else None
        self.assertIsNotNone(tipo)
        ev = Evento.objects.create(titulo='', data_inicio=date(2025, 3, 12), data_fim=date(2025, 3, 12), status=Evento.STATUS_RASCUNHO)

        response = self.client.post(
            reverse('eventos:guiado-etapa-1', kwargs={'pk': ev.pk}),
            {
                'autosave': '1',
                'tipos_demanda': [tipo.pk],
                'data_unica': 'on',
                'data_inicio': '2025-03-20',
                'descricao': '',
                'tem_convite_ou_oficio_evento': 'on',
                'destino_estado_0': self.estado.pk,
                'destino_cidade_0': self.cidade_a.pk,
            },
            HTTP_X_REQUESTED_WITH='XMLHttpRequest',
        )

        self.assertEqual(response.status_code, 200)
        self.assertTrue(response.json().get('ok'))

        ev.refresh_from_db()
        self.assertEqual(ev.data_inicio, date(2025, 3, 20))
        self.assertEqual(ev.data_fim, date(2025, 3, 20))
        self.assertTrue(ev.data_unica)
        self.assertFalse(ev.tem_convite_ou_oficio_evento)
        self.assertEqual(ev.destinos.count(), 1)

    def test_etapa_1_renderiza_hook_de_autosave_e_topo_reativo(self):
        ev = Evento.objects.create(titulo='', data_inicio=date(2025, 3, 12), data_fim=date(2025, 3, 12), status=Evento.STATUS_RASCUNHO)
        response = self.client.get(reverse('eventos:guiado-etapa-1', kwargs={'pk': ev.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'js/oficio_wizard.js')
        self.assertContains(response, 'id="guiado-etapa1-heading"')
        self.assertContains(response, 'id="guiado-etapa1-periodo"')
        self.assertContains(response, 'id="guiado-etapa1-tipos"')
        self.assertContains(response, 'id="guiado-etapa1-destinos"')
        self.assertContains(response, 'id="evento-etapa1-autosave-status"')
        self.assertContains(response, 'data-autosave-link="1"')
        self.assertContains(response, 'data-autosave-navigate=')

    def test_etapa_1_ignora_upload_de_convite(self):
        tipo = self.tipos[0] if self.tipos else None
        self.assertIsNotNone(tipo)
        ev = Evento.objects.create(titulo='', data_inicio=date(2025, 3, 12), data_fim=date(2025, 3, 12), status=Evento.STATUS_RASCUNHO)
        arquivo = SimpleUploadedFile('convite.pdf', b'%PDF-1.4\n%autosave\n', content_type='application/pdf')
        response = self.client.post(
            reverse('eventos:guiado-etapa-1', kwargs={'pk': ev.pk}),
            {
                'autosave': '1',
                'tipos_demanda': [tipo.pk],
                'data_unica': 'on',
                'data_inicio': '2025-03-20',
                'descricao': '',
                'tem_convite_ou_oficio_evento': 'on',
                'destino_estado_0': self.estado.pk,
                'destino_cidade_0': self.cidade_a.pk,
                'convite_documentos': arquivo,
            },
            HTTP_X_REQUESTED_WITH='XMLHttpRequest',
        )
        self.assertEqual(response.status_code, 200)
        self.assertTrue(response.json().get('ok'))
        ev.refresh_from_db()
        self.assertFalse(ev.tem_convite_ou_oficio_evento)
        self.assertEqual(ev.anexos_solicitante.count(), 0)

    def test_etapa_1_titulo_gerado_automaticamente(self):
        tipo = self.tipos[0] if self.tipos else None
        self.assertIsNotNone(tipo)
        ev = Evento.objects.create(titulo='', data_inicio=date(2026, 3, 12), data_fim=date(2026, 3, 12), status=Evento.STATUS_RASCUNHO)
        data = {
            'tipos_demanda': [tipo.pk],
            'data_unica': True,
            'data_inicio': '2026-03-12',
            'data_fim': '2026-03-12',
            'descricao': 'Desc',
            'tem_convite_ou_oficio_evento': False,
            'destino_estado_0': self.estado.pk,
            'destino_cidade_0': self.cidade_a.pk,
        }
        self.client.post(reverse('eventos:guiado-etapa-1', kwargs={'pk': ev.pk}), data)
        ev.refresh_from_db()
        self.assertTrue(ev.titulo)
        self.assertIn('CURITIBA', ev.titulo)
        self.assertIn('12/03/2026', ev.titulo)

    def test_etapa_1_data_unica_ignora_data_fim(self):
        tipo = self.tipos[0] if self.tipos else None
        self.assertIsNotNone(tipo)
        ev = Evento.objects.create(titulo='', data_inicio=date(2025, 6, 10), data_fim=date(2025, 6, 15), status=Evento.STATUS_RASCUNHO)
        data = {
            'tipos_demanda': [tipo.pk],
            'data_unica': True,
            'data_inicio': '2025-06-10',
            'data_fim': '2025-06-10',
            'descricao': 'X',
            'tem_convite_ou_oficio_evento': False,
            'destino_estado_0': self.estado.pk,
            'destino_cidade_0': self.cidade_a.pk,
        }
        response = self.client.post(reverse('eventos:guiado-etapa-1', kwargs={'pk': ev.pk}), data)
        self.assertEqual(response.status_code, 302)
        ev.refresh_from_db()
        self.assertEqual(ev.data_fim, date(2025, 6, 10))

    def test_etapa_1_data_unica_sem_enviar_data_fim_backend_preenche(self):
        """data_unica=True: nÃ£o enviar data_fim; backend deve preencher data_fim = data_inicio."""
        tipo = self.tipos[0] if self.tipos else None
        self.assertIsNotNone(tipo)
        ev = Evento.objects.create(titulo='', data_inicio=date(2025, 7, 20), data_fim=date(2025, 7, 25), status=Evento.STATUS_RASCUNHO)
        data = {
            'tipos_demanda': [tipo.pk],
            'data_unica': True,
            'data_inicio': '2025-07-20',
            'descricao': 'X',
            'tem_convite_ou_oficio_evento': False,
            'destino_estado_0': self.estado.pk,
            'destino_cidade_0': self.cidade_a.pk,
        }
        response = self.client.post(reverse('eventos:guiado-etapa-1', kwargs={'pk': ev.pk}), data)
        self.assertEqual(response.status_code, 302)
        ev.refresh_from_db()
        self.assertEqual(ev.data_fim, date(2025, 7, 20))

    def test_etapa_1_sem_outros_descricao_nao_obrigatoria(self):
        """Sem tipo OUTROS selecionado: descriÃ§Ã£o nÃ£o Ã© exigida e Ã© limpa ao salvar."""
        tipo = next((t for t in (self.tipos or []) if not t.is_outros), None)
        if not tipo:
            self.skipTest('Precisa de um tipo de demanda que nÃ£o seja OUTROS.')
        ev = Evento.objects.create(titulo='', data_inicio=date(2025, 2, 1), data_fim=date(2025, 2, 1), status=Evento.STATUS_RASCUNHO, descricao='Texto antigo')
        data = {
            'tipos_demanda': [tipo.pk],
            'data_unica': True,
            'data_inicio': '2025-02-01',
            'data_fim': '2025-02-01',
            'descricao': '',
            'tem_convite_ou_oficio_evento': False,
            'destino_estado_0': self.estado.pk,
            'destino_cidade_0': self.cidade_a.pk,
        }
        response = self.client.post(reverse('eventos:guiado-etapa-1', kwargs={'pk': ev.pk}), data)
        self.assertEqual(response.status_code, 302)
        ev.refresh_from_db()
        self.assertEqual(ev.descricao, '')

    def test_etapa_1_multiplos_destinos(self):
        tipo = self.tipos[0] if self.tipos else None
        self.assertIsNotNone(tipo)
        ev = Evento.objects.create(titulo='', data_inicio=date(2025, 4, 1), data_fim=date(2025, 4, 3), status=Evento.STATUS_RASCUNHO)
        data = {
            'tipos_demanda': [tipo.pk],
            'data_unica': False,
            'data_inicio': '2025-04-01',
            'data_fim': '2025-04-03',
            'descricao': 'X',
            'tem_convite_ou_oficio_evento': False,
            'destino_estado_0': self.estado.pk,
            'destino_cidade_0': self.cidade_a.pk,
            'destino_estado_1': self.estado.pk,
            'destino_cidade_1': self.cidade_b.pk,
        }
        response = self.client.post(reverse('eventos:guiado-etapa-1', kwargs={'pk': ev.pk}), data)
        self.assertEqual(response.status_code, 302)
        ev.refresh_from_db()
        self.assertEqual(ev.destinos.count(), 2)

    def test_etapa_1_descricao_obrigatoria_quando_outros(self):
        tipo_outros = TipoDemandaEvento.objects.filter(is_outros=True).first()
        if not tipo_outros:
            self.skipTest('Nenhum tipo "Outros" cadastrado.')
        ev = Evento.objects.create(titulo='', data_inicio=date(2025, 1, 1), data_fim=date(2025, 1, 1), status=Evento.STATUS_RASCUNHO)
        data = {
            'tipos_demanda': [tipo_outros.pk],
            'data_unica': True,
            'data_inicio': '2025-01-01',
            'data_fim': '2025-01-01',
            'descricao': '',
            'tem_convite_ou_oficio_evento': False,
            'destino_estado_0': self.estado.pk,
            'destino_cidade_0': self.cidade_a.pk,
        }
        response = self.client.post(reverse('eventos:guiado-etapa-1', kwargs={'pk': ev.pk}), data)
        self.assertEqual(response.status_code, 200)
        self.assertIn('descricao', response.context['form'].errors)

    def test_etapa_1_formulario_abre_com_um_destino(self):
        """Ao abrir a Etapa 1 deve existir pelo menos 1 bloco de destino visÃ­vel."""
        ev = Evento.objects.create(titulo='', data_inicio=date(2025, 1, 1), data_fim=date(2025, 1, 1), status=Evento.STATUS_RASCUNHO)
        response = self.client.get(reverse('eventos:guiado-etapa-1', kwargs={'pk': ev.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'destino_estado_0')
        self.assertContains(response, 'destino_cidade_0')
        self.assertContains(response, 'destinos-container')

    def test_reabrir_etapa_1_apos_salvar_mantem_tipos_destinos_datas(self):
        """ApÃ³s salvar a Etapa 1, reabrir a tela deve exibir os dados persistidos."""
        tipo = self.tipos[0] if self.tipos else None
        self.assertIsNotNone(tipo)
        ev = Evento.objects.create(titulo='', data_inicio=date(2025, 5, 10), data_fim=date(2025, 5, 12), status=Evento.STATUS_RASCUNHO)
        # data_unica=False: nÃ£o enviar a chave no POST (checkbox desmarcado nÃ£o envia nada)
        data = {
            'tipos_demanda': [tipo.pk],
            'data_inicio': '2025-05-10',
            'data_fim': '2025-05-12',
            'descricao': '',
            'tem_convite_ou_oficio_evento': False,
            'destino_estado_0': self.estado.pk,
            'destino_cidade_0': self.cidade_a.pk,
            'destino_estado_1': self.estado.pk,
            'destino_cidade_1': self.cidade_b.pk,
        }
        self.client.post(reverse('eventos:guiado-etapa-1', kwargs={'pk': ev.pk}), data)
        ev.refresh_from_db()
        self.assertEqual(ev.tipos_demanda.count(), 1)
        self.assertEqual(ev.destinos.count(), 2)
        # Datas persistidas: intervalo 10 a 12/05/2025 (data_unica=False; nÃ£o enviar chave data_unica no POST)
        self.assertEqual(ev.data_inicio, date(2025, 5, 10))
        self.assertEqual(ev.data_fim, date(2025, 5, 12))
        self.assertFalse(ev.data_unica)
        response = self.client.get(reverse('eventos:guiado-etapa-1', kwargs={'pk': ev.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertEqual(len(response.context['destinos_atuais']), 2)
        self.assertIn(tipo.pk, response.context['selected_tipos_pks'])
        # Reabrir: formulÃ¡rio deve exibir as datas salvas
        self.assertContains(response, '2025-05-10')
        self.assertContains(response, '2025-05-12')

    def test_etapa_1_datas_persistidas_data_unica_true(self):
        """Com data_unica=True: data_inicio e data_fim persistidas; data_fim = data_inicio."""
        tipo = self.tipos[0] if self.tipos else None
        self.assertIsNotNone(tipo)
        ev = Evento.objects.create(titulo='', data_inicio=date(2025, 1, 15), data_fim=date(2025, 1, 20), status=Evento.STATUS_RASCUNHO)
        data = {
            'tipos_demanda': [tipo.pk],
            'data_unica': True,
            'data_inicio': '2025-08-01',
            'data_fim': '2025-08-01',
            'descricao': '',
            'tem_convite_ou_oficio_evento': False,
            'destino_estado_0': self.estado.pk,
            'destino_cidade_0': self.cidade_a.pk,
        }
        response = self.client.post(reverse('eventos:guiado-etapa-1', kwargs={'pk': ev.pk}), data)
        self.assertEqual(response.status_code, 302)
        ev.refresh_from_db()
        self.assertEqual(ev.data_inicio, date(2025, 8, 1))
        self.assertEqual(ev.data_fim, date(2025, 8, 1))
        self.assertTrue(ev.data_unica)

    def test_etapa_1_datas_persistidas_data_unica_false(self):
        """Com data_unica=False: data_inicio e data_fim persistidas com valores distintos. NÃ£o enviar 'data_unica' no POST."""
        tipo = self.tipos[0] if self.tipos else None
        self.assertIsNotNone(tipo)
        ev = Evento.objects.create(titulo='', data_inicio=date(2025, 1, 1), data_fim=date(2025, 1, 1), status=Evento.STATUS_RASCUNHO)
        data = {
            'tipos_demanda': [tipo.pk],
            'data_inicio': '2025-09-10',
            'data_fim': '2025-09-15',
            'descricao': '',
            'tem_convite_ou_oficio_evento': False,
            'destino_estado_0': self.estado.pk,
            'destino_cidade_0': self.cidade_a.pk,
        }
        response = self.client.post(reverse('eventos:guiado-etapa-1', kwargs={'pk': ev.pk}), data)
        self.assertEqual(response.status_code, 302)
        ev.refresh_from_db()
        self.assertEqual(ev.data_inicio, date(2025, 9, 10))
        self.assertEqual(ev.data_fim, date(2025, 9, 15))
        self.assertFalse(ev.data_unica)

    def test_etapa_1_reabrir_mostra_datas_salvas(self):
        """Reabrir a Etapa 1 apÃ³s salvar deve exibir data_inicio e data_fim nos inputs."""
        tipo = self.tipos[0] if self.tipos else None
        self.assertIsNotNone(tipo)
        ev = Evento.objects.create(titulo='', data_inicio=date(2025, 11, 1), data_fim=date(2025, 11, 5), status=Evento.STATUS_RASCUNHO)
        data = {
            'tipos_demanda': [tipo.pk],
            'data_inicio': '2025-11-01',
            'data_fim': '2025-11-05',
            'descricao': '',
            'tem_convite_ou_oficio_evento': False,
            'destino_estado_0': self.estado.pk,
            'destino_cidade_0': self.cidade_a.pk,
        }
        self.client.post(reverse('eventos:guiado-etapa-1', kwargs={'pk': ev.pk}), data)
        response = self.client.get(reverse('eventos:guiado-etapa-1', kwargs={'pk': ev.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'value="2025-11-01"')
        self.assertContains(response, 'value="2025-11-05"')

    def test_detalhe_evento_mostra_datas_corretas(self):
        """Rota de detalhe redireciona para ediÃ§Ã£o e mantÃ©m datas persistidas na Etapa 1."""
        tipo = self.tipos[0] if self.tipos else None
        self.assertIsNotNone(tipo)
        ev = Evento.objects.create(titulo='', data_inicio=date(2025, 12, 1), data_fim=date(2025, 12, 10), status=Evento.STATUS_RASCUNHO)
        ev.tipos_demanda.add(tipo)
        EventoDestino.objects.create(evento=ev, estado=self.estado, cidade=self.cidade_a, ordem=0)
        data = {
            'tipos_demanda': [tipo.pk],
            'data_inicio': '2025-12-01',
            'data_fim': '2025-12-10',
            'descricao': '',
            'tem_convite_ou_oficio_evento': False,
            'destino_estado_0': self.estado.pk,
            'destino_cidade_0': self.cidade_a.pk,
        }
        self.client.post(reverse('eventos:guiado-etapa-1', kwargs={'pk': ev.pk}), data)
        response = self.client.get(reverse('eventos:detalhe', kwargs={'pk': ev.pk}))
        self.assertEqual(response.status_code, 302)
        self.assertEqual(response.url, reverse('eventos:editar', kwargs={'pk': ev.pk}))

        etapa1 = self.client.get(reverse('eventos:guiado-etapa-1', kwargs={'pk': ev.pk}))
        self.assertEqual(etapa1.status_code, 200)
        self.assertContains(etapa1, 'value="2025-12-01"')
        self.assertContains(etapa1, 'value="2025-12-10"')

    def test_etapa_1_nao_pode_salvar_sem_destino(self):
        """Envio sem nenhum destino vÃ¡lido deve rejeitar com erro."""
        tipo = self.tipos[0] if self.tipos else None
        self.assertIsNotNone(tipo)
        ev = Evento.objects.create(titulo='', data_inicio=date(2025, 1, 1), data_fim=date(2025, 1, 1), status=Evento.STATUS_RASCUNHO)
        data = {
            'tipos_demanda': [tipo.pk],
            'data_unica': True,
            'data_inicio': '2025-01-01',
            'data_fim': '2025-01-01',
            'descricao': '',
            'tem_convite_ou_oficio_evento': False,
        }
        response = self.client.post(reverse('eventos:guiado-etapa-1', kwargs={'pk': ev.pk}), data)
        self.assertEqual(response.status_code, 200)
        self.assertIn('__all__', response.context['form'].errors)
        self.assertIn('destino', str(response.context['form'].errors['__all__']).lower())


class TipoDemandaEventoCRUDTest(TestCase):
    """CRUD de tipos de demanda e bloqueio de exclusÃ£o quando em uso."""

    def setUp(self):
        self.user = User.objects.create_user(username='u', password='p')
        self.client = Client()
        self.client.login(username='u', password='p')

    def test_lista_tipos_demanda(self):
        response = self.client.get(reverse('eventos:tipos-demanda-lista'))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Tipos de demanda')

    def test_cadastrar_tipo_demanda(self):
        data = {'nome': 'NOVO TIPO', 'descricao_padrao': 'Desc padrÃ£o', 'ordem': 50, 'ativo': True, 'is_outros': False}
        response = self.client.post(reverse('eventos:tipos-demanda-cadastrar'), data)
        self.assertEqual(response.status_code, 302)
        self.assertTrue(TipoDemandaEvento.objects.filter(nome='NOVO TIPO').exists())

    def test_excluir_tipo_em_uso_bloqueado(self):
        tipo = TipoDemandaEvento.objects.filter(ativo=True).first()
        self.assertIsNotNone(tipo)
        ev = Evento.objects.create(titulo='E', data_inicio=date(2025, 1, 1), data_fim=date(2025, 1, 1), status=Evento.STATUS_RASCUNHO)
        ev.tipos_demanda.add(tipo)
        response = self.client.post(reverse('eventos:tipos-demanda-excluir', kwargs={'pk': tipo.pk}), {})
        self.assertEqual(response.status_code, 302)
        self.assertTrue(TipoDemandaEvento.objects.filter(pk=tipo.pk).exists())

    def test_excluir_tipo_quando_nao_em_uso_funciona(self):
        """Excluir tipo de demanda quando nÃ£o estÃ¡ em uso deve remover do banco."""
        tipo = TipoDemandaEvento.objects.create(nome='TIPO TESTE EXCLUSAO', ordem=999, ativo=True, is_outros=False)
        self.assertTrue(TipoDemandaEvento.objects.filter(pk=tipo.pk).exists())
        response = self.client.post(reverse('eventos:tipos-demanda-excluir', kwargs={'pk': tipo.pk}), {})
        self.assertEqual(response.status_code, 302)
        self.assertFalse(TipoDemandaEvento.objects.filter(pk=tipo.pk).exists())


class EstimativaLocalServiceTest(TestCase):
    """Testes do serviÃ§o de estimativa local (tempo de viagem vs buffer, corredores)."""

    def test_minutos_para_hhmm(self):
        from eventos.services.estimativa_local import minutos_para_hhmm
        self.assertEqual(minutos_para_hhmm(340), '05:40')
        self.assertEqual(minutos_para_hhmm(65), '01:05')
        self.assertEqual(minutos_para_hhmm(0), '00:00')
        self.assertEqual(minutos_para_hhmm(None), '')

    def test_estimativa_entre_coordenadas_retorna_ok(self):
        from eventos.services.estimativa_local import estimar_distancia_duracao
        out = estimar_distancia_duracao(
            origem_lat=-25.43, origem_lon=-49.27,
            destino_lat=-23.42, destino_lon=-51.94,
        )
        self.assertTrue(out['ok'])
        self.assertIsNotNone(out['distancia_km'])
        self.assertIsNotNone(out['duracao_estimada_min'])
        self.assertIsNotNone(out.get('tempo_viagem_estimado_min'))
        self.assertIsNotNone(out.get('buffer_operacional_sugerido_min'))
        self.assertEqual(out['rota_fonte'], 'ESTIMATIVA_LOCAL')

    def test_tempo_viagem_e_buffer_separados(self):
        """Retorno inclui tempo_viagem_estimado_min (comparÃ¡vel ao Google) e buffer_operacional_sugerido_min."""
        from eventos.services.estimativa_local import estimar_distancia_duracao
        out = estimar_distancia_duracao(-25.43, -49.27, -23.42, -51.94)
        self.assertTrue(out['ok'])
        self.assertIn('tempo_viagem_estimado_min', out)
        self.assertIn('tempo_viagem_estimado_hhmm', out)
        self.assertIn('buffer_operacional_sugerido_min', out)
        tv = out['tempo_viagem_estimado_min']
        buf = out['buffer_operacional_sugerido_min']
        self.assertEqual(out['duracao_estimada_min'], tv + buf)

    def test_duracao_igual_tempo_viagem_mais_buffer(self):
        """duracao_estimada_min = tempo_viagem_estimado_min + buffer_operacional_sugerido_min."""
        from eventos.services.estimativa_local import estimar_distancia_duracao
        out = estimar_distancia_duracao(-25.43, -49.27, -23.42, -51.94)
        self.assertTrue(out['ok'])
        tv = out.get('tempo_viagem_estimado_min', 0) or 0
        buf = out.get('buffer_operacional_sugerido_min', 0) or 0
        self.assertEqual(out['duracao_estimada_min'], tv + buf)

    def test_compatibilidade_tempo_cru_e_adicional(self):
        """Compatibilidade: tempo_cru_estimado_min = tempo_viagem; tempo_adicional_sugerido_min = buffer."""
        from eventos.services.estimativa_local import estimar_distancia_duracao
        out = estimar_distancia_duracao(-25.43, -49.27, -23.42, -51.94)
        self.assertTrue(out['ok'])
        self.assertEqual(out['tempo_cru_estimado_min'], out['tempo_viagem_estimado_min'])
        self.assertEqual(out['tempo_adicional_sugerido_min'], out['buffer_operacional_sugerido_min'])

    def test_classificar_corredor_litoral_curto(self):
        """Curitiba -> Pontal do ParanÃ¡: LITORAL_CURTO com buffer definido pela distÃ¢ncia."""
        from eventos.services.estimativa_local import (
            estimar_distancia_duracao,
            CORREDOR_LITORAL_CURTO,
            FATOR_CORREDOR,
            sugerir_buffer_operacional,
        )
        # Pontal do ParanÃ¡ ~ -25.67, -48.51
        out = estimar_distancia_duracao(
            origem_lat=-25.43, origem_lon=-49.27,
            destino_lat=-25.67, destino_lon=-48.51,
        )
        self.assertTrue(out['ok'])
        self.assertEqual(out['corredor'], CORREDOR_LITORAL_CURTO)
        self.assertEqual(
            out['buffer_operacional_sugerido_min'],
            sugerir_buffer_operacional(out['corredor_macro'], out['distancia_rodoviaria_km']),
        )
        self.assertEqual(FATOR_CORREDOR[CORREDOR_LITORAL_CURTO], 0.86)

    def test_classificar_corredor_campos_gerais(self):
        """Curitiba -> Ponta Grossa: CAMPOS_GERAIS_CURTO com buffer definido pela distÃ¢ncia."""
        from eventos.services.estimativa_local import (
            estimar_distancia_duracao,
            CORREDOR_CAMPOS_GERAIS_CURTO,
            sugerir_buffer_operacional,
        )
        out = estimar_distancia_duracao(
            origem_lat=-25.43, origem_lon=-49.27,
            destino_lat=-25.09, destino_lon=-50.16,
        )
        self.assertTrue(out['ok'])
        self.assertEqual(out['corredor'], CORREDOR_CAMPOS_GERAIS_CURTO)
        self.assertEqual(
            out['buffer_operacional_sugerido_min'],
            sugerir_buffer_operacional(out['corredor_macro'], out['distancia_rodoviaria_km']),
        )

    def test_classificar_corredor_norte_noroeste(self):
        """Curitiba -> MaringÃ¡: NORTE_NOROESTE com buffer definido pela distÃ¢ncia."""
        from eventos.services.estimativa_local import (
            estimar_distancia_duracao,
            CORREDOR_NORTE_NOROESTE,
            sugerir_buffer_operacional,
        )
        out = estimar_distancia_duracao(
            origem_lat=-25.43, origem_lon=-49.27,
            destino_lat=-23.42, destino_lon=-51.93,
        )
        self.assertTrue(out['ok'])
        self.assertEqual(out['corredor'], CORREDOR_NORTE_NOROESTE)
        self.assertEqual(
            out['buffer_operacional_sugerido_min'],
            sugerir_buffer_operacional(out['corredor_macro'], out['distancia_rodoviaria_km']),
        )

    def test_classificar_corredor_oeste_br277(self):
        """Curitiba -> Cascavel: OESTE_BR277 com buffer definido pela distÃ¢ncia."""
        from eventos.services.estimativa_local import (
            estimar_distancia_duracao,
            CORREDOR_OESTE_BR277,
            sugerir_buffer_operacional,
        )
        out = estimar_distancia_duracao(
            origem_lat=-25.43, origem_lon=-49.27,
            destino_lat=-24.96, destino_lon=-53.45,
        )
        self.assertTrue(out['ok'])
        self.assertEqual(out['corredor'], CORREDOR_OESTE_BR277)
        self.assertEqual(
            out['buffer_operacional_sugerido_min'],
            sugerir_buffer_operacional(out['corredor_macro'], out['distancia_rodoviaria_km']),
        )

    def test_fator_corredor_aplicado(self):
        """tempo_viagem_estimado = tempo_cru_base * fator_corredor (arredondado mÃºltiplo 5)."""
        from eventos.services.estimativa_local import (
            estimar_distancia_duracao,
            _velocidade_base_por_faixa,
            arredondar_para_multiplo_5_proximo,
            FATOR_CORREDOR,
        )
        out = estimar_distancia_duracao(-25.43, -49.27, -23.42, -51.93)
        self.assertTrue(out['ok'])
        dist = float(out['distancia_km'])
        vel = _velocidade_base_por_faixa(dist)
        tempo_base_float = (dist / vel) * 60
        tempo_base = arredondar_para_multiplo_5_proximo(tempo_base_float)
        corredor = out['corredor']
        esperado_float = tempo_base * FATOR_CORREDOR[corredor]
        esperado = arredondar_para_multiplo_5_proximo(esperado_float)
        self.assertEqual(out['tempo_viagem_estimado_min'], esperado)
        self.assertEqual(out['tempo_viagem_estimado_min'] % 5, 0)

    def test_buffer_operacional_por_faixa_de_distancia(self):
        from eventos.services.estimativa_local import sugerir_buffer_operacional

        self.assertEqual(sugerir_buffer_operacional('IGNORADO', 60), 15)
        self.assertEqual(sugerir_buffer_operacional('IGNORADO', 61), 20)
        self.assertEqual(sugerir_buffer_operacional('IGNORADO', 120), 20)
        self.assertEqual(sugerir_buffer_operacional('IGNORADO', 121), 25)
        self.assertEqual(sugerir_buffer_operacional('IGNORADO', 200), 25)
        self.assertEqual(sugerir_buffer_operacional('IGNORADO', 201), 35)
        self.assertEqual(sugerir_buffer_operacional('IGNORADO', 300), 35)
        self.assertEqual(sugerir_buffer_operacional('IGNORADO', 301), 45)
        self.assertEqual(sugerir_buffer_operacional('IGNORADO', 450), 45)
        self.assertEqual(sugerir_buffer_operacional('IGNORADO', 451), 60)
        self.assertEqual(sugerir_buffer_operacional('IGNORADO', 600), 60)
        self.assertEqual(sugerir_buffer_operacional('IGNORADO', 601), 75)

    def test_fator_rodoviario_por_faixa(self):
        """Fator rodoviÃ¡rio por linha reta: atÃ© 60â†’1.20; 61-120â†’1.18; 121-250â†’1.17; 251-400â†’1.19; 401-700â†’1.22; >700â†’1.24."""
        from eventos.services.estimativa_local import estimar_distancia_duracao, _fator_rodoviario_por_faixa
        # ~100 km linha reta: faixa 61-120 -> fator 1.18 -> rodoviÃ¡rio ~118 km
        out = estimar_distancia_duracao(0, 0, 0, 0.9)
        self.assertTrue(out['ok'])
        self.assertAlmostEqual(float(out['distancia_km']), 118, delta=8)
        self.assertEqual(float(_fator_rodoviario_por_faixa(50)), 1.20)
        self.assertEqual(float(_fator_rodoviario_por_faixa(60)), 1.20)
        self.assertEqual(float(_fator_rodoviario_por_faixa(100)), 1.18)
        self.assertEqual(float(_fator_rodoviario_por_faixa(200)), 1.17)
        self.assertEqual(float(_fator_rodoviario_por_faixa(350)), 1.19)
        self.assertEqual(float(_fator_rodoviario_por_faixa(500)), 1.22)
        self.assertEqual(float(_fator_rodoviario_por_faixa(800)), 1.24)

    def test_arredondamento_multiplo_5_proximo(self):
        """Arredondar para o mÃºltiplo de 5 mais prÃ³ximo (neutro)."""
        from eventos.services.estimativa_local import arredondar_para_multiplo_5_proximo
        self.assertEqual(arredondar_para_multiplo_5_proximo(362), 360)  # 6h02 -> 6h00
        self.assertEqual(arredondar_para_multiplo_5_proximo(363), 365)  # 6h03 -> 6h05
        self.assertEqual(arredondar_para_multiplo_5_proximo(367), 365)  # 6h07 -> 6h05
        self.assertEqual(arredondar_para_multiplo_5_proximo(368), 370)  # 6h08 -> 6h10
        self.assertEqual(arredondar_para_multiplo_5_proximo(0), 0)
        self.assertEqual(arredondar_para_multiplo_5_proximo(5), 5)
        self.assertEqual(arredondar_para_multiplo_5_proximo(365), 365)

    def test_arredondamento_cima_bloco_5(self):
        from eventos.services.estimativa_local import arredondar_minutos_para_cima_5
        self.assertEqual(arredondar_minutos_para_cima_5(153), 155)  # 2h33 -> 2h35
        self.assertEqual(arredondar_minutos_para_cima_5(155), 155)  # 2h35 -> 2h35
        self.assertEqual(arredondar_minutos_para_cima_5(156), 160)  # 2h36 -> 2h40
        self.assertEqual(arredondar_minutos_para_cima_5(361), 365)  # 6h01 -> 6h05
        self.assertEqual(arredondar_minutos_para_cima_5(0), 0)
        self.assertEqual(arredondar_minutos_para_cima_5(5), 5)

    def test_sem_coordenadas_retorna_erro_amigavel(self):
        from eventos.services.estimativa_local import estimar_distancia_duracao, ERRO_SEM_COORDENADAS
        out = estimar_distancia_duracao(None, -49.27, -23.42, -51.94)
        self.assertFalse(out['ok'])
        self.assertEqual(out['erro'], ERRO_SEM_COORDENADAS)
        out = estimar_distancia_duracao(-25.43, None, -23.42, -51.94)
        self.assertFalse(out['ok'])
        self.assertEqual(out['erro'], ERRO_SEM_COORDENADAS)

    def test_tempo_viagem_multiplo_5_minutos(self):
        """tempo_viagem_estimado_min Ã© mÃºltiplo de 5 (arredondamento neutro)."""
        from eventos.services.estimativa_local import estimar_distancia_duracao
        out = estimar_distancia_duracao(-25.43, -49.27, -23.42, -51.94)
        self.assertTrue(out['ok'])
        self.assertIsNotNone(out['tempo_viagem_estimado_min'])
        self.assertEqual(out['tempo_viagem_estimado_min'] % 5, 0)

    def test_velocidade_base_por_faixa(self):
        """Velocidade base por faixa: atÃ© 60â†’48; 61-120â†’56; 121-250â†’64; 251-400â†’70; 401-700â†’76; >700â†’80."""
        from eventos.services.estimativa_local import _velocidade_base_por_faixa
        self.assertEqual(_velocidade_base_por_faixa(50), 48)
        self.assertEqual(_velocidade_base_por_faixa(60), 48)
        self.assertEqual(_velocidade_base_por_faixa(100), 56)
        self.assertEqual(_velocidade_base_por_faixa(200), 64)
        self.assertEqual(_velocidade_base_por_faixa(350), 70)
        self.assertEqual(_velocidade_base_por_faixa(600), 76)
        self.assertEqual(_velocidade_base_por_faixa(800), 80)

    def test_perfil_rota_retornado_na_estimativa(self):
        """Estimativa retorna perfil_rota (EIXO_PRINCIPAL, DIAGONAL_LONGA, LITORAL_SERRA, URBANA_CURTA ou PADRAO)."""
        from eventos.services.estimativa_local import (
            estimar_distancia_duracao,
            PERFIL_EIXO_PRINCIPAL,
            PERFIL_DIAGONAL_LONGA,
            PERFIL_LITORAL_SERRA,
            PERFIL_URBANA_CURTA,
            PERFIL_PADRAO,
        )
        out = estimar_distancia_duracao(-25.43, -49.27, -23.42, -51.94)
        self.assertTrue(out['ok'])
        self.assertIn(
            out.get('perfil_rota'),
            (PERFIL_EIXO_PRINCIPAL, PERFIL_DIAGONAL_LONGA, PERFIL_LITORAL_SERRA, PERFIL_URBANA_CURTA, PERFIL_PADRAO),
        )

    def test_estimar_tempo_por_distancia_rodoviaria(self):
        """estimar_tempo_por_distancia_rodoviaria retorna tempo_viagem, buffer e duraÃ§Ã£o pela distÃ¢ncia."""
        from eventos.services.estimativa_local import (
            estimar_tempo_por_distancia_rodoviaria,
            CORREDOR_PADRAO,
            CORREDOR_NORTE_NOROESTE,
            sugerir_buffer_operacional,
        )
        out = estimar_tempo_por_distancia_rodoviaria(100.0)
        self.assertIn('tempo_viagem_estimado_min', out)
        self.assertIn('buffer_operacional_sugerido_min', out)
        self.assertIn('duracao_estimada_min', out)
        self.assertIn('corredor', out)
        self.assertEqual(out['corredor'], CORREDOR_PADRAO)
        self.assertEqual(out['buffer_operacional_sugerido_min'], sugerir_buffer_operacional(out['corredor'], 100.0))
        self.assertEqual(out['buffer_operacional_sugerido_min'], 20)
        self.assertEqual(out['duracao_estimada_min'], out['tempo_viagem_estimado_min'] + out['buffer_operacional_sugerido_min'])
        self.assertEqual(out['tempo_viagem_estimado_min'] % 5, 0)
        out2 = estimar_tempo_por_distancia_rodoviaria(100.0, corredor=CORREDOR_NORTE_NOROESTE)
        self.assertEqual(out2['corredor'], CORREDOR_NORTE_NOROESTE)
        self.assertEqual(out2['buffer_operacional_sugerido_min'], 20)

    def test_novos_campos_retorno_route_aware(self):
        """Retorno inclui corredor_macro, corredor_fino, fallback_usado, confianca_estimativa, distancia_linha_reta_km."""
        from eventos.services.estimativa_local import estimar_distancia_duracao
        out = estimar_distancia_duracao(-25.43, -49.27, -23.42, -51.94)
        self.assertTrue(out['ok'])
        self.assertIn('corredor_macro', out)
        self.assertIn('corredor_fino', out)
        self.assertIn('fallback_usado', out)
        self.assertIn('confianca_estimativa', out)
        self.assertIn('distancia_linha_reta_km', out)
        self.assertIn('distancia_rodoviaria_km', out)
        self.assertIn('refs_predominantes', out)
        self.assertIn('pedagio_presente', out)
        self.assertIn('travessia_urbana_presente', out)
        self.assertIn('serra_presente', out)
        self.assertIn(out['confianca_estimativa'], ('alta', 'media', 'baixa'))

    def test_degradacao_limpa_sem_provider(self):
        """Sem OSRM configurado: fallback_usado=True, rota_fonte=ESTIMATIVA_LOCAL, aplicaÃ§Ã£o nÃ£o quebra."""
        from eventos.services.estimativa_local import (
            estimar_distancia_duracao,
            ROTA_FONTE_ESTIMATIVA_LOCAL,
        )
        with patch('eventos.services.estimativa_local.get_default_routing_provider', return_value=None):
            out = estimar_distancia_duracao(-25.43, -49.27, -25.09, -50.16)
        self.assertTrue(out['ok'])
        self.assertTrue(out.get('fallback_usado') is True)
        self.assertEqual(out.get('rota_fonte'), ROTA_FONTE_ESTIMATIVA_LOCAL)
        self.assertEqual(out['duracao_estimada_min'], out['tempo_viagem_estimado_min'] + out['buffer_operacional_sugerido_min'])

    def test_buffer_nao_contamina_tempo_comparavel_maps(self):
        """tempo_viagem_estimado_min Ã© ETA tÃ©cnico (comparÃ¡vel ao Maps); buffer estÃ¡ separado."""
        from eventos.services.estimativa_local import estimar_distancia_duracao
        out = estimar_distancia_duracao(-25.43, -49.27, -24.96, -53.45)
        self.assertTrue(out['ok'])
        self.assertIsNotNone(out['tempo_viagem_estimado_min'])
        self.assertIsNotNone(out['buffer_operacional_sugerido_min'])
        self.assertEqual(out['duracao_estimada_min'], out['tempo_viagem_estimado_min'] + out['buffer_operacional_sugerido_min'])

    def test_rotas_parana_curitiba_ponta_grossa(self):
        """Curitiba -> Ponta Grossa: corredor Campos Gerais, buffer pela distÃ¢ncia."""
        from eventos.services.estimativa_local import estimar_distancia_duracao, sugerir_buffer_operacional
        out = estimar_distancia_duracao(-25.4284, -49.2733, -25.09, -50.16)
        self.assertTrue(out['ok'])
        self.assertEqual(out['corredor'], 'CAMPOS_GERAIS_CURTO')
        self.assertEqual(
            out['buffer_operacional_sugerido_min'],
            sugerir_buffer_operacional(out['corredor_macro'], out['distancia_rodoviaria_km']),
        )

    def test_rotas_parana_curitiba_cascavel(self):
        """Curitiba -> Cascavel: corredor Oeste BR-277, buffer pela distÃ¢ncia."""
        from eventos.services.estimativa_local import estimar_distancia_duracao, sugerir_buffer_operacional
        out = estimar_distancia_duracao(-25.4284, -49.2733, -24.96, -53.45)
        self.assertTrue(out['ok'])
        self.assertEqual(out['corredor'], 'OESTE_BR277')
        self.assertEqual(
            out['buffer_operacional_sugerido_min'],
            sugerir_buffer_operacional(out['corredor_macro'], out['distancia_rodoviaria_km']),
        )

    def test_caminho_com_provider_mock(self):
        """Com provider retornando rota: fallback_usado=False, rota_fonte=OSRM, ETA do provider."""
        from eventos.services.estimativa_local import estimar_distancia_duracao
        from eventos.services.routing_provider import RouteResult
        mock_route = RouteResult(
            distance_km=120.0,
            duration_min=95.0,
            refs_predominantes=['BR-376'],
            steps=[],
            geometry=None,
            raw=None,
        )
        provider = MagicMock()
        provider.route.return_value = mock_route
        with patch('eventos.services.estimativa_local.get_default_routing_provider', return_value=provider):
            out = estimar_distancia_duracao(-25.43, -49.27, -23.42, -51.93)
        self.assertTrue(out['ok'])
        self.assertFalse(out.get('fallback_usado'))
        self.assertEqual(out.get('rota_fonte'), 'OSRM')
        self.assertIsNotNone(out.get('tempo_viagem_estimado_min'))
        self.assertEqual(out['duracao_estimada_min'], out['tempo_viagem_estimado_min'] + out['buffer_operacional_sugerido_min'])


class CalibracaoEstimativaLocalTest(TestCase):
    """Testes da camada de calibraÃ§Ã£o (ajustes por corredor, faixa, atributos)."""

    def test_get_faixa_distancia_key(self):
        from eventos.services.estimativa_local import get_faixa_distancia_key, FAIXA_ATE_60, FAIXA_61_120, FAIXA_121_250, FAIXA_251_400, FAIXA_401_700, FAIXA_ACIMA_700
        self.assertEqual(get_faixa_distancia_key(30), FAIXA_ATE_60)
        self.assertEqual(get_faixa_distancia_key(60), FAIXA_ATE_60)
        self.assertEqual(get_faixa_distancia_key(61), FAIXA_61_120)
        self.assertEqual(get_faixa_distancia_key(120), FAIXA_61_120)
        self.assertEqual(get_faixa_distancia_key(121), FAIXA_121_250)
        self.assertEqual(get_faixa_distancia_key(250), FAIXA_121_250)
        self.assertEqual(get_faixa_distancia_key(251), FAIXA_251_400)
        self.assertEqual(get_faixa_distancia_key(400), FAIXA_251_400)
        self.assertEqual(get_faixa_distancia_key(401), FAIXA_401_700)
        self.assertEqual(get_faixa_distancia_key(700), FAIXA_401_700)
        self.assertEqual(get_faixa_distancia_key(701), FAIXA_ACIMA_700)

    def test_aplicar_calibracao_eta_sem_ajustes(self):
        from eventos.services.estimativa_local import _aplicar_calibracao_eta
        from eventos.services import corredores_pr as corredores
        eta = _aplicar_calibracao_eta(
            100.0, 30.0, corredores.CORREDOR_PADRAO, corredores.CORREDOR_FINO_PADRAO,
            serra_presente=False, travessia_urbana_presente=False, pedagio_presente=False, ref_predominante=None,
        )
        self.assertGreaterEqual(eta, 0)
        self.assertAlmostEqual(eta, 100.0, delta=0.01)

    def test_aplicar_calibracao_eta_com_ajuste_macro(self):
        from eventos.services.estimativa_local import _aplicar_calibracao_eta, AJUSTE_CORREDOR_MACRO_MIN
        from eventos.services import corredores_pr as corredores
        original = AJUSTE_CORREDOR_MACRO_MIN.get(corredores.CORREDOR_PADRAO, 0)
        try:
            AJUSTE_CORREDOR_MACRO_MIN[corredores.CORREDOR_PADRAO] = 10
            eta = _aplicar_calibracao_eta(
                90.0, 30.0, corredores.CORREDOR_PADRAO, corredores.CORREDOR_FINO_PADRAO,
            )
            self.assertAlmostEqual(eta, 100.0, delta=0.01)
        finally:
            AJUSTE_CORREDOR_MACRO_MIN[corredores.CORREDOR_PADRAO] = original

    def test_aplicar_calibracao_eta_com_ajuste_fino(self):
        from eventos.services.estimativa_local import _aplicar_calibracao_eta, AJUSTE_CORREDOR_FINO_MIN
        from eventos.services import corredores_pr as corredores
        original = AJUSTE_CORREDOR_FINO_MIN.get(corredores.MARINGA, 0)
        try:
            AJUSTE_CORREDOR_FINO_MIN[corredores.MARINGA] = -5
            eta = _aplicar_calibracao_eta(
                105.0, 30.0, corredores.CORREDOR_PADRAO, corredores.MARINGA,
            )
            self.assertAlmostEqual(eta, 100.0, delta=0.01)
        finally:
            AJUSTE_CORREDOR_FINO_MIN[corredores.MARINGA] = original

    def test_aplicar_calibracao_eta_com_ajuste_faixa(self):
        from eventos.services.estimativa_local import _aplicar_calibracao_eta, AJUSTE_FAIXA_DISTANCIA_MIN, FAIXA_121_250
        from eventos.services import corredores_pr as corredores
        original = AJUSTE_FAIXA_DISTANCIA_MIN.get(FAIXA_121_250, 0)
        try:
            AJUSTE_FAIXA_DISTANCIA_MIN[FAIXA_121_250] = 3
            eta = _aplicar_calibracao_eta(
                97.0, 200.0, corredores.CORREDOR_PADRAO, corredores.CORREDOR_FINO_PADRAO,
            )
            self.assertAlmostEqual(eta, 100.0, delta=0.01)
        finally:
            AJUSTE_FAIXA_DISTANCIA_MIN[FAIXA_121_250] = original

    def test_aplicar_calibracao_eta_com_ajuste_atributos(self):
        from eventos.services.estimativa_local import _aplicar_calibracao_eta, AJUSTE_ATRIBUTOS_MIN
        from eventos.services import corredores_pr as corredores
        original_serra = AJUSTE_ATRIBUTOS_MIN.get('serra', 0)
        original_urbana = AJUSTE_ATRIBUTOS_MIN.get('travessia_urbana', 0)
        try:
            AJUSTE_ATRIBUTOS_MIN['serra'] = 4
            AJUSTE_ATRIBUTOS_MIN['travessia_urbana'] = 6
            eta = _aplicar_calibracao_eta(
                90.0, 30.0, corredores.CORREDOR_PADRAO, corredores.CORREDOR_FINO_PADRAO,
                serra_presente=True, travessia_urbana_presente=True, pedagio_presente=False,
            )
            self.assertAlmostEqual(eta, 100.0, delta=0.01)
        finally:
            AJUSTE_ATRIBUTOS_MIN['serra'] = original_serra
            AJUSTE_ATRIBUTOS_MIN['travessia_urbana'] = original_urbana

    def test_aplicar_calibracao_eta_com_ajuste_ref_predominante(self):
        from eventos.services.estimativa_local import _aplicar_calibracao_eta, AJUSTE_REF_PREDOMINANTE_MIN
        from eventos.services import corredores_pr as corredores
        original = AJUSTE_REF_PREDOMINANTE_MIN.get('BR-376', 0)
        try:
            AJUSTE_REF_PREDOMINANTE_MIN['BR-376'] = -8
            eta = _aplicar_calibracao_eta(
                108.0, 30.0, corredores.CORREDOR_PADRAO, corredores.CORREDOR_FINO_PADRAO,
                ref_predominante='BR-376',
            )
            self.assertAlmostEqual(eta, 100.0, delta=0.01)
        finally:
            AJUSTE_REF_PREDOMINANTE_MIN['BR-376'] = original

    def test_duracao_estimada_igual_tempo_viagem_mais_buffer_com_osrm(self):
        """duracao_estimada_min = tempo_viagem_estimado_min + buffer_operacional_sugerido_min (com provider)."""
        from eventos.services.estimativa_local import estimar_distancia_duracao
        from eventos.services.routing_provider import RouteResult
        mock_route = RouteResult(
            distance_km=200.0,
            duration_min=120.0,
            refs_predominantes=['BR-376'],
            steps=[],
            geometry=None,
            raw=None,
        )
        provider = MagicMock()
        provider.route.return_value = mock_route
        with patch('eventos.services.estimativa_local.get_default_routing_provider', return_value=provider):
            out = estimar_distancia_duracao(-25.43, -49.27, -23.42, -51.93)
        self.assertTrue(out['ok'])
        self.assertEqual(
            out['duracao_estimada_min'],
            out['tempo_viagem_estimado_min'] + out['buffer_operacional_sugerido_min'],
        )

    def test_buffer_nao_contamina_eta_com_osrm(self):
        """tempo_viagem_estimado_min Ã© ETA tÃ©cnico; buffer separado (com provider)."""
        from eventos.services.estimativa_local import estimar_distancia_duracao
        from eventos.services.routing_provider import RouteResult
        mock_route = RouteResult(
            distance_km=200.0,
            duration_min=120.0,
            refs_predominantes=[],
            steps=[],
            geometry=None,
            raw=None,
        )
        provider = MagicMock()
        provider.route.return_value = mock_route
        with patch('eventos.services.estimativa_local.get_default_routing_provider', return_value=provider):
            out = estimar_distancia_duracao(-25.43, -49.27, -23.42, -51.93)
        self.assertTrue(out['ok'])
        self.assertIsNotNone(out['tempo_viagem_estimado_min'])
        self.assertIsNotNone(out['buffer_operacional_sugerido_min'])
        self.assertEqual(out['duracao_estimada_min'], out['tempo_viagem_estimado_min'] + out['buffer_operacional_sugerido_min'])


class EstimativaParanaProviderClassificacaoTest(TestCase):
    """Cobertura dos corredores/atributos do provider route-aware do Parana."""

    def test_provider_classifica_paranagua_sem_confundir_com_pontal(self):
        from eventos.services.estimativa_local import estimar_distancia_duracao
        from eventos.services.routing_provider import RouteResult

        mock_route = RouteResult(
            distance_km=90.0,
            duration_min=65.0,
            refs_predominantes=['BR-277'],
            steps=[{'distance': 90000, 'duration': 3900, 'name': 'BR-277', 'ref': 'BR-277', 'road_refs': ['BR-277']}],
            geometry=[
                (-49.264622, -25.419547),
                (-49.05, -25.33),
                (-48.83, -25.40),
                (-48.67, -25.46),
                (-48.522528, -25.516078),
            ],
            raw=None,
        )
        provider = MagicMock()
        provider.route.return_value = mock_route

        with patch('eventos.services.estimativa_local.get_default_routing_provider', return_value=provider):
            out = estimar_distancia_duracao(-25.419547, -49.264622, -25.516078, -48.522528)

        self.assertEqual(out['rota_fonte'], 'OSRM')
        self.assertEqual(out['corredor_macro'], 'BR277_LITORAL')
        self.assertEqual(out['corredor_fino'], 'PARANAGUA')
        self.assertTrue(out['serra_presente'])

    def test_provider_classifica_cruzeiro_do_sul_no_noroeste(self):
        from eventos.services.estimativa_local import estimar_distancia_duracao
        from eventos.services.routing_provider import RouteResult

        mock_route = RouteResult(
            distance_km=495.0,
            duration_min=410.0,
            refs_predominantes=['BR-376', 'PR-463'],
            steps=[
                {'distance': 380000, 'duration': 20000, 'name': 'BR-376', 'ref': 'BR-376', 'road_refs': ['BR-376']},
                {'distance': 115000, 'duration': 4600, 'name': 'PR-463', 'ref': 'PR-463', 'road_refs': ['PR-463']},
            ],
            geometry=[
                (-49.264622, -25.419547),
                (-50.5, -24.2),
                (-51.4, -23.6),
                (-51.9, -23.1),
                (-52.162210, -22.962440),
            ],
            raw=None,
        )
        provider = MagicMock()
        provider.route.return_value = mock_route

        with patch('eventos.services.estimativa_local.get_default_routing_provider', return_value=provider):
            out = estimar_distancia_duracao(-25.419547, -49.264622, -22.962440, -52.162210)

        self.assertEqual(out['corredor_macro'], 'NOROESTE_INTERIOR')
        self.assertEqual(out['corredor_fino'], 'CRUZEIRO_DO_SUL')
        self.assertFalse(out['serra_presente'])

    def test_serra_presente_fica_false_no_oeste_com_br277(self):
        from eventos.services.estimativa_local import estimar_distancia_duracao
        from eventos.services.routing_provider import RouteResult

        mock_route = RouteResult(
            distance_km=500.0,
            duration_min=395.0,
            refs_predominantes=['BR-277'],
            steps=[{'distance': 500000, 'duration': 23700, 'name': 'BR-277', 'ref': 'BR-277', 'road_refs': ['BR-277']}],
            geometry=[
                (-49.264622, -25.419547),
                (-50.1, -25.30),
                (-51.2, -25.18),
                (-52.3, -25.04),
                (-53.459005, -24.957301),
            ],
            raw=None,
        )
        provider = MagicMock()
        provider.route.return_value = mock_route

        with patch('eventos.services.estimativa_local.get_default_routing_provider', return_value=provider):
            out = estimar_distancia_duracao(-25.419547, -49.264622, -24.957301, -53.459005)

        self.assertEqual(out['corredor_macro'], 'BR277_OESTE')
        self.assertEqual(out['corredor_fino'], 'CASCAVEL')
        self.assertFalse(out['serra_presente'])

    def test_benchmark_oficial_contem_rotas_minimas_do_parana(self):
        from scripts.analisar_estimativa_pr import load_benchmark

        destinos_esperados = {
            'Pontal do ParanÃ¡',
            'ParanaguÃ¡',
            'Ponta Grossa',
            'TelÃªmaco Borba',
            'Guarapuava',
            'Londrina',
            'Apucarana',
            'MaringÃ¡',
            'Cianorte',
            'Umuarama',
            'Cruzeiro do Sul',
            'Cascavel',
            'Palotina',
            'Foz do IguaÃ§u',
            'Francisco BeltrÃ£o',
        }

        destinos = {item['destino_nome'] for item in load_benchmark()}
        self.assertTrue(destinos_esperados.issubset(destinos))


class ScriptAnalisarEstimativaPrTest(TestCase):
    """Testes do script de benchmark/calibraÃ§Ã£o (mÃ©tricas e sugestÃ£o)."""

    def test_script_calcula_metricas_com_benchmark_temporario(self):
        import tempfile
        from scripts.analisar_estimativa_pr import load_benchmark, tempo_referencia
        # Benchmark com 2 rotas e tempo de referÃªncia: erros absolutos 10 e 20 â†’ MAE 15
        records = [
            {
                'origem_lat': -25.43, 'origem_lon': -49.27, 'destino_lat': -25.09, 'destino_lon': -50.16,
                'tempo_referencia_min': 100,
            },
            {
                'origem_lat': -25.43, 'origem_lon': -49.27, 'destino_lat': -23.42, 'destino_lon': -51.93,
                'tempo_referencia_min': 200,
            },
        ]
        with tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False, encoding='utf-8') as f:
            json.dump(records, f, ensure_ascii=False)
            path = Path(f.name)
        try:
            loaded = load_benchmark(path)
            self.assertEqual(len(loaded), 2)
            self.assertEqual(tempo_referencia(loaded[0]), 100)
            self.assertEqual(tempo_referencia(loaded[1]), 200)
        finally:
            path.unlink(missing_ok=True)

    def test_script_sugerir_calibracao_emite_saida_esperada(self):
        import subprocess
        import tempfile
        BASE_DIR = Path(settings.BASE_DIR)
        # Dois registros mesmo corredor: erro +12 e +12 â†’ sugestÃ£o -12
        records = [
            {
                'origem_lat': -25.4284, 'origem_lon': -49.2733, 'destino_lat': -25.09, 'destino_lon': -50.16,
                'tempo_referencia_min': 90,
            },
            {
                'origem_lat': -25.4284, 'origem_lon': -49.2733, 'destino_lat': -25.08, 'destino_lon': -50.15,
                'tempo_referencia_min': 90,
            },
        ]
        with tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False, encoding='utf-8') as f:
            json.dump(records, f, ensure_ascii=False)
            path = Path(f.name)
        try:
            result = subprocess.run(
                [sys.executable, str(BASE_DIR / 'scripts' / 'analisar_estimativa_pr.py'), '--sugerir-calibracao', '--benchmark-file', str(path)],
                capture_output=True,
                text=True,
                cwd=str(BASE_DIR),
                env={**os.environ, 'DJANGO_SETTINGS_MODULE': 'config.settings'},
                timeout=30,
            )
            out = result.stdout + result.stderr
            self.assertIn('Sugestoes de calibracao', out or '')
            self.assertIn('AJUSTE_', out or '')
        finally:
            path.unlink(missing_ok=True)


class OSRMRoutingProviderHTTPTest(TestCase):
    """Testes do OSRMRoutingProvider com mock de HTTP (requests)."""

    def test_osrm_retorna_rota_quando_http_200_e_json_valido(self):
        """Quando requests.get retorna 200 e JSON com routes[0], provider retorna RouteResult."""
        from eventos.services.routing_provider import OSRMRoutingProvider, RouteResult
        response = MagicMock()
        response.raise_for_status = MagicMock()
        response.json.return_value = {
            'routes': [{
                'distance': 120000,
                'duration': 5700,
                'legs': [{
                    'steps': [
                        {'distance': 1000, 'duration': 60, 'name': 'BR-376', 'ref': 'BR-376'},
                    ],
                }],
            }],
        }
        with patch('requests.get', return_value=response):
            provider = OSRMRoutingProvider('http://localhost:5000', timeout_seconds=5)
            result = provider.route(-25.43, -49.27, -23.42, -51.93)
        self.assertIsInstance(result, RouteResult)
        self.assertEqual(result.distance_km, 120.0)
        self.assertEqual(result.duration_min, 95.0)
        self.assertIn('BR-376', result.refs_predominantes)

    def test_osrm_retorna_none_em_timeout(self):
        """Quando requests.get dÃ¡ timeout, provider retorna None (sistema usa fallback)."""
        import requests
        from eventos.services.routing_provider import OSRMRoutingProvider
        with patch('requests.get', side_effect=requests.Timeout('timeout')):
            provider = OSRMRoutingProvider('http://localhost:5000', timeout_seconds=5)
            result = provider.route(-25.43, -49.27, -23.42, -51.93)
        self.assertIsNone(result)

    def test_osrm_retorna_none_em_erro_http(self):
        """Quando requests.get retorna HTTP 500, provider retorna None."""
        from eventos.services.routing_provider import OSRMRoutingProvider
        response = MagicMock()
        response.raise_for_status.side_effect = Exception('500 Server Error')
        with patch('requests.get', return_value=response):
            provider = OSRMRoutingProvider('http://localhost:5000', timeout_seconds=5)
            result = provider.route(-25.43, -49.27, -23.42, -51.93)
        self.assertIsNone(result)

    def test_osrm_retorna_none_quando_sem_rotas(self):
        """Quando JSON tem routes vazio, provider retorna None."""
        from eventos.services.routing_provider import OSRMRoutingProvider
        response = MagicMock()
        response.raise_for_status = MagicMock()
        response.json.return_value = {'routes': []}
        with patch('requests.get', return_value=response):
            provider = OSRMRoutingProvider('http://localhost:5000', timeout_seconds=5)
            result = provider.route(-25.43, -49.27, -23.42, -51.93)
        self.assertIsNone(result)

    def test_osrm_retorna_none_quando_base_url_vazia(self):
        """Quando base_url Ã© vazio, provider.route retorna None sem chamar HTTP."""
        from eventos.services.routing_provider import OSRMRoutingProvider
        provider = OSRMRoutingProvider('', timeout_seconds=5)
        with patch('requests.get') as mock_get:
            result = provider.route(-25.43, -49.27, -23.42, -51.93)
        self.assertIsNone(result)
        mock_get.assert_not_called()


class TrechoCalcularKmEndpointTest(TestCase):
    """Testes do endpoint POST trechos/<pk>/calcular-km/ (estimativa local)."""

    def setUp(self):
        self.user = User.objects.create_user(username='u', password='p')
        self.client = Client()
        self.estado = Estado.objects.create(nome='ParanÃ¡', sigla='PR', codigo_ibge='41')
        self.cidade_a = Cidade.objects.create(
            nome='Curitiba', estado=self.estado, codigo_ibge='4106902',
            latitude=Decimal('-25.4284'), longitude=Decimal('-49.2733'),
        )
        self.cidade_b = Cidade.objects.create(
            nome='MaringÃ¡', estado=self.estado, codigo_ibge='4115200',
            latitude=Decimal('-23.4205'), longitude=Decimal('-51.9332'),
        )
        self.evento = Evento.objects.create(
            titulo='Evento Teste',
            data_inicio=date(2025, 1, 1),
            data_fim=date(2025, 1, 1),
            status=Evento.STATUS_RASCUNHO,
        )
        self.roteiro = RoteiroEvento.objects.create(
            evento=self.evento,
            origem_estado=self.estado,
            origem_cidade=self.cidade_a,
        )
        RoteiroEventoDestino.objects.create(
            roteiro=self.roteiro,
            estado=self.estado,
            cidade=self.cidade_b,
            ordem=0,
        )
        self.trecho = RoteiroEventoTrecho.objects.create(
            roteiro=self.roteiro,
            ordem=0,
            tipo=RoteiroEventoTrecho.TIPO_IDA,
            origem_estado=self.estado,
            origem_cidade=self.cidade_a,
            destino_estado=self.estado,
            destino_cidade=self.cidade_b,
        )

    def test_endpoint_exige_login(self):
        url = reverse('eventos:trecho-calcular-km', kwargs={'pk': self.trecho.pk})
        response = self.client.post(url, {}, HTTP_X_REQUESTED_WITH='XMLHttpRequest')
        self.assertEqual(response.status_code, 302)
        self.assertIn('login', response.url)

    def test_endpoint_retorna_km_e_duracao_estimativa_local(self):
        self.client.login(username='u', password='p')
        url = reverse('eventos:trecho-calcular-km', kwargs={'pk': self.trecho.pk})
        response = self.client.post(
            url,
            {},
            content_type='application/json',
            HTTP_X_REQUESTED_WITH='XMLHttpRequest',
        )
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertTrue(data['ok'], data.get('erro'))
        self.assertIsNotNone(data['distancia_km'])
        self.assertIsNotNone(data['duracao_estimada_min'])
        self.assertIn('duracao_estimada_hhmm', data)
        self.assertIn('tempo_cru_estimado_min', data)
        self.assertIn('tempo_adicional_sugerido_min', data)
        self.assertEqual(data.get('rota_fonte'), 'ESTIMATIVA_LOCAL')
        self.trecho.refresh_from_db()
        self.assertIsNotNone(self.trecho.distancia_km)
        self.assertIsNotNone(self.trecho.duracao_estimada_min)
        self.assertIsNotNone(self.trecho.tempo_cru_estimado_min)
        self.assertIsNotNone(self.trecho.tempo_adicional_min)
        # duracao_estimada_min persiste o total retornado pelo serviÃ§o (cru + adicional + correÃ§Ã£o final)
        self.assertEqual(self.trecho.duracao_estimada_min, data['duracao_estimada_min'])
        self.assertEqual(self.trecho.rota_fonte, 'ESTIMATIVA_LOCAL')
        self.assertIsNotNone(self.trecho.rota_calculada_em)
        self.assertEqual(self.trecho.tempo_cru_estimado_min % 5, 0)

    def test_endpoint_erro_sem_coordenadas(self):
        self.cidade_b.latitude = None
        self.cidade_b.longitude = None
        self.cidade_b.save()
        self.client.login(username='u', password='p')
        url = reverse('eventos:trecho-calcular-km', kwargs={'pk': self.trecho.pk})
        response = self.client.post(
            url,
            {},
            content_type='application/json',
            HTTP_X_REQUESTED_WITH='XMLHttpRequest',
        )
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertFalse(data['ok'])
        self.assertIn('coordenadas', data['erro'].lower())
        self.trecho.refresh_from_db()
        self.assertIsNone(self.trecho.distancia_km)
        self.assertIsNone(self.trecho.duracao_estimada_min)

    @patch('eventos.views.estimar_distancia_duracao')
    def test_endpoint_trata_erro_estimativa(self, mock_estimar):
        mock_estimar.return_value = {
            'ok': False,
            'distancia_km': None,
            'duracao_estimada_min': None,
            'duracao_estimada_hhmm': '',
            'rota_fonte': 'ESTIMATIVA_LOCAL',
            'erro': 'Cidade sem coordenadas para estimativa.',
        }
        self.client.login(username='u', password='p')
        url = reverse('eventos:trecho-calcular-km', kwargs={'pk': self.trecho.pk})
        response = self.client.post(
            url,
            {},
            content_type='application/json',
            HTTP_X_REQUESTED_WITH='XMLHttpRequest',
        )
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertFalse(data['ok'])
        self.assertEqual(data['erro'], 'Cidade sem coordenadas para estimativa.')
        self.trecho.refresh_from_db()
        self.assertIsNone(self.trecho.distancia_km)
        self.assertIsNone(self.trecho.duracao_estimada_min)


class EstimarKmPorCidadesEndpointTest(TestCase):
    """Testes do endpoint POST trechos/estimar/ (para cadastro, sem trecho salvo)."""

    def setUp(self):
        self.user = User.objects.create_user(username='u', password='p')
        self.client = Client()
        self.estado = Estado.objects.create(nome='ParanÃ¡', sigla='PR', codigo_ibge='41')
        self.cidade_a = Cidade.objects.create(
            nome='Curitiba', estado=self.estado, codigo_ibge='4106902',
            latitude=Decimal('-25.4284'), longitude=Decimal('-49.2733'),
        )
        self.cidade_b = Cidade.objects.create(
            nome='MaringÃ¡', estado=self.estado, codigo_ibge='4115200',
            latitude=Decimal('-23.4205'), longitude=Decimal('-51.9332'),
        )

    def test_estimar_km_por_cidades_exige_login(self):
        url = reverse('eventos:trechos-estimar')
        response = self.client.post(
            url,
            json.dumps({'origem_cidade_id': self.cidade_a.pk, 'destino_cidade_id': self.cidade_b.pk}),
            content_type='application/json',
        )
        self.assertEqual(response.status_code, 302)
        self.assertIn('login', response.url)

    def test_estimar_km_por_cidades_retorna_mesmo_formato_que_trecho_calcular(self):
        """Endpoint retorna ok, distancia_km, tempo_cru_estimado_min, tempo_adicional_sugerido_min (nÃ£o persiste)."""
        self.client.login(username='u', password='p')
        url = reverse('eventos:trechos-estimar')
        response = self.client.post(
            url,
            json.dumps({'origem_cidade_id': self.cidade_a.pk, 'destino_cidade_id': self.cidade_b.pk}),
            content_type='application/json',
        )
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertTrue(data['ok'], data.get('erro'))
        self.assertIn('distancia_km', data)
        self.assertIn('duracao_estimada_min', data)
        self.assertIn('tempo_cru_estimado_min', data)
        self.assertIn('tempo_adicional_sugerido_min', data)
        self.assertEqual(data.get('rota_fonte'), 'ESTIMATIVA_LOCAL')
        if data.get('tempo_viagem_estimado_min') is not None and data.get('buffer_operacional_sugerido_min') is not None:
            self.assertEqual(
                data['duracao_estimada_min'],
                data['tempo_viagem_estimado_min'] + data['buffer_operacional_sugerido_min'],
            )

    def test_estimar_km_por_cidades_sem_ids_retorna_erro(self):
        self.client.login(username='u', password='p')
        url = reverse('eventos:trechos-estimar')
        response = self.client.post(url, json.dumps({}), content_type='application/json')
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertFalse(data['ok'])
        self.assertIn('origem_cidade_id', data.get('erro', '').lower())

    def test_trechos_estimar_trecho_novo_sem_pk(self):
        """Endpoint trechos-estimar/ permite estimar por origem/destino sem trecho salvo (nÃ£o persiste)."""
        self.client.login(username='u', password='p')
        url = reverse('eventos:trechos-estimar')
        n_before = RoteiroEventoTrecho.objects.count()
        response = self.client.post(
            url,
            json.dumps({'origem_cidade_id': self.cidade_a.pk, 'destino_cidade_id': self.cidade_b.pk}),
            content_type='application/json',
        )
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertTrue(data['ok'], data.get('erro'))
        self.assertIn('distancia_km', data)
        self.assertIn('duracao_estimada_min', data)
        self.assertIn('tempo_cru_estimado_min', data)
        self.assertIn('tempo_adicional_sugerido_min', data)
        self.assertIn('duracao_estimada_hhmm', data)
        self.assertEqual(data.get('rota_fonte'), 'ESTIMATIVA_LOCAL')
        self.assertEqual(RoteiroEventoTrecho.objects.count(), n_before)


class PainelBlocosClicaveisTest(TestCase):
    """Painel: ordem de neg?cio e links clic?veis."""

    def setUp(self):
        self.user = User.objects.create_user(username='u', password='p')
        self.client = Client()
        self.client.login(username='u', password='p')
        self.evento = Evento.objects.create(
            titulo='Evento Painel',
            data_inicio=date(2025, 1, 1),
            data_fim=date(2025, 1, 5),
            status=Evento.STATUS_RASCUNHO,
        )

    def test_telas_reais_das_etapas(self):
        urls = [
            reverse('eventos:guiado-etapa-1', kwargs={'pk': self.evento.pk}),
            reverse('eventos:guiado-etapa-2', kwargs={'evento_id': self.evento.pk}),
            reverse('eventos:guiado-etapa-5', kwargs={'evento_id': self.evento.pk}),
            reverse('eventos:guiado-etapa-4', kwargs={'evento_id': self.evento.pk}),
            reverse('eventos:guiado-etapa-5', kwargs={'evento_id': self.evento.pk}),
            reverse('eventos:guiado-etapa-6', kwargs={'evento_id': self.evento.pk}),
            reverse('eventos:guiado-etapa-7', kwargs={'evento_id': self.evento.pk}),
        ]
        for url in urls:
            response = self.client.get(url)
            self.assertEqual(response.status_code, 200)


class EventoEtapa4PtOsTest(TestCase):
    """Etapa 4: usa cadastro real de Plano de Trabalho e Ordem de ServiÃ§o."""

    def setUp(self):
        self.user = User.objects.create_user(username='u', password='p')
        self.client = Client()
        self.client.login(username='u', password='p')
        self.evento = Evento.objects.create(
            titulo='Evento Etapa 4',
            data_inicio=date(2025, 1, 1),
            data_fim=date(2025, 1, 5),
            status=Evento.STATUS_RASCUNHO,
        )

    def test_etapa_4_exige_login(self):
        self.client.logout()
        response = self.client.get(reverse('eventos:guiado-etapa-4', kwargs={'evento_id': self.evento.pk}))
        self.assertEqual(response.status_code, 302)
        self.assertIn('login', response.url)

    def test_etapa_4_get_mostra_links_para_cadastro_real(self):
        response = self.client.get(reverse('eventos:guiado-etapa-4', kwargs={'evento_id': self.evento.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Etapa 4 - PT / OS')
        self.assertContains(response, 'Cadastrar novo')
        self.assertContains(response, 'Plano de Trabalho')
        self.assertContains(response, 'bi-clipboard-check', html=False)
        self.assertContains(response, 'data-scroll-top', html=False)
        self.assertContains(response, 'data-fab-toggle', html=False)
        self.assertContains(response, 'data-fab-menu', html=False)
        self.assertContains(response, 'context_source=evento', html=False)
        self.assertContains(response, f'preselected_event_id={self.evento.pk}')
        self.assertNotContains(response, 'data-list-view-toggle', html=False)

    def test_etapa_4_lista_pt_os_reais_vinculados_ao_evento(self):
        pt = PlanoTrabalho.objects.create(evento=self.evento)
        os = OrdemServico.objects.create(evento=self.evento, finalidade='OS evento')
        PlanoTrabalho.objects.filter(pk=pt.pk).update(updated_at=tz.now() - timedelta(days=1))
        OrdemServico.objects.filter(pk=os.pk).update(updated_at=tz.now())
        response = self.client.get(reverse('eventos:guiado-etapa-4', kwargs={'evento_id': self.evento.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, str(pt.pk))
        self.assertContains(response, str(os.pk))
        self.assertEqual(
            [item['item_type'] for item in response.context['itens_unificados']],
            ['os', 'pt'],
        )

class EventoEtapa5TermosTest(TestCase):
    """Etapa 3 ? Termos no n?vel do evento: status, modalidade e gera??o por participante."""

    def setUp(self):
        self.user = get_user_model().objects.create_user(username='u5', password='p5')
        self.client = Client()
        self.client.login(username='u5', password='p5')
        self.evento = Evento.objects.create(
            titulo='Evento Etapa Termos',
            data_inicio=date(2025, 1, 1),
            data_fim=date(2025, 1, 5),
            status=Evento.STATUS_RASCUNHO,
        )
        self.estado = Estado.objects.create(nome='Paran?', sigla='PR', codigo_ibge='41')
        self.cidade = Cidade.objects.create(nome='Curitiba', estado=self.estado, codigo_ibge='4106902')
        EventoDestino.objects.create(evento=self.evento, estado=self.estado, cidade=self.cidade, ordem=0)

    def _criar_viajante(self, nome='Viajante A', completo=False, cpf='12345678909'):
        kwargs = {
            'nome': nome,
            'cpf': cpf,
        }
        if completo:
            cargo = Cargo.objects.create(nome=f'CARGO {nome}')
            unidade = UnidadeLotacao.objects.create(nome=f'UNIDADE {nome}')
            kwargs.update(
                {
                    'status': Viajante.STATUS_FINALIZADO,
                    'cargo': cargo,
                    'telefone': f'4199{(cpf or "0")[-7:]:0>7}',
                    'unidade_lotacao': unidade,
                    'rg': f'1{(cpf or "0")[-9:]:0>9}',
                }
            )
        return Viajante.objects.create(**kwargs)

    def _vincular_viajante_em_oficio(self, viajante, numero=1):
        oficio = Oficio.objects.create(evento=self.evento, numero=numero, status=Oficio.STATUS_RASCUNHO)
        oficio.viajantes.add(viajante)
        return oficio

    def _criar_veiculo_finalizado(self):
        combustivel = CombustivelVeiculo.objects.create(nome='Gasolina', is_padrao=True)
        return Veiculo.objects.create(
            placa='ABC1D23',
            modelo='Viatura Teste',
            combustivel=combustivel,
            tipo=Veiculo.TIPO_DESCARACTERIZADO,
            status=Veiculo.STATUS_FINALIZADO,
        )

    def test_etapa_5_exige_login(self):
        self.client.logout()
        response = self.client.get(reverse('eventos:guiado-etapa-5', kwargs={'evento_id': self.evento.pk}))
        self.assertEqual(response.status_code, 302)
        self.assertIn('login', response.url)

    def test_etapa_5_get_sem_oficios_exibe_mensagem(self):
        response = self.client.get(reverse('eventos:guiado-etapa-5', kwargs={'evento_id': self.evento.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Nenhum participante no evento')

    def test_etapa_5_get_com_oficio_e_viajante_exibe_tabela(self):
        viajante = self._criar_viajante(nome='Viajante A')
        self._vincular_viajante_em_oficio(viajante)
        response = self.client.get(reverse('eventos:guiado-etapa-5', kwargs={'evento_id': self.evento.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, viajante.nome)
        self.assertContains(response, 'value="PENDENTE"')
        self.assertContains(response, 'value="GERADO"')
        self.assertContains(response, 'value="COMPLETO"')
        self.assertContains(response, 'value="SEMIPREENCHIDO"')

    def test_etapa_5_restaura_toolbar_propria_com_toggle_sem_filtros_globais(self):
        response = self.client.get(reverse('eventos:guiado-etapa-5', kwargs={'evento_id': self.evento.pk}))
        self.assertEqual(response.status_code, 200)
        content = response.content.decode('utf-8')

        self.assertContains(response, 'Completa')
        self.assertContains(response, 'Visualizacao unica')
        self.assertContains(response, 'Novo ofÃ­cio')
        self.assertIn('data-list-view-root', content)
        self.assertIn('data-view-mode="rich"', content)
        self.assertIn('data-view-storage-key="central-viagens.guiado.etapa3.oficios.view-mode"', content)
        self.assertIn('<script src="/static/js/list_standard.js"></script>', content)
        self.assertIn('guiado-etapa3-oficios-empty', content)
        self.assertIn('oficios-view-pane--basic', content)
        self.assertIn('oficios-view-pane--rich', content)
        self.assertIn('guiado-etapa3-oficios-stack', content)
        self.assertNotContains(response, 'Oficios salvos')
        self.assertNotContains(response, 'name="q"', html=False)
        self.assertNotContains(response, 'name="status"', html=False)
        self.assertNotContains(response, 'placeholder="Oficio, protocolo, destino ou servidor"', html=False)

    def test_etapa_5_post_salva_status_e_modalidade(self):
        viajante = self._criar_viajante(nome='Viajante B')
        self._vincular_viajante_em_oficio(viajante)
        response = self.client.post(
            reverse('eventos:guiado-etapa-5', kwargs={'evento_id': self.evento.pk}),
            {
                f'status_{viajante.pk}': EventoTermoParticipante.STATUS_DISPENSADO,
                f'modalidade_{viajante.pk}': EventoTermoParticipante.MODALIDADE_SEMIPREENCHIDO,
            },
        )
        self.assertEqual(response.status_code, 302)
        termo = EventoTermoParticipante.objects.get(evento=self.evento, viajante=viajante)
        self.assertEqual(termo.status, EventoTermoParticipante.STATUS_DISPENSADO)
        self.assertEqual(termo.modalidade, EventoTermoParticipante.MODALIDADE_SEMIPREENCHIDO)

    def test_etapa_5_autosave_persiste_status_e_modalidade(self):
        viajante = self._criar_viajante(nome='Viajante Autosave')
        self._vincular_viajante_em_oficio(viajante)

        response = self.client.post(
            reverse('eventos:guiado-etapa-5', kwargs={'evento_id': self.evento.pk}),
            {
                'autosave': '1',
                f'status_{viajante.pk}': EventoTermoParticipante.STATUS_CONCLUIDO,
                f'modalidade_{viajante.pk}': EventoTermoParticipante.MODALIDADE_SEMIPREENCHIDO,
            },
            HTTP_X_REQUESTED_WITH='XMLHttpRequest',
        )

        self.assertEqual(response.status_code, 200)
        self.assertTrue(response.json().get('ok'))
        termo = EventoTermoParticipante.objects.get(evento=self.evento, viajante=viajante)
        self.assertEqual(termo.status, EventoTermoParticipante.STATUS_CONCLUIDO)
        self.assertEqual(termo.modalidade, EventoTermoParticipante.MODALIDADE_SEMIPREENCHIDO)

    def test_etapa_5_acoes_participante_dispensar_reabrir(self):
        viajante = self._criar_viajante(nome='Viajante C')
        self._vincular_viajante_em_oficio(viajante)
        url = reverse('eventos:guiado-etapa-5', kwargs={'evento_id': self.evento.pk})

        post_dispensa = self.client.post(url, {'acao_participante': f'dispensar:{viajante.pk}'})
        self.assertEqual(post_dispensa.status_code, 302)
        termo = EventoTermoParticipante.objects.get(evento=self.evento, viajante=viajante)
        self.assertEqual(termo.status, EventoTermoParticipante.STATUS_DISPENSADO)

        post_reabrir = self.client.post(url, {'acao_participante': f'reabrir:{viajante.pk}'})
        self.assertEqual(post_reabrir.status_code, 302)
        termo.refresh_from_db()
        self.assertEqual(termo.status, EventoTermoParticipante.STATUS_PENDENTE)

    def test_etapa_5_download_docx_nao_exige_assinatura_configurada(self):
        viajante = self._criar_viajante(nome='Servidor Completo', completo=True, cpf='98765432100')
        self._vincular_viajante_em_oficio(viajante)
        url = reverse(
            'eventos:guiado-etapa-3-termo-download',
            kwargs={'evento_id': self.evento.pk, 'viajante_id': viajante.pk, 'formato': 'docx'},
        )

        response = self.client.get(url)

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response['Content-Type'],
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        termo = EventoTermoParticipante.objects.get(evento=self.evento, viajante=viajante)
        self.assertEqual(termo.status, EventoTermoParticipante.STATUS_GERADO)

    def test_etapa_5_download_termo_padrao_branco_docx_sem_servidor(self):
        url = reverse(
            'eventos:guiado-etapa-3-termo-padrao-download',
            kwargs={'evento_id': self.evento.pk, 'formato': 'docx'},
        )

        response = self.client.get(url)

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response['Content-Type'],
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        self.assertIn('padrao_branco.docx', response['Content-Disposition'])

    def test_etapa_5_download_termo_padrao_branco_pdf_quando_backend_disponivel(self):
        url = reverse(
            'eventos:guiado-etapa-3-termo-padrao-download',
            kwargs={'evento_id': self.evento.pk, 'formato': 'pdf'},
        )

        with patch('eventos.views.convert_docx_bytes_to_pdf_bytes', return_value=b'%PDF-1.4 termo branco'):
            response = self.client.get(url)

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response['Content-Type'], 'application/pdf')
        self.assertTrue(response.content.startswith(b'%PDF-1.4'))
        self.assertIn('padrao_branco.pdf', response['Content-Disposition'])

    def test_etapa_5_modo_viatura_lote_gera_zip_individual_por_servidor(self):
        veiculo = self._criar_veiculo_finalizado()
        viajante1 = self._criar_viajante(nome='Servidor ZIP 1', completo=True, cpf='30020010001')
        viajante2 = self._criar_viajante(nome='Servidor ZIP 2', completo=True, cpf='30020010002')
        url = reverse(
            'eventos:guiado-etapa-3-termo-viatura-download',
            kwargs={'evento_id': self.evento.pk, 'formato': 'docx'},
        )

        response = self.client.post(
            url,
            {
                'veiculo_id': str(veiculo.pk),
                'viajantes': [str(viajante1.pk), str(viajante2.pk)],
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response['Content-Type'], 'application/zip')
        with ZipFile(BytesIO(response.content)) as zip_file:
            names = zip_file.namelist()
        self.assertEqual(len(names), 2)
        self.assertTrue(all(name.endswith('.docx') for name in names))

        termo1 = EventoTermoParticipante.objects.get(evento=self.evento, viajante=viajante1)
        termo2 = EventoTermoParticipante.objects.get(evento=self.evento, viajante=viajante2)
        self.assertEqual(termo1.modalidade, EventoTermoParticipante.MODALIDADE_COMPLETO)
        self.assertEqual(termo2.modalidade, EventoTermoParticipante.MODALIDADE_COMPLETO)
        self.assertEqual(termo1.status, EventoTermoParticipante.STATUS_GERADO)
        self.assertEqual(termo2.status, EventoTermoParticipante.STATUS_GERADO)

    def test_etapa_5_download_docx_semipreenchido_funciona_sem_dados_pessoais(self):
        viajante = self._criar_viajante(nome='Servidor Sem Dados', completo=False, cpf='')
        self._vincular_viajante_em_oficio(viajante)
        EventoTermoParticipante.objects.update_or_create(
            evento=self.evento,
            viajante=viajante,
            defaults={'modalidade': EventoTermoParticipante.MODALIDADE_SEMIPREENCHIDO},
        )
        url = reverse(
            'eventos:guiado-etapa-3-termo-download',
            kwargs={'evento_id': self.evento.pk, 'viajante_id': viajante.pk, 'formato': 'docx'},
        )

        response = self.client.get(url)

        self.assertEqual(response.status_code, 200)
        termo = EventoTermoParticipante.objects.get(evento=self.evento, viajante=viajante)
        self.assertEqual(termo.status, EventoTermoParticipante.STATUS_GERADO)

    def test_etapa_5_download_pdf_quando_backend_disponivel(self):
        viajante = self._criar_viajante(nome='Servidor PDF', completo=True, cpf='11122233399')
        self._vincular_viajante_em_oficio(viajante)
        url = reverse(
            'eventos:guiado-etapa-3-termo-download',
            kwargs={'evento_id': self.evento.pk, 'viajante_id': viajante.pk, 'formato': 'pdf'},
        )

        with patch('eventos.views.convert_docx_bytes_to_pdf_bytes', return_value=b'%PDF-1.4 termo'):
            response = self.client.get(url)

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response['Content-Type'], 'application/pdf')
        self.assertTrue(response.content.startswith(b'%PDF-1.4'))
        termo = EventoTermoParticipante.objects.get(evento=self.evento, viajante=viajante)
        self.assertEqual(termo.status, EventoTermoParticipante.STATUS_GERADO)
        self.assertEqual(termo.ultimo_formato_gerado, 'pdf')

    def test_etapa_5_funciona_com_multiplos_servidores_e_multiplos_oficios(self):
        viajante1 = self._criar_viajante(nome='Servidor 1', cpf='11122233344')
        viajante2 = self._criar_viajante(nome='Servidor 2', cpf='11122233345')
        oficio1 = self._vincular_viajante_em_oficio(viajante1, numero=1)
        oficio2 = self._vincular_viajante_em_oficio(viajante2, numero=2)
        oficio2.viajantes.add(viajante1)

        response = self.client.get(reverse('eventos:guiado-etapa-5', kwargs={'evento_id': self.evento.pk}))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, viajante1.nome)
        self.assertContains(response, viajante2.nome)
        self.assertContains(response, oficio1.numero_formatado)
        self.assertContains(response, oficio2.numero_formatado)

class EventoEtapa6FinalizacaoTest(TestCase):
    """Etapa 6 â€” FinalizaÃ§Ã£o: checklist, pendÃªncias, observaÃ§Ãµes, finalizar evento, painel."""

    def setUp(self):
        self.user = get_user_model().objects.create_user(username='u6', password='p6')
        self.client = Client()
        self.client.login(username='u6', password='p6')
        self.evento = Evento.objects.create(
            titulo='Evento Etapa 6',
            data_inicio=date(2025, 1, 1),
            data_fim=date(2025, 1, 5),
            status=Evento.STATUS_RASCUNHO,
        )

    def test_etapa_6_exige_login(self):
        self.client.logout()
        response = self.client.get(reverse('eventos:guiado-etapa-7', kwargs={'evento_id': self.evento.pk}))
        self.assertEqual(response.status_code, 302)
        self.assertIn('login', response.url)

    def test_etapa_6_get_exibe_checklist_e_pendencias(self):
        response = self.client.get(reverse('eventos:guiado-etapa-7', kwargs={'evento_id': self.evento.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'SituaÃ§Ã£o das etapas')
        self.assertContains(response, 'Etapa 4')
        self.assertContains(response, 'Etapa 5')
        self.assertContains(response, 'Etapa 6')
        self.assertContains(response, 'PendÃªncias para finalizar')

    def test_etapa_6_post_salva_observacoes(self):
        response = self.client.post(
            reverse('eventos:guiado-etapa-7', kwargs={'evento_id': self.evento.pk}),
            {'observacoes_finais': 'Obs finais do evento.'},
        )
        self.assertEqual(response.status_code, 302)
        fin = EventoFinalizacao.objects.get(evento=self.evento)
        self.assertEqual((fin.observacoes_finais or '').strip(), 'Obs finais do evento.')
        self.assertIsNone(fin.finalizado_em)

    def test_etapa_6_autosave_persiste_observacoes(self):
        response = self.client.post(
            reverse('eventos:guiado-finalizacao', kwargs={'evento_id': self.evento.pk}),
            {
                'autosave': '1',
                'observacoes_finais': 'ObservaÃ§Ã£o salva automaticamente.',
            },
            HTTP_X_REQUESTED_WITH='XMLHttpRequest',
        )

        self.assertEqual(response.status_code, 200)
        self.assertTrue(response.json().get('ok'))
        fin = EventoFinalizacao.objects.get(evento=self.evento)
        self.assertEqual((fin.observacoes_finais or '').strip(), 'ObservaÃ§Ã£o salva automaticamente.')

    def test_etapa_6_finalizar_com_criterios_atendidos(self):
        tipo = TipoDemandaEvento.objects.create(nome='AÃ‡ÃƒO TESTE FINALIZAR', ordem=999, ativo=True, is_outros=False)
        estado = Estado.objects.create(nome='ParanÃ¡', sigla='PR', codigo_ibge='41')
        cidade = Cidade.objects.create(nome='Curitiba', estado=estado, codigo_ibge='4106902')
        etapa1_resp = self.client.post(
            reverse('eventos:guiado-etapa-1', kwargs={'pk': self.evento.pk}),
            {
                'tipos_demanda': [tipo.pk],
                'data_inicio': '2025-01-01',
                'data_fim': '2025-01-05',
                'descricao': '',
                'tem_convite_ou_oficio_evento': False,
                'destino_estado_0': estado.pk,
                'destino_cidade_0': cidade.pk,
            },
        )
        self.assertEqual(etapa1_resp.status_code, 302)
        PlanoTrabalho.objects.create(
            evento=self.evento,
            status=PlanoTrabalho.STATUS_FINALIZADO,
            objetivo='PT finalizado',
        )
        roteiro = RoteiroEvento.objects.create(evento=self.evento)
        RoteiroEvento.objects.filter(pk=roteiro.pk).update(status=RoteiroEvento.STATUS_FINALIZADO)
        viajante = Viajante.objects.create(nome='V', cpf='11122233344')
        oficio = Oficio.objects.create(
            evento=self.evento,
            status=Oficio.STATUS_FINALIZADO,
            data_criacao=date(2024, 12, 1),
        )
        oficio.viajantes.add(viajante)
        EventoTermoParticipante.objects.create(
            evento=self.evento, viajante=viajante, status=EventoTermoParticipante.STATUS_CONCLUIDO
        )
        response = self.client.post(
            reverse('eventos:guiado-etapa-7', kwargs={'evento_id': self.evento.pk}),
            {'observacoes_finais': '', 'finalizar': '1'},
        )
        self.assertEqual(response.status_code, 302)
        self.assertIn('guiado-finalizacao', response.url)
        fin = EventoFinalizacao.objects.get(evento=self.evento)
        self.assertIsNotNone(fin.finalizado_em)
        self.assertEqual(fin.finalizado_por_id, self.user.pk)
        self.evento.refresh_from_db()
        self.assertEqual(self.evento.status, Evento.STATUS_FINALIZADO)

class EventoFinalizadoTravasTest(TestCase):
    """Travas pÃ³s-finalizaÃ§Ã£o: bloquear exclusÃ£o, mas permitir ediÃ§Ã£o do evento e dos ofÃ­cios."""

    def setUp(self):
        self.user = get_user_model().objects.create_user(username='u_trava', password='p_trava')
        self.client = Client()
        self.client.login(username='u_trava', password='p_trava')
        self.evento = Evento.objects.create(
            titulo='Evento Finalizado',
            data_inicio=date(2025, 1, 1),
            data_fim=date(2025, 1, 5),
            status=Evento.STATUS_FINALIZADO,
        )

    def test_excluir_evento_finalizado_bloqueado(self):
        response = self.client.post(
            reverse('eventos:excluir', kwargs={'pk': self.evento.pk}),
            {},
            follow=True,
        )
        self.assertEqual(response.status_code, 200)
        self.assertTrue(Evento.objects.filter(pk=self.evento.pk).exists())
        self.assertContains(response, 'finalizado')

    def test_get_etapa_1_finalizado_retorna_200_consulta(self):
        response = self.client.get(reverse('eventos:guiado-etapa-1', kwargs={'pk': self.evento.pk}))
        self.assertEqual(response.status_code, 200)

    def test_post_etapa_1_finalizado_bloqueado(self):
        response = self.client.post(
            reverse('eventos:guiado-etapa-1', kwargs={'pk': self.evento.pk}),
            {'data_inicio': '2025-01-10', 'data_fim': '2025-01-15', 'tipos_demanda': []},
        )
        # Continua sujeito Ã s validaÃ§Ãµes normais do formulÃ¡rio,
        # mas nÃ£o deve ser bloqueado apenas por o evento estar finalizado.
        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, 'Evento finalizado. NÃ£o Ã© possÃ­vel alterar os dados do evento.')

    def test_post_etapa_4_finalizado_bloqueado(self):
        PlanoTrabalho.objects.create(evento=self.evento)
        before_pt = PlanoTrabalho.objects.filter(evento=self.evento).count()
        before_os = OrdemServico.objects.filter(evento=self.evento).count()
        response = self.client.post(
            reverse('eventos:guiado-etapa-4', kwargs={'evento_id': self.evento.pk}),
            {'qualquer': 'valor'},
        )
        self.assertEqual(response.status_code, 302)
        self.assertEqual(PlanoTrabalho.objects.filter(evento=self.evento).count(), before_pt)
        self.assertEqual(OrdemServico.objects.filter(evento=self.evento).count(), before_os)

    def test_get_etapa_5_finalizado_retorna_200_consulta(self):
        viajante = Viajante.objects.create(nome='V', cpf='11122233344')
        oficio = Oficio.objects.create(evento=self.evento, status=Oficio.STATUS_RASCUNHO)
        oficio.viajantes.add(viajante)
        EventoTermoParticipante.objects.create(
            evento=self.evento, viajante=viajante, status=EventoTermoParticipante.STATUS_PENDENTE,
        )
        response = self.client.get(reverse('eventos:guiado-etapa-5', kwargs={'evento_id': self.evento.pk}))
        self.assertEqual(response.status_code, 200)
        termo = self.evento.termos_participantes.get(viajante=viajante)
        self.assertEqual(termo.status, EventoTermoParticipante.STATUS_PENDENTE)

    def test_download_lote_viatura_em_evento_finalizado_nao_permite_novos_participantes(self):
        combustivel = CombustivelVeiculo.objects.create(nome='Diesel', is_padrao=True)
        veiculo = Veiculo.objects.create(
            placa='DEF2G34',
            modelo='Viatura Finalizada',
            combustivel=combustivel,
            tipo=Veiculo.TIPO_DESCARACTERIZADO,
            status=Veiculo.STATUS_FINALIZADO,
        )
        viajante_existente = Viajante.objects.create(nome='V Existente', cpf='99988877766')
        viajante_novo = Viajante.objects.create(nome='V Novo', cpf='88877766655')
        EventoParticipante.objects.create(evento=self.evento, viajante=viajante_existente, ordem=0)
        EventoTermoParticipante.objects.create(
            evento=self.evento,
            viajante=viajante_existente,
            status=EventoTermoParticipante.STATUS_PENDENTE,
        )

        response = self.client.post(
            reverse('eventos:guiado-etapa-3-termo-viatura-download', kwargs={'evento_id': self.evento.pk, 'formato': 'docx'}),
            {'veiculo_id': str(veiculo.pk), 'viajantes': [str(viajante_existente.pk), str(viajante_novo.pk)]},
        )

        self.assertEqual(response.status_code, 302)
        self.assertIn('guiado/termos', response.url)
        self.assertFalse(
            EventoParticipante.objects.filter(evento=self.evento, viajante=viajante_novo).exists()
        )

    def test_oficio_excluir_bloqueado_quando_evento_finalizado(self):
        oficio = Oficio.objects.create(evento=self.evento, status=Oficio.STATUS_RASCUNHO)
        response = self.client.post(
            reverse('eventos:oficio-excluir', kwargs={'pk': oficio.pk}),
            {},
            follow=True,
        )
        self.assertEqual(response.status_code, 200)
        self.assertTrue(Oficio.objects.filter(pk=oficio.pk).exists())
        self.assertContains(response, 'finalizado')

    def test_oficio_step1_post_bloqueado_quando_evento_finalizado(self):
        oficio = Oficio.objects.create(evento=self.evento, status=Oficio.STATUS_RASCUNHO)
        response = self.client.post(
            reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk}),
            {'protocolo': 'X', 'motivo': 'Teste', 'custeio_tipo': Oficio.CUSTEIO_UNIDADE, 'viajantes': []},
            follow=True,
        )
        self.assertEqual(response.status_code, 200)
        # NÃ£o deve haver mensagem de bloqueio apenas por o evento estar finalizado.
        self.assertNotContains(response, 'Evento finalizado. NÃ£o Ã© possÃ­vel editar ofÃ­cios vinculados.')


class EventoEtapa3OficiosTest(TestCase):
    """Etapa 5 â€” OfÃ­cios do evento (hub): listar, criar, status OK/Pendente. URL: guiado-etapa-5."""

    def setUp(self):
        self.user = User.objects.create_user(username='u', password='p')
        self.client = Client()
        self.client.login(username='u', password='p')
        self.estado = Estado.objects.create(nome='ParanÃ¡', sigla='PR', codigo_ibge='41')
        self.cidade = Cidade.objects.create(nome='Curitiba', estado=self.estado, codigo_ibge='4106902')
        self.evento = Evento.objects.create(
            titulo='Evento Etapa 3',
            data_inicio=date(2025, 1, 1),
            data_fim=date(2025, 1, 5),
            status=Evento.STATUS_RASCUNHO,
        )

    def test_etapa_3_exige_login(self):
        self.client.logout()
        response = self.client.get(reverse('eventos:guiado-etapa-5', kwargs={'evento_id': self.evento.pk}))
        self.assertEqual(response.status_code, 302)
        self.assertIn('login', response.url)

    def test_etapa_3_lista_oficios_do_evento(self):
        oficio1 = Oficio.objects.create(evento=self.evento, status=Oficio.STATUS_RASCUNHO)
        oficio2 = Oficio.objects.create(evento=self.evento, numero=2, ano=2025, status=Oficio.STATUS_FINALIZADO)
        response = self.client.get(reverse('eventos:guiado-etapa-5', kwargs={'evento_id': self.evento.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'OfÃ­cios do evento')
        self.assertContains(response, 'Rascunho')
        self.assertContains(response, 'Finalizado')
        self.assertContains(response, '02/2025')
        self.assertContains(response, reverse('eventos:oficio-editar', kwargs={'pk': oficio1.pk}))

    def test_etapa_3_mostra_botao_criar_oficio(self):
        response = self.client.get(reverse('eventos:guiado-etapa-5', kwargs={'evento_id': self.evento.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Criar OfÃ­cio neste Evento')
        url_criar = reverse('eventos:guiado-etapa-5-criar-oficio', kwargs={'evento_id': self.evento.pk})
        self.assertContains(response, url_criar)

    def test_get_criar_oficio_nao_cria_registro(self):
        url_criar = reverse('eventos:guiado-etapa-5-criar-oficio', kwargs={'evento_id': self.evento.pk})
        response = self.client.get(url_criar)
        self.assertEqual(response.status_code, 302)
        self.assertEqual(response.url, reverse('eventos:guiado-etapa-5', kwargs={'evento_id': self.evento.pk}))
        self.assertEqual(self.evento.oficios.count(), 0)

    def test_post_criar_oficio_preserva_vinculo_com_evento(self):
        url_criar = reverse('eventos:guiado-etapa-5-criar-oficio', kwargs={'evento_id': self.evento.pk})
        response = self.client.post(url_criar)
        self.assertEqual(response.status_code, 302)
        self.assertIn('oficio/', response.url)
        self.assertIn('/step1/', response.url)
        self.assertEqual(self.evento.oficios.count(), 1)
        oficio = self.evento.oficios.get()
        self.assertEqual(oficio.evento_id, self.evento.pk)
        self.assertEqual(oficio.status, Oficio.STATUS_RASCUNHO)

class OficioWizardTest(TestCase):
    """Wizard do OfÃ­cio: Step 1 (viajantes), Step 2 (transporte/motorista), fluxo."""

    def setUp(self):
        self.user = User.objects.create_user(username='u2', password='p2')
        self.client = Client()
        self.client.login(username='u2', password='p2')
        self.estado = Estado.objects.create(nome='ParanÃ¡', sigla='PR', codigo_ibge='41')
        self.cidade = Cidade.objects.create(nome='Curitiba', estado=self.estado, codigo_ibge='4106902')
        self.evento = Evento.objects.create(
            titulo='Evento Teste',
            data_inicio=date(2025, 1, 1),
            data_fim=date(2025, 1, 5),
            status=Evento.STATUS_RASCUNHO,
        )
        self.cargo = Cargo.objects.create(nome='ANALISTA', is_padrao=True)
        self.unidade = UnidadeLotacao.objects.create(nome='DPC')
        self.viajante = Viajante.objects.create(
            nome='Fulano Teste',
            status=Viajante.STATUS_FINALIZADO,
            cargo=self.cargo,
            cpf='12345678901',
            telefone='41999999999',
            unidade_lotacao=self.unidade,
            rg='1234567890',
        )
        self.combustivel = CombustivelVeiculo.objects.create(nome='GASOLINA', is_padrao=True)
        self.oficio = Oficio.objects.create(evento=self.evento, status=Oficio.STATUS_RASCUNHO)

    def test_wizard_step1_exige_ao_menos_um_viajante(self):
        url = reverse('eventos:oficio-step1', kwargs={'pk': self.oficio.pk})
        get_resp = self.client.get(url)
        self.assertEqual(get_resp.status_code, 200)
        csrf = get_resp.context.get('csrf_token', '')
        if hasattr(csrf, '__str__'):
            csrf = str(csrf)
        self.oficio.refresh_from_db()
        post_resp = self.client.post(url, data={
            'oficio_numero': self.oficio.numero_formatado,
            'protocolo': '12.345.678-9',
            'data_criacao': self.oficio.data_criacao.strftime('%d/%m/%Y'),
            'motivo': 'Motivo teste',
            'assunto_tipo': Oficio.ASSUNTO_TIPO_AUTORIZACAO,
            'custeio_tipo': Oficio.CUSTEIO_UNIDADE,
            'viajantes': [],  # nenhum viajante
            'csrfmiddlewaretoken': csrf,
        })
        self.assertEqual(post_resp.status_code, 200)
        self.assertFormError(post_resp.context['form'], 'viajantes', 'Selecione ao menos um viajante.')

    def test_wizard_step1_salva_viajantes(self):
        url = reverse('eventos:oficio-step1', kwargs={'pk': self.oficio.pk})
        get_resp = self.client.get(url)
        csrf = str(get_resp.context['csrf_token']) if get_resp.context.get('csrf_token') else ''
        self.oficio.refresh_from_db()
        post_resp = self.client.post(url, data={
            'oficio_numero': self.oficio.numero_formatado,
            'protocolo': '12.345.678-9',
            'data_criacao': self.oficio.data_criacao.strftime('%d/%m/%Y'),
            'motivo': 'Motivo',
            'assunto_tipo': Oficio.ASSUNTO_TIPO_AUTORIZACAO,
            'custeio_tipo': Oficio.CUSTEIO_UNIDADE,
            'viajantes': [self.viajante.pk],
            'csrfmiddlewaretoken': csrf,
        })
        self.assertEqual(post_resp.status_code, 302)
        self.oficio.refresh_from_db()
        self.assertEqual(list(self.oficio.viajantes.values_list('pk', flat=True)), [self.viajante.pk])
        self.assertIsNotNone(self.oficio.numero)
        self.assertEqual(self.oficio.ano, tz.localdate().year)

    def test_wizard_step2_exige_placa_modelo_combustivel(self):
        url = reverse('eventos:oficio-step2', kwargs={'pk': self.oficio.pk})
        get_resp = self.client.get(url)
        csrf = str(get_resp.context['csrf_token']) if get_resp.context.get('csrf_token') else ''
        post_resp = self.client.post(url, data={
            'placa': '',
            'modelo': '',
            'combustivel': '',
            'tipo_viatura': Oficio.TIPO_VIATURA_DESCARACTERIZADA,
            'motorista_carona': False,
            'csrfmiddlewaretoken': csrf,
        })
        self.assertEqual(post_resp.status_code, 200)
        form = post_resp.context['form']
        self.assertTrue(form.errors.get('placa'))
        self.assertTrue(form.errors.get('modelo'))
        self.assertTrue(form.errors.get('combustivel'))

    def test_wizard_step2_motorista_carona_exige_oficio_protocolo(self):
        url = reverse('eventos:oficio-step2', kwargs={'pk': self.oficio.pk})
        get_resp = self.client.get(url)
        csrf = str(get_resp.context['csrf_token']) if get_resp.context.get('csrf_token') else ''
        post_resp = self.client.post(url, data={
            'placa': 'ABC1234',
            'modelo': 'Modelo X',
            'combustivel': 'GASOLINA',
            'tipo_viatura': Oficio.TIPO_VIATURA_DESCARACTERIZADA,
            'motorista_viajante': '__manual__',
            'motorista_nome': 'Motorista Externo',
            'motorista_oficio_numero': '',
            'motorista_oficio_ano': str(tz.localdate().year),
            'motorista_protocolo': '',
            'csrfmiddlewaretoken': csrf,
        })
        self.assertEqual(post_resp.status_code, 200)
        form = post_resp.context['form']
        self.assertTrue(form.errors.get('motorista_oficio_numero') or form.errors.get('motorista_protocolo'))

    def test_oficio_model_exige_motorista_protocolo_9_digitos_quando_carona(self):
        """Model Oficio.clean() deve exigir motorista_protocolo com 9 dÃ­gitos quando motorista_carona=True."""
        self.oficio.motorista_carona = True
        self.oficio.motorista_protocolo = '12345678'  # 8 dÃ­gitos
        with self.assertRaises(ValidationError) as ctx:
            self.oficio.full_clean()
        self.assertIn('motorista_protocolo', ctx.exception.message_dict)
        self.oficio.motorista_protocolo = '123456789'
        self.oficio.full_clean()  # nÃ£o deve levantar

    def test_wizard_step2_motorista_sem_cadastro_exige_nome_manual(self):
        url = reverse('eventos:oficio-step2', kwargs={'pk': self.oficio.pk})
        response = self.client.post(url, data={
            'placa': 'ABC1234',
            'modelo': 'Modelo X',
            'combustivel': 'GASOLINA',
            'tipo_viatura': Oficio.TIPO_VIATURA_DESCARACTERIZADA,
            'motorista_viajante': '__manual__',
            'motorista_nome': '',
            'motorista_oficio_numero': '1',
            'motorista_oficio_ano': str(tz.localdate().year),
            'motorista_protocolo': '12.345.678-9',
        })
        self.assertEqual(response.status_code, 200)
        self.assertFormError(
            response.context['form'],
            'motorista_nome',
            'Informe o nome do motorista sem cadastro.',
        )

    def test_editar_oficio_redireciona_para_step1(self):
        response = self.client.get(reverse('eventos:oficio-editar', kwargs={'pk': self.oficio.pk}))
        self.assertEqual(response.status_code, 302)
        self.assertIn('/step1/', response.url)

    def test_oficio_step4_finalizar_marca_finalizado(self):
        data_saida = tz.localdate() + timedelta(days=20)
        data_retorno = data_saida + timedelta(days=1)
        self.client.post(
            reverse('eventos:oficio-step1', kwargs={'pk': self.oficio.pk}),
            data={
                'oficio_numero': self.oficio.numero_formatado,
                'protocolo': '12.345.678-9',
                'data_criacao': self.oficio.data_criacao.strftime('%d/%m/%Y'),
                'motivo': 'Motivo teste',
                'assunto_tipo': Oficio.ASSUNTO_TIPO_AUTORIZACAO,
                'custeio_tipo': Oficio.CUSTEIO_UNIDADE,
                'viajantes': [self.viajante.pk],
            },
        )
        self.client.post(
            reverse('eventos:oficio-step2', kwargs={'pk': self.oficio.pk}),
            data={
                'placa': 'ABC1234',
                'modelo': 'Modelo X',
                'combustivel': 'GASOLINA',
                'tipo_viatura': Oficio.TIPO_VIATURA_DESCARACTERIZADA,
                'motorista_viajante': '__manual__',
                'motorista_nome': 'Motorista Externo',
                'motorista_oficio_numero': '1',
                'motorista_oficio_ano': str(tz.localdate().year),
                'motorista_protocolo': '12.345.678-9',
            },
        )
        cidade_destino = Cidade.objects.create(nome='Destino Wizard', estado=self.estado, codigo_ibge='4106999')
        self.client.post(
            reverse('eventos:oficio-step3', kwargs={'pk': self.oficio.pk}),
            data={
                'roteiro_modo': Oficio.ROTEIRO_MODO_PROPRIO,
                'sede_estado': str(self.estado.pk),
                'sede_cidade': str(self.cidade.pk),
                'destino_estado_0': str(self.estado.pk),
                'destino_cidade_0': str(cidade_destino.pk),
                'trecho_0_saida_data': data_saida.strftime('%Y-%m-%d'),
                'trecho_0_saida_hora': '08:00',
                'trecho_0_chegada_data': data_saida.strftime('%Y-%m-%d'),
                'trecho_0_chegada_hora': '12:00',
                'retorno_saida_data': data_retorno.strftime('%Y-%m-%d'),
                'retorno_saida_hora': '09:00',
                'retorno_chegada_data': data_retorno.strftime('%Y-%m-%d'),
                'retorno_chegada_hora': '18:00',
            },
        )
        self.oficio.refresh_from_db()
        before_update = self.oficio.updated_at
        url = reverse('eventos:oficio-step4', kwargs={'pk': self.oficio.pk})
        get_resp = self.client.get(url)
        csrf = str(get_resp.context['csrf_token']) if get_resp.context.get('csrf_token') else ''
        post_resp = self.client.post(url, data={'finalizar': '1', 'csrfmiddlewaretoken': csrf})
        self.assertEqual(post_resp.status_code, 302)
        self.oficio.refresh_from_db()
        self.assertEqual(self.oficio.status, Oficio.STATUS_FINALIZADO)
        self.assertGreater(self.oficio.updated_at, before_update)


class OficioStep1AcceptanceTest(TestCase):
    """Aceite do Step 1 do OfÃ­cio (fidelidade ao legado)."""

    def setUp(self):
        self.user = User.objects.create_user(username='step1', password='step1')
        self.client = Client()
        self.client.login(username='step1', password='step1')
        self.estado = Estado.objects.create(nome='ParanÃ¡', sigla='PR', codigo_ibge='41')
        self.cidade = Cidade.objects.create(nome='Curitiba', estado=self.estado, codigo_ibge='4106902')
        self.evento = Evento.objects.create(
            titulo='Evento Step 1',
            data_inicio=date(2026, 1, 10),
            data_fim=date(2026, 1, 12),
            status=Evento.STATUS_RASCUNHO,
        )
        self.cargo = Cargo.objects.create(nome='ANALISTA STEP1', is_padrao=True)
        self.unidade = UnidadeLotacao.objects.create(nome='DPC STEP1')
        self.viajante_final = Viajante.objects.create(
            nome='Viajante Finalizado',
            status=Viajante.STATUS_FINALIZADO,
            cargo=self.cargo,
            cpf='11122233344',
            telefone='41999990001',
            unidade_lotacao=self.unidade,
            rg='1234567890',
        )
        self.viajante_final_2 = Viajante.objects.create(
            nome='Outro Viajante Finalizado',
            status=Viajante.STATUS_FINALIZADO,
            cargo=self.cargo,
            cpf='22233344455',
            telefone='41999990004',
            unidade_lotacao=self.unidade,
            rg='2234567890',
        )
        self.viajante_rascunho = Viajante.objects.create(
            nome='Viajante Rascunho',
            status=Viajante.STATUS_RASCUNHO,
            cargo=self.cargo,
            cpf='55566677788',
            telefone='41999990002',
            unidade_lotacao=self.unidade,
            rg='1234567891',
        )
        self.motorista_externo = Viajante.objects.create(
            nome='Motorista Externo Cadastro',
            status=Viajante.STATUS_FINALIZADO,
            cargo=self.cargo,
            cpf='99988877766',
            telefone='41999990003',
            unidade_lotacao=self.unidade,
            rg='1234567892',
        )
        self.combustivel = CombustivelVeiculo.objects.create(nome='GASOLINA', is_padrao=True)

    def _criar_oficio(self, ano=None):
        kwargs = {'evento': self.evento, 'status': Oficio.STATUS_RASCUNHO}
        if ano is not None:
            kwargs['ano'] = ano
        return Oficio.objects.create(**kwargs)

    def _payload_step1(self, oficio, **overrides):
        oficio.refresh_from_db()
        data = {
            'oficio_numero': oficio.numero_formatado,
            'protocolo': '12.345.678-9',
            'data_criacao': oficio.data_criacao.strftime('%d/%m/%Y'),
            'modelo_motivo': '',
            'motivo': 'Motivo base',
            'assunto_tipo': Oficio.ASSUNTO_TIPO_AUTORIZACAO,
            'custeio_tipo': Oficio.CUSTEIO_UNIDADE,
            'nome_instituicao_custeio': '',
            'viajantes': [self.viajante_final.pk],
        }
        data.update(overrides)
        return data

    def _payload_step2(self, **overrides):
        data = {
            'placa': 'ABC-1234',
            'modelo': 'VIATURA TESTE',
            'combustivel': 'GASOLINA',
            'tipo_viatura': Oficio.TIPO_VIATURA_DESCARACTERIZADA,
            'porte_transporte_armas': '1',
            'motorista_viajante': '__manual__',
            'motorista_nome': 'Motorista Externo',
            'motorista_oficio_numero': '7',
            'motorista_oficio_ano': str(tz.localdate().year),
            'motorista_protocolo': '12.345.678-9',
        }
        data.update(overrides)
        return data

    def _payload_step3(self, destinos, trechos=None, retorno=None, **overrides):
        data = {
            'roteiro_modo': Oficio.ROTEIRO_MODO_PROPRIO,
            'sede_estado': str(self.estado.pk),
            'sede_cidade': str(self.cidade.pk),
        }
        for idx, cidade in enumerate(destinos):
            data[f'destino_estado_{idx}'] = str(cidade.estado_id)
            data[f'destino_cidade_{idx}'] = str(cidade.pk)
        trechos = trechos or []
        for idx, trecho in enumerate(trechos):
            for key, value in trecho.items():
                data[f'trecho_{idx}_{key}'] = value
        retorno = retorno or {}
        data.update(
            {
                'retorno_saida_data': retorno.get('saida_data', ''),
                'retorno_saida_hora': retorno.get('saida_hora', ''),
                'retorno_chegada_data': retorno.get('chegada_data', ''),
                'retorno_chegada_hora': retorno.get('chegada_hora', ''),
            }
        )
        data.update(overrides)
        return data

    def _criar_destino(self, nome='Destino Finalizacao'):
        codigo = str(4107000 + Cidade.objects.count())
        return Cidade.objects.create(nome=nome, estado=self.estado, codigo_ibge=codigo)

    def _salvar_steps_1_e_2(self, oficio, step1_overrides=None, step2_overrides=None):
        step1_overrides = step1_overrides or {}
        step2_overrides = step2_overrides or {}
        response_step1 = self.client.post(
            reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk}),
            data=self._payload_step1(oficio, **step1_overrides),
        )
        self.assertEqual(response_step1.status_code, 302)
        response_step2 = self.client.post(
            reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}),
            data=self._payload_step2(**step2_overrides),
        )
        self.assertEqual(response_step2.status_code, 302)

    def _salvar_oficio_finalizavel(self, oficio, destino=None, step1_overrides=None, step2_overrides=None, step3_overrides=None):
        self._salvar_steps_1_e_2(oficio, step1_overrides=step1_overrides, step2_overrides=step2_overrides)
        destino = destino or self._criar_destino()
        payload_step3 = self._payload_step3(
            [destino],
            trechos=[
                {
                    'saida_data': '2026-10-10',
                    'saida_hora': '08:00',
                    'chegada_data': '2026-10-10',
                    'chegada_hora': '12:00',
                }
            ],
            retorno={
                'saida_data': '2026-10-11',
                'saida_hora': '08:30',
                'chegada_data': '2026-10-11',
                'chegada_hora': '16:30',
            },
            **(step3_overrides or {}),
        )
        response_step3 = self.client.post(
            reverse('eventos:oficio-step3', kwargs={'pk': oficio.pk}),
            data=payload_step3,
        )
        self.assertEqual(response_step3.status_code, 302)
        oficio.refresh_from_db()
        return destino

    def _post_finalizar_oficio(self, oficio):
        return self.client.post(
            reverse('eventos:oficio-step4', kwargs={'pk': oficio.pk}),
            data={'finalizar': '1'},
        )

    def _criar_roteiro_evento_base(self, destinos, data_base=None):
        data_base = data_base or datetime(2026, 2, 10, 8, 0)
        if tz.is_naive(data_base):
            data_base = tz.make_aware(data_base)
        self.evento.cidade_base = self.cidade
        self.evento.save(update_fields=['cidade_base'])
        roteiro = RoteiroEvento.objects.create(
            evento=self.evento,
            origem_estado=self.estado,
            origem_cidade=self.cidade,
            saida_dt=data_base,
            retorno_saida_dt=data_base + timedelta(days=1, hours=1),
            retorno_chegada_dt=data_base + timedelta(days=1, hours=4),
            status=RoteiroEvento.STATUS_FINALIZADO,
        )
        origem_estado_id = self.estado.pk
        origem_cidade_id = self.cidade.pk
        for ordem, cidade_destino in enumerate(destinos):
            RoteiroEventoDestino.objects.create(
                roteiro=roteiro,
                estado=self.estado,
                cidade=cidade_destino,
                ordem=ordem,
            )
            RoteiroEventoTrecho.objects.create(
                roteiro=roteiro,
                ordem=ordem,
                tipo=RoteiroEventoTrecho.TIPO_IDA,
                origem_estado_id=origem_estado_id,
                origem_cidade_id=origem_cidade_id,
                destino_estado_id=self.estado.pk,
                destino_cidade_id=cidade_destino.pk,
                saida_dt=data_base + timedelta(hours=ordem * 3),
                chegada_dt=data_base + timedelta(hours=ordem * 3 + 2),
            )
            origem_estado_id = self.estado.pk
            origem_cidade_id = cidade_destino.pk
        RoteiroEventoTrecho.objects.create(
            roteiro=roteiro,
            ordem=len(destinos),
            tipo=RoteiroEventoTrecho.TIPO_RETORNO,
            origem_estado_id=origem_estado_id,
            origem_cidade_id=origem_cidade_id,
            destino_estado_id=self.estado.pk,
            destino_cidade_id=self.cidade.pk,
            saida_dt=data_base + timedelta(days=1, hours=1),
            chegada_dt=data_base + timedelta(days=1, hours=4),
        )
        roteiro.chegada_dt = data_base + timedelta(hours=((len(destinos) - 1) * 3) + 2)
        roteiro.save()
        return roteiro

    def _split_step2_response(self, response):
        html = response.content.decode('utf-8')
        if '<script>' not in html:
            self.fail('Template do Step 2 sem bloco de script esperado.')
        return html.split('<script>', 1)

    def test_1_novo_oficio_gera_numero_automatico_formato_xx_ano(self):
        oficio = self._criar_oficio(ano=2026)
        self.assertEqual(oficio.numero, 1)
        self.assertEqual(oficio.ano, 2026)
        self.assertEqual(oficio.numero_formatado, '01/2026')

    def test_2_sequencia_anual_funciona_e_reinicia_no_novo_ano(self):
        oficio_1 = self._criar_oficio(ano=2026)
        oficio_2 = self._criar_oficio(ano=2026)
        oficio_3 = self._criar_oficio(ano=2027)
        self.assertEqual(oficio_1.numero_formatado, '01/2026')
        self.assertEqual(oficio_2.numero_formatado, '02/2026')
        self.assertEqual(oficio_3.numero_formatado, '01/2027')

    def test_2b_criacao_usa_menor_numero_livre_do_ano(self):
        oficio_1 = self._criar_oficio(ano=2026)
        oficio_2 = self._criar_oficio(ano=2026)
        oficio_3 = self._criar_oficio(ano=2026)
        oficio_2.delete()
        oficio_novo = self._criar_oficio(ano=2026)
        self.assertEqual(oficio_1.numero_formatado, '01/2026')
        self.assertEqual(oficio_3.numero_formatado, '03/2026')
        self.assertEqual(oficio_novo.numero_formatado, '02/2026')

    def test_3_ao_editar_numero_nao_muda(self):
        oficio = self._criar_oficio(ano=2026)
        numero_antes = oficio.numero_formatado
        url = reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk})
        response = self.client.post(url, data=self._payload_step1(oficio))
        self.assertEqual(response.status_code, 302)
        oficio.refresh_from_db()
        self.assertEqual(oficio.numero_formatado, numero_antes)

    def test_4_e_5_protocolo_aplica_aceita_mascara_e_reabre_mascarado(self):
        oficio = self._criar_oficio(ano=2026)
        url = reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk})
        response = self.client.post(url, data=self._payload_step1(oficio, protocolo='12.345.678-9'))
        self.assertEqual(response.status_code, 302)
        oficio.refresh_from_db()
        self.assertEqual(oficio.protocolo, '123456789')
        reaberto = self.client.get(url)
        self.assertEqual(reaberto.status_code, 200)
        self.assertContains(reaberto, '12.345.678-9')

    def test_6_data_criacao_vem_preenchida_com_data_atual(self):
        oficio = self._criar_oficio()
        url = reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk})
        response = self.client.get(url)
        self.assertEqual(response.status_code, 200)
        oficio.refresh_from_db()
        self.assertEqual(oficio.data_criacao, tz.localdate())
        self.assertContains(response, tz.localdate().strftime('%d/%m/%Y'))

    def test_7_selecionar_modelo_preenche_motivo(self):
        oficio = self._criar_oficio(ano=2026)
        modelo = ModeloMotivoViagem.objects.create(
            codigo='modelo_step1',
            nome='Modelo Step 1',
            texto='Texto padrÃ£o do modelo',
            ordem=1,
            ativo=True,
        )
        url = reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk})
        response = self.client.post(url, data=self._payload_step1(oficio, modelo_motivo=modelo.pk, motivo=''))
        self.assertEqual(response.status_code, 302)
        oficio.refresh_from_db()
        self.assertEqual(oficio.motivo, 'Texto padrÃ£o do modelo')
        self.assertEqual(oficio.modelo_motivo_id, modelo.pk)

    def test_8_motivo_pode_ser_editado_manualmente(self):
        oficio = self._criar_oficio(ano=2026)
        modelo = ModeloMotivoViagem.objects.create(
            codigo='modelo_editavel',
            nome='Modelo EditÃ¡vel',
            texto='Texto original',
            ordem=2,
            ativo=True,
        )
        url = reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk})
        response = self.client.post(
            url,
            data=self._payload_step1(oficio, modelo_motivo=modelo.pk, motivo='Texto alterado manualmente'),
        )
        self.assertEqual(response.status_code, 302)
        oficio.refresh_from_db()
        self.assertEqual(oficio.motivo, 'Texto alterado manualmente')

    def test_9_salvar_motivo_como_novo_modelo_funciona(self):
        oficio = self._criar_oficio(ano=2026)
        url = reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk})
        response = self.client.post(
            url,
            data=self._payload_step1(
                oficio,
                motivo='Motivo para virar modelo',
                salvar_modelo_motivo='1',
                novo_modelo_nome='Novo Modelo Step1',
            ),
        )
        self.assertEqual(response.status_code, 302)
        modelo = ModeloMotivoViagem.objects.get(nome='Novo Modelo Step1')
        self.assertEqual(modelo.texto, 'Motivo para virar modelo')
        oficio.refresh_from_db()
        self.assertEqual(oficio.modelo_motivo_id, modelo.pk)

    def test_10_e_11_nome_instituicao_aparece_so_em_outra_instituicao_e_eh_obrigatorio(self):
        oficio = self._criar_oficio(ano=2026)
        url = reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk})

        get_resp = self.client.get(url)
        self.assertEqual(get_resp.status_code, 200)
        self.assertContains(get_resp, 'id="nome-instituicao-wrapper" style="display:none;"')

        post_resp = self.client.post(
            url,
            data=self._payload_step1(
                oficio,
                custeio_tipo=Oficio.CUSTEIO_OUTRA_INSTITUICAO,
                nome_instituicao_custeio='',
            ),
        )
        self.assertEqual(post_resp.status_code, 200)
        self.assertFormError(
            post_resp.context['form'],
            'nome_instituicao_custeio',
            'Informe a instituiÃ§Ã£o de custeio.',
        )

    def test_custeio_diferente_de_outra_instituicao_limpa_nome(self):
        oficio = self._criar_oficio(ano=2026)
        oficio.custeio_tipo = Oficio.CUSTEIO_OUTRA_INSTITUICAO
        oficio.nome_instituicao_custeio = 'InstituiÃ§Ã£o X'
        oficio.save(update_fields=['custeio_tipo', 'nome_instituicao_custeio'])
        url = reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk})
        response = self.client.post(
            url,
            data=self._payload_step1(
                oficio,
                custeio_tipo=Oficio.CUSTEIO_UNIDADE,
                nome_instituicao_custeio='Deve limpar',
            ),
        )
        self.assertEqual(response.status_code, 302)
        oficio.refresh_from_db()
        self.assertEqual(oficio.nome_instituicao_custeio, '')

    def test_12_viajantes_exige_ao_menos_um(self):
        oficio = self._criar_oficio(ano=2026)
        url = reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk})
        response = self.client.post(url, data=self._payload_step1(oficio, viajantes=[]))
        self.assertEqual(response.status_code, 200)
        self.assertFormError(response.context['form'], 'viajantes', 'Selecione ao menos um viajante.')

    def test_13_autocomplete_busca_por_nome_e_filtra_rascunho(self):
        response = self.client.get(
            reverse('eventos:oficio-step1-viajantes-api'),
            {'q': 'Finalizado'},
        )
        self.assertEqual(response.status_code, 200)
        payload = response.json()
        nomes = [item['nome'] for item in payload['results']]
        self.assertIn(self.viajante_final.nome, nomes)
        self.assertIn(self.viajante_final_2.nome, nomes)
        self.assertNotIn(self.viajante_rascunho.nome, nomes)

    def test_13b_autocomplete_busca_por_rg(self):
        response = self.client.get(
            reverse('eventos:oficio-step1-viajantes-api'),
            {'q': self.viajante_final.rg},
        )
        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload['results'][0]['id'], self.viajante_final.pk)

    def test_13c_autocomplete_busca_por_cpf_mascarado(self):
        response = self.client.get(
            reverse('eventos:oficio-step1-viajantes-api'),
            {'q': '111.222.333-44'},
        )
        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload['results'][0]['id'], self.viajante_final.pk)
        self.assertEqual(payload['results'][0]['cpf'], '111.222.333-44')

    def test_14_botao_cadastrar_novo_viajante_existe(self):
        oficio = self._criar_oficio(ano=2026)
        url = reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk})
        response = self.client.get(url)
        self.assertEqual(response.status_code, 200)
        expected = (
            f"{reverse('cadastros:viajante-cadastrar')}?next="
            f"{quote(reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk}))}"
        )
        self.assertContains(response, expected)

    def test_15_reabrir_step1_mostra_dados_salvos(self):
        oficio = self._criar_oficio(ano=2026)
        url = reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk})
        post = self.client.post(
            url,
            data=self._payload_step1(
                oficio,
                protocolo='98.765.432-1',
                motivo='Motivo salvo',
                custeio_tipo=Oficio.CUSTEIO_OUTRA_INSTITUICAO,
                nome_instituicao_custeio='InstituiÃ§Ã£o de Teste',
                viajantes=[self.viajante_final.pk],
            ),
        )
        self.assertEqual(post.status_code, 302)
        reaberto = self.client.get(url)
        self.assertEqual(reaberto.status_code, 200)
        self.assertContains(reaberto, '98.765.432-1')
        self.assertContains(reaberto, 'Motivo salvo')
        self.assertContains(reaberto, 'InstituiÃ§Ã£o de Teste')
        self.assertContains(reaberto, f'data-viajante-chip="{self.viajante_final.pk}"')
        self.assertContains(reaberto, f'data-viajante-hidden="{self.viajante_final.pk}"')
        self.assertContains(reaberto, self.viajante_final.nome)

    def test_step1_cabecalho_novo_remove_wizard_visual_e_ativa_autosave(self):
        oficio = self._criar_oficio(ano=2026)
        response = self.client.get(reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Voltar para lista')
        self.assertContains(response, '>Excluir<', html=False)
        self.assertContains(response, 'oficio-sticky-header')
        self.assertContains(response, 'js/oficio_wizard.js')
        self.assertContains(response, 'data-autosave-link="1"')
        self.assertContains(response, 'data-oficio-glance-toggle')
        self.assertContains(response, 'oficio-wizard-shell')
        self.assertContains(response, 'data-oficio-glance-rail="1"')
        self.assertContains(response, 'data-oficio-footer-actions="1"')
        self.assertContains(response, 'Resumo essencial')
        self.assertContains(response, 'Equipe vinculada a este ofÃ­cio')
        self.assertNotContains(response, 'oficio-wizard-submenu-actions')
        self.assertNotContains(response, 'oficio-autosave-status')
        self.assertNotContains(response, 'Wizard do OfÃ­cio')
        self.assertNotContains(response, 'Lista de eventos')
        self.assertNotContains(response, 'form-actions')

    def test_step1_post_duplicado_nao_duplica_viajante_salvo(self):
        oficio = self._criar_oficio(ano=2026)
        response = self.client.post(
            reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk}),
            data=self._payload_step1(oficio, viajantes=[self.viajante_final.pk, self.viajante_final.pk]),
        )
        self.assertEqual(response.status_code, 302)
        oficio.refresh_from_db()
        self.assertEqual(list(oficio.viajantes.values_list('pk', flat=True)), [self.viajante_final.pk])

    def test_step1_regravar_remove_viajante_que_saiu_dos_chips(self):
        oficio = self._criar_oficio(ano=2026)
        oficio.viajantes.set([self.viajante_final, self.viajante_final_2])
        response = self.client.post(
            reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk}),
            data=self._payload_step1(oficio, viajantes=[self.viajante_final_2.pk]),
        )
        self.assertEqual(response.status_code, 302)
        oficio.refresh_from_db()
        self.assertEqual(list(oficio.viajantes.values_list('pk', flat=True)), [self.viajante_final_2.pk])

    def test_step2_reabre_motorista_protocolo_mascarado(self):
        oficio = self._criar_oficio(ano=2026)
        url = reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk})
        response = self.client.post(url, data=self._payload_step2())
        self.assertEqual(response.status_code, 302)
        oficio.refresh_from_db()
        self.assertEqual(oficio.motorista_protocolo, '123456789')
        reaberto = self.client.get(url)
        self.assertEqual(reaberto.status_code, 200)
        self.assertContains(reaberto, '12.345.678-9')

    def test_step2_preenche_carona_oficio_referencia_quando_numero_ano_existem(self):
        """Step 2 deve preencher carona_oficio_referencia quando motorista carona e nÃºmero/ano batem com outro ofÃ­cio."""
        oficio_atual = self._criar_oficio()
        oficio_ref = self._criar_oficio()
        self.client.post(
            reverse('eventos:oficio-step1', kwargs={'pk': oficio_atual.pk}),
            data=self._payload_step1(oficio_atual, viajantes=[self.viajante_final.pk]),
        )
        self.assertEqual(self.client.post(
            reverse('eventos:oficio-step2', kwargs={'pk': oficio_atual.pk}),
            data=self._payload_step2(
                motorista_oficio_numero=str(oficio_ref.numero),
                motorista_oficio_ano=str(oficio_ref.ano),
            ),
        ).status_code, 302)
        oficio_atual.refresh_from_db()
        self.assertEqual(oficio_atual.carona_oficio_referencia_id, oficio_ref.pk)

    def test_step2_busca_viatura_por_placa_retorna_modelo_combustivel_e_tipo(self):
        veiculo = Veiculo.objects.create(
            placa='ABC1234',
            modelo='SPIN',
            combustivel=self.combustivel,
            tipo=Veiculo.TIPO_CARACTERIZADO,
            status=Veiculo.STATUS_FINALIZADO,
        )
        response = self.client.get(
            reverse('eventos:oficio-step2-veiculo-api'),
            {'placa': 'ABC-1234'},
        )
        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertTrue(payload['ok'])
        self.assertTrue(payload['found'])
        self.assertEqual(payload['modelo'], veiculo.modelo)
        self.assertEqual(payload['combustivel'], 'GASOLINA')
        self.assertEqual(payload['tipo_viatura'], Oficio.TIPO_VIATURA_CARACTERIZADA)

    def test_step2_autocomplete_de_viatura_busca_por_modelo_e_filtra_rascunho(self):
        finalizado = Veiculo.objects.create(
            placa='ABC1234',
            modelo='SPIN LT',
            combustivel=self.combustivel,
            tipo=Veiculo.TIPO_CARACTERIZADO,
            status=Veiculo.STATUS_FINALIZADO,
        )
        Veiculo.objects.create(
            placa='ABC9999',
            modelo='SPIN LTZ',
            combustivel=self.combustivel,
            tipo=Veiculo.TIPO_CARACTERIZADO,
            status=Veiculo.STATUS_RASCUNHO,
        )

        response = self.client.get(
            reverse('eventos:oficio-step2-veiculos-busca-api'),
            {'q': 'spin'},
        )
        self.assertEqual(response.status_code, 200)
        payload = response.json()
        returned_ids = [item['id'] for item in payload['results']]
        self.assertIn(finalizado.pk, returned_ids)
        self.assertContains(response, 'ABC-1234')
        self.assertNotContains(response, 'ABC-9999')

    def test_step2_placa_inexistente_retorna_found_false(self):
        response = self.client.get(
            reverse('eventos:oficio-step2-veiculo-api'),
            {'placa': 'ZZZ-9999'},
        )
        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertTrue(payload['ok'])
        self.assertFalse(payload['found'])

    def test_step2_exibe_porte_transporte_armas_com_sim_por_padrao(self):
        oficio = self._criar_oficio(ano=2026)
        response = self.client.get(reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.context['form'].initial['porte_transporte_armas'], '1')
        self.assertEqual(response.context['step2_preview']['porte_transporte_armas_label'], 'Sim')

    def test_step2_salva_e_reabre_porte_transporte_armas_nao(self):
        oficio = self._criar_oficio(ano=2026)
        response = self.client.post(
            reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}),
            data=self._payload_step2(porte_transporte_armas='0'),
        )
        self.assertEqual(response.status_code, 302)
        oficio.refresh_from_db()
        self.assertFalse(oficio.porte_transporte_armas)
        reaberto = self.client.get(reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}))
        self.assertEqual(reaberto.status_code, 200)
        self.assertEqual(reaberto.context['form'].initial['porte_transporte_armas'], '0')
        self.assertEqual(reaberto.context['step2_preview']['porte_transporte_armas_label'], 'NÃ£o')

    def test_step2_template_da_placa_tem_protecao_contra_loop(self):
        oficio = self._criar_oficio(ano=2026)
        response = self.client.get(reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}))
        self.assertEqual(response.status_code, 200)
        _, script = self._split_step2_response(response)
        self.assertIn("form.dataset.step2JsInitialized === 'true'", script)
        self.assertIn("form.dataset.step2JsInitialized = 'true'", script)
        self.assertIn('var exactLookupDebounceMs = 380;', script)
        self.assertIn('var exactLookupController = null;', script)
        self.assertIn('AbortController', script)
        self.assertIn('normalized.length !== 7', script)
        self.assertIn('normalized === lastExactLookup || normalized === exactLookupInFlightPlaca', script)
        self.assertIn('exactLookupController.abort()', script)
        self.assertNotIn("placaInput.dispatchEvent(new Event('change'", script)

    def test_step2_template_da_placa_reabre_com_bootstrap_controlado(self):
        oficio = self._criar_oficio(ano=2026)
        response = self.client.get(reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}))
        self.assertEqual(response.status_code, 200)
        _, script = self._split_step2_response(response)
        self.assertIn('function bootstrapVehicleLookupIfNeeded()', script)
        self.assertIn('modeloPreenchido &&', script)
        self.assertIn('combustivelPreenchido &&', script)
        self.assertIn('tipoViaturaPreenchido', script)
        self.assertIn('bootstrapVehicleLookupIfNeeded();', script)

    def test_step2_exibe_botao_cadastrar_nova_viatura_com_retorno(self):
        oficio = self._criar_oficio(ano=2026)
        response = self.client.get(reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}))
        self.assertEqual(response.status_code, 200)
        expected = (
            f"{reverse('cadastros:veiculo-cadastrar')}?next="
            f"{quote(reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}))}"
        )
        self.assertContains(response, expected)

    def test_step2_exibe_botao_adicionar_servidor_com_retorno(self):
        oficio = self._criar_oficio(ano=2026)
        response = self.client.get(reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}))
        self.assertEqual(response.status_code, 200)
        expected = (
            f"{reverse('cadastros:viajante-cadastrar')}?next="
            f"{quote(reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}))}"
        )
        self.assertContains(response, expected)

    def test_step2_oculta_motorista_manual_por_padrao(self):
        oficio = self._criar_oficio(ano=2026)
        response = self.client.get(reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'id="motorista-manual-wrapper"')
        self.assertContains(response, 'class="row g-3 mt-1 d-none" id="motorista-manual-wrapper"')
        self.assertContains(response, 'id="btn-toggle-motorista-mode"')
        self.assertContains(response, 'id="motorista-autocomplete-input"')

    def test_step2_cabecalho_fixo_e_toggle_unico_de_motorista(self):
        oficio = self._criar_oficio(ano=2026)
        response = self.client.get(reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Voltar para lista')
        self.assertContains(response, '>Excluir<', html=False)
        self.assertContains(response, 'oficio-sticky-header')
        self.assertContains(response, 'data-oficio-glance-toggle')
        self.assertContains(response, 'oficio-wizard-shell')
        self.assertContains(response, 'data-oficio-glance-rail="1"')
        self.assertContains(response, 'data-oficio-footer-actions="1"')
        self.assertNotContains(response, 'oficio-wizard-submenu-actions')
        self.assertContains(response, 'id="btn-toggle-motorista-mode"')
        self.assertNotContains(response, 'id="btn-motorista-manual"')
        self.assertNotContains(response, 'id="btn-motorista-servidor"')
        self.assertNotContains(response, 'Wizard do OfÃ­cio')
        self.assertNotContains(response, 'form-actions')

    def test_step2_reexibe_campo_manual_quando_seleciona_motorista_sem_cadastro(self):
        oficio = self._criar_oficio(ano=2026)
        response = self.client.post(
            reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}),
            data=self._payload_step2(motorista_nome=''),
        )
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'class="row g-3 mt-1" id="motorista-manual-wrapper"')
        self.assertContains(response, 'Informe o nome do motorista sem cadastro.')

    def test_step2_motoristas_api_busca_por_nome(self):
        response = self.client.get(
            reverse('eventos:oficio-step2-motoristas-api'),
            {'q': 'Motorista Externo'},
        )
        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload['results'][0]['id'], self.motorista_externo.pk)

    def test_step2_motoristas_api_busca_por_rg(self):
        response = self.client.get(
            reverse('eventos:oficio-step2-motoristas-api'),
            {'q': self.motorista_externo.rg},
        )
        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload['results'][0]['id'], self.motorista_externo.pk)

    def test_step2_motoristas_api_busca_por_cpf_mascarado(self):
        response = self.client.get(
            reverse('eventos:oficio-step2-motoristas-api'),
            {'q': '999.888.777-66'},
        )
        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload['results'][0]['id'], self.motorista_externo.pk)
        self.assertEqual(payload['results'][0]['cpf'], '999.888.777-66')

    def test_step2_motorista_de_fora_do_oficio_vira_carona(self):
        oficio = self._criar_oficio(ano=2026)
        oficio.viajantes.add(self.viajante_final)
        response = self.client.post(
            reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}),
            data=self._payload_step2(
                motorista_viajante=str(self.motorista_externo.pk),
                motorista_nome='',
                motorista_oficio_numero='',
                motorista_protocolo='',
            ),
        )
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'class="row g-3 mt-1" id="motorista-carona-wrapper"')
        self.assertFormError(
            response.context['form'],
            'motorista_oficio_numero',
            'Informe o nÃºmero do ofÃ­cio do motorista.',
        )
        self.assertFormError(
            response.context['form'],
            'motorista_protocolo',
            'Informe o protocolo do motorista.',
        )

    def test_step2_motorista_do_oficio_nao_exige_carona(self):
        oficio = self._criar_oficio(ano=2026)
        oficio.viajantes.add(self.viajante_final)
        response = self.client.post(
            reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}),
            data=self._payload_step2(
                motorista_viajante=str(self.viajante_final.pk),
                motorista_nome='',
                motorista_oficio_numero='',
                motorista_protocolo='',
            ),
        )
        self.assertEqual(response.status_code, 302)
        oficio.refresh_from_db()
        self.assertFalse(oficio.motorista_carona)
        self.assertEqual(oficio.motorista_viajante_id, self.viajante_final.pk)
        self.assertEqual(oficio.motorista_oficio, '')
        self.assertEqual(oficio.motorista_protocolo, '')

    def test_step2_reabre_motorista_cadastrado_com_chip_e_hidden_corretos(self):
        oficio = self._criar_oficio(ano=2026)
        oficio.viajantes.add(self.viajante_final)
        self.client.post(
            reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}),
            data=self._payload_step2(
                motorista_viajante=str(self.viajante_final.pk),
                motorista_nome='',
                motorista_oficio_numero='',
                motorista_protocolo='',
            ),
        )
        response = self.client.get(reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, f'data-motorista-chip="{self.viajante_final.pk}"')
        self.assertContains(response, f'name="motorista_viajante" value="{self.viajante_final.pk}"')

    def test_step2_persistencia_numero_oficio_motorista_com_ano_atual(self):
        oficio = self._criar_oficio(ano=2026)
        response = self.client.post(
            reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}),
            data=self._payload_step2(motorista_oficio_numero='12'),
        )
        self.assertEqual(response.status_code, 302)
        oficio.refresh_from_db()
        self.assertEqual(oficio.motorista_oficio_numero, 12)
        self.assertEqual(oficio.motorista_oficio_ano, tz.localdate().year)
        self.assertEqual(oficio.motorista_oficio, f'12/{tz.localdate().year}')

    def test_step2_reabre_dados_salvos_de_carona(self):
        oficio = self._criar_oficio(ano=2026)
        self.client.post(
            reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}),
            data=self._payload_step2(motorista_oficio_numero='12'),
        )
        response = self.client.get(reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'ABC-1234')
        self.assertContains(response, 'VIATURA TESTE')
        self.assertContains(response, 'GASOLINA')
        self.assertContains(response, f'/{tz.localdate().year}')
        self.assertContains(response, '12.345.678-9')

    def test_step2_reabre_resumo_essencial_com_step1_e_transporte(self):
        oficio = self._criar_oficio(ano=2026)
        self.client.post(
            reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk}),
            data=self._payload_step1(
                oficio,
                protocolo='98.765.432-1',
                motivo='Motivo acumulado',
                custeio_tipo=Oficio.CUSTEIO_OUTRA_INSTITUICAO,
                nome_instituicao_custeio='InstituiÃ§Ã£o Acumulada',
                viajantes=[self.viajante_final.pk],
            ),
        )
        self.client.post(
            reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}),
            data=self._payload_step2(
                motorista_viajante=str(self.motorista_externo.pk),
                motorista_nome='',
                motorista_oficio_numero='12',
            ),
        )
        response = self.client.get(reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}))
        self.assertEqual(response.status_code, 200)
        self.motorista_externo.refresh_from_db()
        self.assertContains(response, 'id="summary-oficio"')
        self.assertContains(response, oficio.numero_formatado)
        self.assertContains(response, 'id="summary-protocolo"')
        self.assertContains(response, '98.765.432-1')
        self.assertContains(response, 'id="summary-viajantes-meta"')
        self.assertContains(response, 'id="summary-data-evento"')
        self.assertIn('10/01/2026', response.context['wizard_glance']['data_evento'])
        self.assertIn('12/01/2026', response.context['wizard_glance']['data_evento'])
        self.assertContains(response, f'data-motorista-chip="{self.motorista_externo.pk}"')
        self.assertContains(response, self.motorista_externo.nome)
        self.assertContains(response, 'class="oficio-carona-pill">Carona</span>', html=False)
        self.assertContains(response, 'id="motorista-carona-wrapper"')
        self.assertContains(response, 'value="12"')
        self.assertContains(response, 'value="12.345.678-9"')
        self.assertNotContains(response, 'RelatÃ³rio rÃ¡pido')
        self.assertNotContains(response, 'Leitura rÃ¡pida do motorista')

    def test_step2_resumo_manual_carona_mostra_documentos_sem_dados_cadastrais(self):
        oficio = self._criar_oficio(ano=2026)
        self.client.post(
            reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}),
            data=self._payload_step2(motorista_nome='Motorista Manual', motorista_oficio_numero='12'),
        )
        response = self.client.get(reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}))
        self.assertEqual(response.status_code, 200)
        preview = response.context['step2_preview']['motorista']
        self.assertTrue(preview['is_manual'])
        self.assertFalse(preview['show_cargo'])
        self.assertFalse(preview['show_rg'])
        self.assertFalse(preview['show_cpf'])
        self.assertTrue(preview['show_oficio'])
        self.assertTrue(preview['show_protocolo'])
        server_html, _ = self._split_step2_response(response)
        self.assertIn('name="motorista_nome"', server_html)
        self.assertIn('value="MOTORISTA MANUAL"', server_html)
        self.assertIn(f'/{tz.localdate().year}</span>', server_html)
        self.assertIn('value="12.345.678-9"', server_html)
        self.assertIn('id="motorista-carona-wrapper"', server_html)
        self.assertNotIn('preview-motorista-nome"', server_html)
        self.assertNotIn('Leitura rÃ¡pida do motorista', server_html)

    def test_step2_motorista_do_oficio_nao_exibe_linhas_vazias_de_carona_no_resumo(self):
        oficio = self._criar_oficio(ano=2026)
        oficio.viajantes.add(self.viajante_final)
        self.client.post(
            reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}),
            data=self._payload_step2(
                motorista_viajante=str(self.viajante_final.pk),
                motorista_nome='',
                motorista_oficio_numero='',
                motorista_protocolo='',
            ),
        )
        response = self.client.get(reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}))
        self.assertEqual(response.status_code, 200)
        preview = response.context['step2_preview']['motorista']
        self.assertFalse(preview['is_carona'])
        self.assertFalse(preview['show_oficio'])
        self.assertFalse(preview['show_protocolo'])
        server_html, _ = self._split_step2_response(response)
        self.assertNotIn('id="preview-motorista-oficio-row"', server_html)
        self.assertNotIn('id="preview-motorista-protocolo-row"', server_html)

    def test_step2_reabre_modo_manual_quando_motorista_manual_foi_salvo(self):
        oficio = self._criar_oficio(ano=2026)
        self.client.post(
            reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}),
            data=self._payload_step2(motorista_nome='Motorista Manual', motorista_oficio_numero='12'),
        )
        response = self.client.get(reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'name="motorista_viajante" value="__manual__"')
        self.assertContains(response, 'id="motorista-manual-wrapper"')
        self.assertContains(response, 'MOTORISTA MANUAL')

    def test_step2_salva_manual_quando_placa_existe_e_campos_vao_vazios(self):
        Veiculo.objects.create(
            placa='ABC1234',
            modelo='TRAILBLAZER',
            combustivel=self.combustivel,
            tipo=Veiculo.TIPO_DESCARACTERIZADO,
            status=Veiculo.STATUS_FINALIZADO,
        )
        oficio = self._criar_oficio(ano=2026)
        response = self.client.post(
            reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}),
            data=self._payload_step2(modelo='', combustivel='', tipo_viatura=''),
        )
        self.assertEqual(response.status_code, 302)
        oficio.refresh_from_db()
        self.assertEqual(oficio.modelo, 'TRAILBLAZER')
        self.assertEqual(oficio.combustivel, 'GASOLINA')
        self.assertEqual(oficio.tipo_viatura, Oficio.TIPO_VIATURA_DESCARACTERIZADA)

    def test_step4_exibe_protocolos_mascarados(self):
        oficio = self._criar_oficio(ano=2026)
        self.client.post(reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk}), data=self._payload_step1(oficio))
        self.client.post(reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}), data=self._payload_step2())
        response = self.client.get(reverse('eventos:oficio-step4', kwargs={'pk': oficio.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, '12.345.678-9')
        self.assertContains(response, 'ABC-1234')

    def test_step1_exibe_botao_gerenciar_modelos_de_motivo(self):
        oficio = self._criar_oficio(ano=2026)
        url = reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk})
        response = self.client.get(url)
        self.assertEqual(response.status_code, 200)
        modelos_url = reverse('eventos:modelos-motivo-lista')
        self.assertContains(response, f'{modelos_url}?volta_step1={oficio.pk}')

    def test_modelos_motivo_via_wizard_remove_header_duplicado_e_usa_botao_volta_ao_oficio(self):
        oficio = self._criar_oficio(ano=2026)
        response = self.client.get(
            f"{reverse('eventos:modelos-motivo-lista')}?volta_step1={oficio.pk}"
        )
        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, '<header class="main-header">', html=True)
        self.assertContains(response, 'Voltar ao OfÃ­cio')
        self.assertNotContains(response, 'Voltar para Step 1')

    def test_modelo_motivo_texto_api_retorna_texto(self):
        modelo = ModeloMotivoViagem.objects.create(
            codigo='api_modelo',
            nome='Modelo API',
            texto='Texto API',
            ordem=1,
            ativo=True,
        )
        url = reverse('eventos:modelos-motivo-texto-api', kwargs={'pk': modelo.pk})
        response = self.client.get(url)
        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertTrue(payload.get('ok'))
        self.assertEqual(payload.get('texto'), 'Texto API')

    def test_excluir_oficio_reutiliza_lacuna_no_mesmo_ano(self):
        oficios = [self._criar_oficio(ano=2026) for _ in range(5)]
        oficios[4].delete()
        proximo = self._criar_oficio(ano=2026)
        self.assertEqual(proximo.numero_formatado, '05/2026')

    def test_oficio_pode_ser_excluido_com_redirecionamento_coerente(self):
        oficio = self._criar_oficio(ano=2026)
        url = reverse('eventos:oficio-excluir', kwargs={'pk': oficio.pk})
        response = self.client.post(url)
        self.assertRedirects(response, reverse('eventos:oficios-global'))
        self.assertFalse(Oficio.objects.filter(pk=oficio.pk).exists())

    def test_step3_get_seed_do_roteiro_do_evento(self):
        oficio = self._criar_oficio(ano=2026)
        cidade_destino = Cidade.objects.create(nome='Londrina Step3', estado=self.estado, codigo_ibge='4113701')
        self._criar_roteiro_evento_base([cidade_destino])

        response = self.client.get(reverse('eventos:oficio-step3', kwargs={'pk': oficio.pk}))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'PrÃ©-preenchido com o roteiro salvo.')
        self.assertContains(response, cidade_destino.nome)
        self.assertEqual(len(response.context['trechos_state']), 1)
        self.assertEqual(response.context['trechos_state'][0]['saida_data'], '2026-02-10')
        self.assertEqual(response.context['retorno_state']['saida_data'], '2026-02-11')
        self.assertEqual(response.context['retorno_state']['saida_hora'], '09:00')

    def test_step3_get_seed_do_roteiro_do_evento_carrega_tempo_e_fonte_de_ida_e_retorno(self):
        oficio = self._criar_oficio(ano=2026)
        cidade_destino = Cidade.objects.create(nome='Apucarana Seed', estado=self.estado, codigo_ibge='4101408')
        roteiro = self._criar_roteiro_evento_base([cidade_destino])
        trecho_ida = roteiro.trechos.get(ordem=0)
        trecho_retorno = roteiro.trechos.get(ordem=1)
        trecho_ida.tempo_cru_estimado_min = 180
        trecho_ida.tempo_adicional_min = 20
        trecho_ida.duracao_estimada_min = 200
        trecho_ida.rota_fonte = 'OSRM'
        trecho_ida.save(update_fields=['tempo_cru_estimado_min', 'tempo_adicional_min', 'duracao_estimada_min', 'rota_fonte'])
        trecho_retorno.tempo_cru_estimado_min = 210
        trecho_retorno.tempo_adicional_min = 15
        trecho_retorno.duracao_estimada_min = 225
        trecho_retorno.rota_fonte = 'ESTIMATIVA_LOCAL'
        trecho_retorno.save(update_fields=['tempo_cru_estimado_min', 'tempo_adicional_min', 'duracao_estimada_min', 'rota_fonte'])

        response = self.client.get(reverse('eventos:oficio-step3', kwargs={'pk': oficio.pk}))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.context['trechos_state'][0]['tempo_cru_estimado_min'], 180)
        self.assertEqual(response.context['trechos_state'][0]['rota_fonte'], 'OSRM')
        self.assertEqual(response.context['retorno_state']['tempo_cru_estimado_min'], 210)
        self.assertEqual(response.context['retorno_state']['rota_fonte'], 'ESTIMATIVA_LOCAL')

    def test_step3_permite_escolher_qualquer_roteiro_salvo_do_banco(self):
        oficio = self._criar_oficio(ano=2026)
        cidade_a = Cidade.objects.create(nome='Destino A Step3', estado=self.estado, codigo_ibge='4113703')
        cidade_b = Cidade.objects.create(nome='Destino B Step3', estado=self.estado, codigo_ibge='4113704')
        roteiro_a = self._criar_roteiro_evento_base([cidade_a], data_base=datetime(2026, 2, 10, 8, 0))
        data_base = tz.make_aware(datetime(2026, 2, 12, 7, 30))
        roteiro_b = RoteiroEvento.objects.create(
            evento=None,
            origem_estado=self.estado,
            origem_cidade=self.cidade,
            saida_dt=data_base,
            retorno_saida_dt=data_base + timedelta(days=1, hours=1),
            retorno_chegada_dt=data_base + timedelta(days=1, hours=4),
            status=RoteiroEvento.STATUS_FINALIZADO,
            tipo=RoteiroEvento.TIPO_AVULSO,
        )
        RoteiroEventoDestino.objects.create(
            roteiro=roteiro_b,
            estado=self.estado,
            cidade=cidade_b,
            ordem=0,
        )
        RoteiroEventoTrecho.objects.create(
            roteiro=roteiro_b,
            ordem=0,
            tipo=RoteiroEventoTrecho.TIPO_IDA,
            origem_estado_id=self.estado.pk,
            origem_cidade_id=self.cidade.pk,
            destino_estado_id=self.estado.pk,
            destino_cidade_id=cidade_b.pk,
            saida_dt=data_base,
            chegada_dt=data_base + timedelta(hours=2),
        )
        RoteiroEventoTrecho.objects.create(
            roteiro=roteiro_b,
            ordem=1,
            tipo=RoteiroEventoTrecho.TIPO_RETORNO,
            origem_estado_id=self.estado.pk,
            origem_cidade_id=cidade_b.pk,
            destino_estado_id=self.estado.pk,
            destino_cidade_id=self.cidade.pk,
            saida_dt=data_base + timedelta(days=1, hours=1),
            chegada_dt=data_base + timedelta(days=1, hours=4),
        )
        roteiro_b.chegada_dt = data_base + timedelta(hours=2)
        roteiro_b.save()

        response = self.client.get(reverse('eventos:oficio-step3', kwargs={'pk': oficio.pk}))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(len(response.context['roteiros_evento']), 2)

        post = self.client.post(
            reverse('eventos:oficio-step3', kwargs={'pk': oficio.pk}),
            data={
                'roteiro_modo': Oficio.ROTEIRO_MODO_EVENTO,
                'roteiro_evento_id': str(roteiro_b.pk),
            },
        )

        self.assertEqual(post.status_code, 302)
        oficio.refresh_from_db()
        self.assertEqual(oficio.roteiro_modo, Oficio.ROTEIRO_MODO_EVENTO)
        self.assertEqual(oficio.roteiro_evento_id, roteiro_b.pk)
        self.assertEqual(OficioTrecho.objects.filter(oficio=oficio).count(), 1)
        self.assertEqual(OficioTrecho.objects.get(oficio=oficio).destino_cidade_id, cidade_b.pk)
        self.assertEqual(RoteiroEventoTrecho.objects.filter(roteiro=roteiro_a).count(), 2)
        self.assertEqual(RoteiroEventoTrecho.objects.filter(roteiro=roteiro_b).count(), 2)

        reopened = self.client.get(reverse('eventos:oficio-step3', kwargs={'pk': oficio.pk}))

        self.assertEqual(reopened.status_code, 200)
        self.assertEqual(reopened.context['roteiro_evento_id'], roteiro_b.pk)
        self.assertEqual(reopened.context['roteiro_modo'], Oficio.ROTEIRO_MODO_EVENTO)
        self.assertEqual(reopened.context['trechos_state'][0]['destino_cidade_id'], cidade_b.pk)
        self.assertContains(reopened, cidade_b.nome)

    def test_step3_salva_trechos_no_proprio_oficio_sem_corromper_roteiro_evento(self):
        oficio = self._criar_oficio(ano=2026)
        cidade_destino = Cidade.objects.create(nome='MaringÃ¡ Step3', estado=self.estado, codigo_ibge='4115201')
        roteiro = self._criar_roteiro_evento_base([cidade_destino])

        response = self.client.post(
            reverse('eventos:oficio-step3', kwargs={'pk': oficio.pk}),
            data=self._payload_step3(
                [cidade_destino],
                trechos=[
                    {
                        'saida_data': '2026-03-10',
                        'saida_hora': '08:00',
                        'chegada_data': '2026-03-10',
                        'chegada_hora': '12:00',
                    }
                ],
                retorno={
                    'saida_data': '2026-03-11',
                    'saida_hora': '09:00',
                    'chegada_data': '2026-03-11',
                    'chegada_hora': '18:00',
                },
            ),
        )

        self.assertEqual(response.status_code, 302)
        oficio.refresh_from_db()
        trechos = list(OficioTrecho.objects.filter(oficio=oficio).order_by('ordem'))
        self.assertEqual(len(trechos), 1)
        self.assertEqual(trechos[0].destino_cidade, cidade_destino)
        self.assertEqual(trechos[0].saida_hora.strftime('%H:%M'), '08:00')
        self.assertEqual(oficio.retorno_saida_cidade, f'{cidade_destino.nome}/{self.estado.sigla}')
        self.assertEqual(oficio.retorno_chegada_cidade, f'{self.cidade.nome}/{self.estado.sigla}')
        trecho_roteiro = RoteiroEventoTrecho.objects.get(roteiro=roteiro, ordem=0)
        self.assertEqual(tz.localtime(trecho_roteiro.saida_dt).replace(tzinfo=None), datetime(2026, 2, 10, 8, 0))
        self.assertEqual(RoteiroEventoTrecho.objects.filter(roteiro=roteiro).count(), 2)

    def test_step3_persiste_estimativa_local_no_roteiro_proprio(self):
        oficio = self._criar_oficio(ano=2026)
        cidade_destino = Cidade.objects.create(nome='Apucarana Step3', estado=self.estado, codigo_ibge='4101409')

        response = self.client.post(
            reverse('eventos:oficio-step3', kwargs={'pk': oficio.pk}),
            data=self._payload_step3(
                [cidade_destino],
                trechos=[
                    {
                        'saida_data': '2026-03-10',
                        'saida_hora': '08:00',
                        'chegada_data': '2026-03-10',
                        'chegada_hora': '10:45',
                    }
                ],
                retorno={
                    'saida_data': '2026-03-11',
                    'saida_hora': '09:00',
                    'chegada_data': '2026-03-11',
                    'chegada_hora': '12:00',
                },
                trecho_0_distancia_km='123.45',
                trecho_0_tempo_cru_estimado_min='150',
                trecho_0_tempo_adicional_min='15',
                trecho_0_duracao_estimada_min='165',
                trecho_0_rota_fonte='ESTIMATIVA_LOCAL',
            ),
        )

        self.assertEqual(response.status_code, 302)
        trecho = OficioTrecho.objects.get(oficio=oficio)
        self.assertEqual(str(trecho.distancia_km), '123.45')
        self.assertEqual(trecho.tempo_cru_estimado_min, 150)
        self.assertEqual(trecho.tempo_adicional_min, 15)
        self.assertEqual(trecho.duracao_estimada_min, 165)
        self.assertEqual(trecho.rota_fonte, 'ESTIMATIVA_LOCAL')

        reopened = self.client.get(reverse('eventos:oficio-step3', kwargs={'pk': oficio.pk}))

        self.assertEqual(reopened.context['trechos_state'][0]['distancia_km'], '123.45')
        self.assertEqual(reopened.context['trechos_state'][0]['tempo_cru_estimado_min'], 150)
        self.assertEqual(reopened.context['trechos_state'][0]['tempo_adicional_min'], 15)
        self.assertEqual(reopened.context['trechos_state'][0]['duracao_estimada_min'], 165)
        self.assertEqual(reopened.context['trechos_state'][0]['rota_fonte'], 'ESTIMATIVA_LOCAL')

    def test_step3_reabre_com_mesma_ordem_cidades_e_horarios(self):
        oficio = self._criar_oficio(ano=2026)
        cidade_b = Cidade.objects.create(nome='Londrina Reabertura', estado=self.estado, codigo_ibge='4113702')
        cidade_c = Cidade.objects.create(nome='Foz Reabertura', estado=self.estado, codigo_ibge='4108301')
        self.client.post(
            reverse('eventos:oficio-step3', kwargs={'pk': oficio.pk}),
            data=self._payload_step3(
                [cidade_b, cidade_c],
                trechos=[
                    {
                        'saida_data': '2026-04-01',
                        'saida_hora': '08:00',
                        'chegada_data': '2026-04-01',
                        'chegada_hora': '10:00',
                    },
                    {
                        'saida_data': '2026-04-01',
                        'saida_hora': '11:00',
                        'chegada_data': '2026-04-01',
                        'chegada_hora': '15:00',
                    },
                ],
                retorno={
                    'saida_data': '2026-04-02',
                    'saida_hora': '09:00',
                    'chegada_data': '2026-04-02',
                    'chegada_hora': '18:00',
                },
            ),
        )

        response = self.client.get(reverse('eventos:oficio-step3', kwargs={'pk': oficio.pk}))

        self.assertEqual(response.status_code, 200)
        trechos_state = response.context['trechos_state']
        self.assertEqual(len(trechos_state), 2)
        self.assertEqual(trechos_state[0]['destino_nome'], f'{cidade_b.nome}/{self.estado.sigla}')
        self.assertEqual(trechos_state[1]['destino_nome'], f'{cidade_c.nome}/{self.estado.sigla}')
        self.assertEqual(trechos_state[0]['saida_hora'], '08:00')
        self.assertEqual(trechos_state[1]['saida_hora'], '11:00')
        self.assertEqual(response.context['retorno_state']['saida_data'], '2026-04-02')
        self.assertEqual(response.context['retorno_state']['chegada_hora'], '18:00')

    def test_step3_reabre_bate_volta_diario_com_trechos_explicitos_e_destino_unico(self):
        oficio = self._criar_oficio(ano=2026)
        cidade_destino = Cidade.objects.create(nome='Colombo Loop', estado=self.estado, codigo_ibge='4105806')
        self._salvar_steps_1_e_2(oficio)

        response_post = self.client.post(
            reverse('eventos:oficio-step3', kwargs={'pk': oficio.pk}),
            data=self._payload_step3(
                [cidade_destino],
                trechos=[
                    {
                        'origem_nome': f'{self.cidade.nome}/{self.estado.sigla}',
                        'destino_nome': f'{cidade_destino.nome}/{self.estado.sigla}',
                        'origem_estado_id': str(self.estado.pk),
                        'origem_cidade_id': str(self.cidade.pk),
                        'destino_estado_id': str(self.estado.pk),
                        'destino_cidade_id': str(cidade_destino.pk),
                        'saida_data': '2026-04-23',
                        'saida_hora': '08:00',
                        'chegada_data': '2026-04-23',
                        'chegada_hora': '08:40',
                        'tempo_cru_estimado_min': '40',
                        'tempo_adicional_min': '0',
                        'duracao_estimada_min': '40',
                    },
                    {
                        'origem_nome': f'{cidade_destino.nome}/{self.estado.sigla}',
                        'destino_nome': f'{self.cidade.nome}/{self.estado.sigla}',
                        'origem_estado_id': str(self.estado.pk),
                        'origem_cidade_id': str(cidade_destino.pk),
                        'destino_estado_id': str(self.estado.pk),
                        'destino_cidade_id': str(self.cidade.pk),
                        'saida_data': '2026-04-23',
                        'saida_hora': '17:30',
                        'chegada_data': '2026-04-23',
                        'chegada_hora': '18:10',
                        'tempo_cru_estimado_min': '40',
                        'tempo_adicional_min': '0',
                        'duracao_estimada_min': '40',
                    },
                    {
                        'origem_nome': f'{self.cidade.nome}/{self.estado.sigla}',
                        'destino_nome': f'{cidade_destino.nome}/{self.estado.sigla}',
                        'origem_estado_id': str(self.estado.pk),
                        'origem_cidade_id': str(self.cidade.pk),
                        'destino_estado_id': str(self.estado.pk),
                        'destino_cidade_id': str(cidade_destino.pk),
                        'saida_data': '2026-04-24',
                        'saida_hora': '08:00',
                        'chegada_data': '2026-04-24',
                        'chegada_hora': '08:40',
                        'tempo_cru_estimado_min': '40',
                        'tempo_adicional_min': '0',
                        'duracao_estimada_min': '40',
                    },
                    {
                        'origem_nome': f'{cidade_destino.nome}/{self.estado.sigla}',
                        'destino_nome': f'{self.cidade.nome}/{self.estado.sigla}',
                        'origem_estado_id': str(self.estado.pk),
                        'origem_cidade_id': str(cidade_destino.pk),
                        'destino_estado_id': str(self.estado.pk),
                        'destino_cidade_id': str(self.cidade.pk),
                        'saida_data': '2026-04-24',
                        'saida_hora': '17:30',
                        'chegada_data': '2026-04-24',
                        'chegada_hora': '18:10',
                        'tempo_cru_estimado_min': '40',
                        'tempo_adicional_min': '0',
                        'duracao_estimada_min': '40',
                    },
                ],
                retorno={
                    'saida_data': '2026-04-24',
                    'saida_hora': '17:30',
                    'chegada_data': '2026-04-24',
                    'chegada_hora': '18:10',
                    'saida_cidade': f'{cidade_destino.nome}/{self.estado.sigla}',
                    'chegada_cidade': f'{self.cidade.nome}/{self.estado.sigla}',
                },
                bate_volta_diario_ativo='1',
                bate_volta_data_inicio='2026-04-23',
                bate_volta_data_fim='2026-04-24',
                bate_volta_ida_saida_hora='08:00',
                bate_volta_ida_tempo_min='40',
                bate_volta_volta_saida_hora='17:30',
                bate_volta_volta_tempo_min='40',
            ),
        )

        self.assertEqual(response_post.status_code, 302)

        response = self.client.get(reverse('eventos:oficio-step3', kwargs={'pk': oficio.pk}))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(len(response.context['trechos_state']), 4)
        self.assertEqual(len(response.context['destinos_atuais']), 1)
        self.assertEqual(response.context['destinos_atuais'][0]['cidade_id'], cidade_destino.pk)
        self.assertTrue(response.context['step3_state_json']['bate_volta_diario']['ativo'])

    def test_step3_calculadora_diarias_classifica_por_destino_operacional_ignorando_retorno_para_sede(self):
        oficio = self._criar_oficio(ano=2026)
        cidade_destino = Cidade.objects.create(nome='Colombo Classificacao', estado=self.estado, codigo_ibge='4105807')
        oficio.viajantes.add(self.viajante_final)
        self._salvar_steps_1_e_2(oficio)

        response = self.client.post(
            reverse('eventos:oficio-step3-calcular-diarias', kwargs={'pk': oficio.pk}),
            data=self._payload_step3(
                [cidade_destino],
                trechos=[
                    {
                        'origem_nome': f'{self.cidade.nome}/{self.estado.sigla}',
                        'destino_nome': f'{cidade_destino.nome}/{self.estado.sigla}',
                        'origem_estado_id': str(self.estado.pk),
                        'origem_cidade_id': str(self.cidade.pk),
                        'destino_estado_id': str(self.estado.pk),
                        'destino_cidade_id': str(cidade_destino.pk),
                        'saida_data': '2026-04-23',
                        'saida_hora': '08:00',
                        'chegada_data': '2026-04-23',
                        'chegada_hora': '08:40',
                        'tempo_cru_estimado_min': '40',
                        'tempo_adicional_min': '0',
                        'duracao_estimada_min': '40',
                    },
                    {
                        'origem_nome': f'{cidade_destino.nome}/{self.estado.sigla}',
                        'destino_nome': f'{self.cidade.nome}/{self.estado.sigla}',
                        'origem_estado_id': str(self.estado.pk),
                        'origem_cidade_id': str(cidade_destino.pk),
                        'destino_estado_id': str(self.estado.pk),
                        'destino_cidade_id': str(self.cidade.pk),
                        'saida_data': '2026-04-23',
                        'saida_hora': '17:30',
                        'chegada_data': '2026-04-23',
                        'chegada_hora': '18:10',
                        'tempo_cru_estimado_min': '40',
                        'tempo_adicional_min': '0',
                        'duracao_estimada_min': '40',
                    },
                    {
                        'origem_nome': f'{self.cidade.nome}/{self.estado.sigla}',
                        'destino_nome': f'{cidade_destino.nome}/{self.estado.sigla}',
                        'origem_estado_id': str(self.estado.pk),
                        'origem_cidade_id': str(self.cidade.pk),
                        'destino_estado_id': str(self.estado.pk),
                        'destino_cidade_id': str(cidade_destino.pk),
                        'saida_data': '2026-04-24',
                        'saida_hora': '08:00',
                        'chegada_data': '2026-04-24',
                        'chegada_hora': '08:40',
                        'tempo_cru_estimado_min': '40',
                        'tempo_adicional_min': '0',
                        'duracao_estimada_min': '40',
                    },
                    {
                        'origem_nome': f'{cidade_destino.nome}/{self.estado.sigla}',
                        'destino_nome': f'{self.cidade.nome}/{self.estado.sigla}',
                        'origem_estado_id': str(self.estado.pk),
                        'origem_cidade_id': str(cidade_destino.pk),
                        'destino_estado_id': str(self.estado.pk),
                        'destino_cidade_id': str(self.cidade.pk),
                        'saida_data': '2026-04-24',
                        'saida_hora': '17:30',
                        'chegada_data': '2026-04-24',
                        'chegada_hora': '18:10',
                        'tempo_cru_estimado_min': '40',
                        'tempo_adicional_min': '0',
                        'duracao_estimada_min': '40',
                    },
                ],
                retorno={
                    'saida_data': '2026-04-24',
                    'saida_hora': '17:30',
                    'chegada_data': '2026-04-24',
                    'chegada_hora': '18:10',
                    'saida_cidade': f'{cidade_destino.nome}/{self.estado.sigla}',
                    'chegada_cidade': f'{self.cidade.nome}/{self.estado.sigla}',
                },
                bate_volta_diario_ativo='1',
                bate_volta_data_inicio='2026-04-23',
                bate_volta_data_fim='2026-04-24',
                bate_volta_ida_saida_hora='08:00',
                bate_volta_ida_tempo_min='40',
                bate_volta_volta_saida_hora='17:30',
                bate_volta_volta_tempo_min='40',
            ),
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertTrue(payload['ok'])
        self.assertEqual(payload['tipo_destino'], Oficio.TIPO_DESTINO_INTERIOR)
        self.assertEqual(payload['totais']['total_diarias'], '2 x 30%')

    def test_step3_calcular_diarias_salva_roteiro_sem_duplicar(self):
        ConfiguracaoSistema.objects.create(cidade_sede_padrao=self.cidade, prazo_justificativa_dias=10)
        oficio = self._criar_oficio(ano=2026)
        Oficio.objects.filter(pk=oficio.pk).update(data_criacao=date(2026, 9, 20))
        oficio.refresh_from_db()
        cidade_destino = Cidade.objects.create(nome='Destino Salvar Roteiro', estado=self.estado, codigo_ibge='4113707')
        self._salvar_steps_1_e_2(oficio)

        payload = self._payload_step3(
            [cidade_destino],
            trechos=[
                {
                    'saida_data': '2026-10-10',
                    'saida_hora': '08:00',
                    'chegada_data': '2026-10-10',
                    'chegada_hora': '12:00',
                }
            ],
            retorno={
                'saida_data': '2026-10-11',
                'saida_hora': '09:00',
                'chegada_data': '2026-10-11',
                'chegada_hora': '13:30',
            },
            trecho_0_distancia_km='123.45',
            trecho_0_tempo_cru_estimado_min='210',
            trecho_0_tempo_adicional_min='25',
            trecho_0_duracao_estimada_min='235',
            trecho_0_rota_fonte='OSRM',
            retorno_distancia_km='123.45',
            retorno_tempo_cru_estimado_min='220',
            retorno_tempo_adicional_min='25',
            retorno_duracao_estimada_min='245',
            retorno_rota_fonte='ESTIMATIVA_LOCAL',
        )

        response = self.client.post(
            reverse('eventos:oficio-step3-calcular-diarias', kwargs={'pk': oficio.pk}),
            data=payload,
        )

        self.assertEqual(response.status_code, 200)
        payload_response = response.json()
        self.assertTrue(payload_response['ok'])
        self.assertTrue(payload_response['roteiro_salvo_criado'])

        roteiros_salvos = RoteiroEvento.objects.filter(origem_cidade=self.cidade)
        self.assertEqual(roteiros_salvos.count(), 1)
        roteiro = roteiros_salvos.get()
        self.assertEqual(roteiro.tipo, RoteiroEvento.TIPO_EVENTO)
        self.assertEqual(roteiro.destinos.count(), 1)
        self.assertEqual(roteiro.destinos.get().cidade_id, cidade_destino.pk)
        self.assertEqual(roteiro.trechos.count(), 2)
        trecho_ida = roteiro.trechos.get(ordem=0)
        trecho_retorno = roteiro.trechos.get(ordem=1)
        self.assertEqual(str(trecho_ida.distancia_km), '123.45')
        self.assertEqual(trecho_ida.tempo_cru_estimado_min, 210)
        self.assertEqual(trecho_ida.tempo_adicional_min, 25)
        self.assertEqual(trecho_ida.duracao_estimada_min, 235)
        self.assertEqual(trecho_ida.rota_fonte, 'OSRM')
        self.assertEqual(trecho_retorno.tempo_cru_estimado_min, 220)
        self.assertEqual(trecho_retorno.tempo_adicional_min, 25)
        self.assertEqual(trecho_retorno.duracao_estimada_min, 245)
        self.assertEqual(trecho_retorno.rota_fonte, 'ESTIMATIVA_LOCAL')

        response_repeat = self.client.post(
            reverse('eventos:oficio-step3-calcular-diarias', kwargs={'pk': oficio.pk}),
            data=self._payload_step3(
                [cidade_destino],
                trechos=[
                    {
                        'saida_data': '2026-10-10',
                        'saida_hora': '08:00',
                        'chegada_data': '2026-10-10',
                        'chegada_hora': '12:00',
                    }
                ],
                retorno={
                    'saida_data': '2026-10-11',
                    'saida_hora': '09:00',
                    'chegada_data': '2026-10-11',
                    'chegada_hora': '13:30',
                },
                trecho_0_distancia_km='123.45',
                trecho_0_tempo_cru_estimado_min='210',
                trecho_0_tempo_adicional_min='25',
                trecho_0_duracao_estimada_min='235',
                trecho_0_rota_fonte='OSRM',
                retorno_distancia_km='123.45',
                retorno_tempo_cru_estimado_min='220',
                retorno_tempo_adicional_min='25',
                retorno_duracao_estimada_min='245',
                retorno_rota_fonte='ESTIMATIVA_LOCAL',
            ),
        )

        self.assertEqual(response_repeat.status_code, 200)
        self.assertFalse(response_repeat.json()['roteiro_salvo_criado'])
        self.assertEqual(RoteiroEvento.objects.filter(origem_cidade=self.cidade).count(), 1)

    def test_step3_calcular_diarias_ativa_fluxo_de_justificativa_quando_obrigatoria(self):
        ConfiguracaoSistema.objects.create(cidade_sede_padrao=self.cidade, prazo_justificativa_dias=10)
        oficio = self._criar_oficio(ano=2026)
        Oficio.objects.filter(pk=oficio.pk).update(data_criacao=date(2026, 10, 5))
        oficio.refresh_from_db()
        cidade_destino = Cidade.objects.create(nome='Destino Justificativa Roteiro', estado=self.estado, codigo_ibge='4113708')
        self._salvar_steps_1_e_2(oficio)

        response = self.client.post(
            reverse('eventos:oficio-step3-calcular-diarias', kwargs={'pk': oficio.pk}),
            data=self._payload_step3(
                [cidade_destino],
                trechos=[
                    {
                        'saida_data': '2026-10-10',
                        'saida_hora': '08:00',
                        'chegada_data': '2026-10-10',
                        'chegada_hora': '12:00',
                    }
                ],
                retorno={
                    'saida_data': '2026-10-11',
                    'saida_hora': '09:00',
                    'chegada_data': '2026-10-11',
                    'chegada_hora': '13:30',
                },
            ),
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertTrue(payload['ok'])
        self.assertTrue(payload['justificativa_required'])
        self.assertFalse(payload['justificativa_filled'])
        self.assertIn(reverse('eventos:oficio-justificativa', kwargs={'pk': oficio.pk}), payload['justificativa_url'])
        self.assertEqual(RoteiroEvento.objects.filter(evento=self.evento).count(), 1)

    def test_step3_template_nao_forca_redirecionamento_automatico_para_justificativa(self):
        oficio = self._criar_oficio(ano=2026)
        response = self.client.get(reverse('eventos:oficio-step3', kwargs={'pk': oficio.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, 'window.location.href = data.justificativa_url;')

    def test_step3_exibe_etapa_justificativa_no_stepper_apos_calculo_quando_obrigatoria(self):
        ConfiguracaoSistema.objects.create(cidade_sede_padrao=self.cidade, prazo_justificativa_dias=10)
        oficio = self._criar_oficio(ano=2026)
        Oficio.objects.filter(pk=oficio.pk).update(data_criacao=date(2026, 10, 5))
        oficio.refresh_from_db()
        cidade_destino = Cidade.objects.create(nome='Destino Stepper Justificativa', estado=self.estado, codigo_ibge='4113788')
        self._salvar_steps_1_e_2(oficio)
        self.client.post(
            reverse('eventos:oficio-step3-calcular-diarias', kwargs={'pk': oficio.pk}),
            data=self._payload_step3(
                [cidade_destino],
                trechos=[
                    {
                        'saida_data': '2026-10-10',
                        'saida_hora': '08:00',
                        'chegada_data': '2026-10-10',
                        'chegada_hora': '12:00',
                    }
                ],
                retorno={
                    'saida_data': '2026-10-11',
                    'saida_hora': '09:00',
                    'chegada_data': '2026-10-11',
                    'chegada_hora': '13:30',
                },
            ),
        )

        reaberta = self.client.get(reverse('eventos:oficio-step3', kwargs={'pk': oficio.pk}))
        self.assertEqual(reaberta.status_code, 200)
        self.assertContains(reaberta, 'data-step-key="justificativa"')

    def test_step3_valida_retorno_obrigatorio(self):
        oficio = self._criar_oficio(ano=2026)
        cidade_destino = Cidade.objects.create(nome='Ponta Grossa Step3', estado=self.estado, codigo_ibge='4119906')

        response = self.client.post(
            reverse('eventos:oficio-step3', kwargs={'pk': oficio.pk}),
            data=self._payload_step3(
                [cidade_destino],
                trechos=[
                    {
                        'saida_data': '2026-05-10',
                        'saida_hora': '08:00',
                        'chegada_data': '2026-05-10',
                        'chegada_hora': '11:00',
                    }
                ],
                retorno={},
            ),
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Informe a saÃ­da do retorno')
        self.assertContains(response, 'Informe a chegada do retorno')

    def test_step3_resumo_rapido_acumula_steps_1_2_e_3(self):
        oficio = self._criar_oficio(ano=2026)
        cidade_destino = Cidade.objects.create(nome='Guarapuava Step3', estado=self.estado, codigo_ibge='4109400')
        self.client.post(reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk}), data=self._payload_step1(oficio))
        self.client.post(reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}), data=self._payload_step2())
        self.client.post(
            reverse('eventos:oficio-step3', kwargs={'pk': oficio.pk}),
            data=self._payload_step3(
                [cidade_destino],
                trechos=[
                    {
                        'saida_data': '2026-06-01',
                        'saida_hora': '08:30',
                        'chegada_data': '2026-06-01',
                        'chegada_hora': '12:00',
                    }
                ],
                retorno={
                    'saida_data': '2026-06-02',
                    'saida_hora': '07:00',
                    'chegada_data': '2026-06-02',
                    'chegada_hora': '14:00',
                },
            ),
        )

        response = self.client.get(reverse('eventos:oficio-step3', kwargs={'pk': oficio.pk}))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'data-oficio-glance')
        self.assertContains(response, '12.345.678-9')
        self.assertContains(response, 'id="summary-oficio"')
        self.assertContains(response, 'id="summary-destino"')
        self.assertContains(response, 'id="summary-data-evento"')
        self.assertContains(response, '10/01/2026 atÃ© 12/01/2026')
        self.assertContains(response, cidade_destino.nome)
        self.assertNotContains(response, 'RelatÃ³rio rÃ¡pido')

    def test_step3_cabecalho_fixo_remove_texto_legado_e_mostra_nova_ui(self):
        oficio = self._criar_oficio(ano=2026)
        response = self.client.get(reverse('eventos:oficio-step3', kwargs={'pk': oficio.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Voltar para lista')
        self.assertContains(response, '>Excluir<', html=False)
        self.assertContains(response, 'data-oficio-sticky-header')
        self.assertContains(response, 'data-oficio-glance-toggle')
        self.assertContains(response, 'data-oficio-glance-rail="1"')
        self.assertContains(response, 'oficio-wizard-shell')
        self.assertContains(response, 'data-oficio-footer-actions="1"')
        self.assertContains(response, 'Equipe vinculada a este ofÃ­cio')
        self.assertNotContains(response, 'oficio-wizard-submenu-actions')
        self.assertContains(response, 'Dados e viajantes')
        self.assertContains(response, 'Roteiro e diÃ¡rias')
        self.assertNotContains(response, 'Tempo estimado de viagem')
        self.assertContains(response, 'Tempo de viagem')
        self.assertContains(response, 'Tempo adicional')
        self.assertContains(response, 'Tempo total')
        self.assertContains(response, 'ParanÃ¡ (PR)')
        self.assertContains(response, 'Valor total')
        self.assertNotContains(response, 'Valor por extenso')
        self.assertNotContains(response, 'Step 1')
        self.assertNotContains(response, 'Salvar roteiro')
        self.assertNotContains(response, 'CÃ¡lculo periodizado do legado com base nos horÃ¡rios do Step 3.')
        self.assertNotContains(response, 'Calculo atualizado a partir dos horarios do Step 3.')
        self.assertNotContains(response, 'Wizard do OfÃ­cio')
        self.assertNotContains(response, 'form-actions')

    def test_step3_calculadora_diarias_endpoint_usa_dados_dos_trechos(self):
        oficio = self._criar_oficio(ano=2026)
        cidade_destino = Cidade.objects.create(nome='Interior Diarias Step3', estado=self.estado, codigo_ibge='4113705')
        oficio.viajantes.add(self.viajante_final)

        response = self.client.post(
            reverse('eventos:oficio-step3-calcular-diarias', kwargs={'pk': oficio.pk}),
            data=self._payload_step3(
                [cidade_destino],
                trechos=[
                    {
                        'saida_data': '2026-08-10',
                        'saida_hora': '08:00',
                        'chegada_data': '2026-08-10',
                        'chegada_hora': '12:00',
                    }
                ],
                retorno={
                    'saida_data': '2026-08-11',
                    'saida_hora': '08:00',
                    'chegada_data': '2026-08-11',
                    'chegada_hora': '10:00',
                },
            ),
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertTrue(payload['ok'])
        self.assertEqual(payload['tipo_destino'], Oficio.TIPO_DESTINO_INTERIOR)
        self.assertEqual(payload['totais']['total_diarias'], '1 x 100%')
        self.assertEqual(payload['totais']['total_valor'], '290,55')
        self.assertTrue(payload['totais']['valor_extenso'])
        oficio.refresh_from_db()
        self.assertEqual(oficio.valor_diarias, '290,55')
        self.assertTrue(oficio.valor_diarias_extenso)

    def test_step3_calculadora_diarias_ignora_motorista_externo_ao_contar_servidores(self):
        oficio = self._criar_oficio(ano=2026)
        cidade_destino = Cidade.objects.create(nome='Interior Sem Motorista', estado=self.estado, codigo_ibge='4113715')
        self.client.post(
            reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk}),
            data=self._payload_step1(oficio, viajantes=[self.viajante_final.pk]),
        )
        self.client.post(
            reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}),
            data=self._payload_step2(
                motorista_viajante=str(self.motorista_externo.pk),
                motorista_nome='',
                motorista_oficio_numero='12',
                motorista_protocolo='12.345.678-9',
            ),
        )
        response = self.client.post(
            reverse('eventos:oficio-step3-calcular-diarias', kwargs={'pk': oficio.pk}),
            data=self._payload_step3(
                [cidade_destino],
                trechos=[
                    {
                        'saida_data': '2026-08-20',
                        'saida_hora': '08:00',
                        'chegada_data': '2026-08-20',
                        'chegada_hora': '12:00',
                    }
                ],
                retorno={
                    'saida_data': '2026-08-21',
                    'saida_hora': '08:00',
                    'chegada_data': '2026-08-21',
                    'chegada_hora': '10:00',
                },
            ),
        )
        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertTrue(payload['ok'])
        self.assertEqual(payload['totais']['total_diarias'], '1 x 100%')
        self.assertTrue(payload['totais']['valor_extenso'])

    def test_step4_exibe_trechos_salvos_do_step3(self):
        oficio = self._criar_oficio(ano=2026)
        cidade_destino = Cidade.objects.create(nome='Cascavel Step4', estado=self.estado, codigo_ibge='4104809')
        self.client.post(
            reverse('eventos:oficio-step3', kwargs={'pk': oficio.pk}),
            data=self._payload_step3(
                [cidade_destino],
                trechos=[
                    {
                        'saida_data': '2026-07-01',
                        'saida_hora': '08:00',
                        'chegada_data': '2026-07-01',
                        'chegada_hora': '12:30',
                    }
                ],
                retorno={
                    'saida_data': '2026-07-02',
                    'saida_hora': '09:15',
                    'chegada_data': '2026-07-02',
                    'chegada_hora': '18:45',
                },
            ),
        )

        response = self.client.get(reverse('eventos:oficio-step4', kwargs={'pk': oficio.pk}))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, f'{self.cidade.nome}/{self.estado.sigla} â†’ {cidade_destino.nome}/{self.estado.sigla}')
        self.assertContains(response, 'Retorno:')
        self.assertContains(response, '02/07/2026 09:15')

    def test_step4_exibe_modo_fonte_e_diarias_do_step3(self):
        oficio = self._criar_oficio(ano=2026)
        cidade_destino = Cidade.objects.create(nome='Umuarama Step4', estado=self.estado, codigo_ibge='4128105')
        oficio.viajantes.add(self.viajante_final)

        self.client.post(
            reverse('eventos:oficio-step3', kwargs={'pk': oficio.pk}),
            data=self._payload_step3(
                [cidade_destino],
                trechos=[
                    {
                        'saida_data': '2026-09-01',
                        'saida_hora': '08:00',
                        'chegada_data': '2026-09-01',
                        'chegada_hora': '12:00',
                    }
                ],
                retorno={
                    'saida_data': '2026-09-02',
                    'saida_hora': '08:00',
                    'chegada_data': '2026-09-02',
                    'chegada_hora': '10:00',
                },
            ),
        )

        response = self.client.get(reverse('eventos:oficio-step4', kwargs={'pk': oficio.pk}))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.context['step3_preview']['roteiro_modo'], Oficio.ROTEIRO_MODO_PROPRIO)
        self.assertEqual(response.context['step3_preview']['diarias']['quantidade'], '1 x 100%')
        self.assertEqual(response.context['step3_preview']['diarias']['valor_total'], '290,55')

    def test_step4_resumo_remove_relatorio_rapido_e_exibe_opcao_de_termo(self):
        oficio = self._criar_oficio(ano=2026)
        cidade_destino = Cidade.objects.create(nome='Campo MourÃ£o Step4', estado=self.estado, codigo_ibge='4104303')
        oficio.viajantes.add(self.viajante_final)

        self.client.post(
            reverse('eventos:oficio-step3', kwargs={'pk': oficio.pk}),
            data=self._payload_step3(
                [cidade_destino],
                trechos=[
                    {
                        'saida_data': '2026-09-01',
                        'saida_hora': '08:00',
                        'chegada_data': '2026-09-01',
                        'chegada_hora': '12:00',
                    }
                ],
                retorno={
                    'saida_data': '2026-09-02',
                    'saida_hora': '08:00',
                    'chegada_data': '2026-09-02',
                    'chegada_hora': '10:00',
                },
            ),
        )

        response = self.client.get(reverse('eventos:oficio-step4', kwargs={'pk': oficio.pk}))

        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, 'RelatÃ³rio rÃ¡pido')
        self.assertContains(response, 'Gerar termo de autorizaÃ§Ã£o jÃ¡ preenchido?')
        self.assertContains(response, 'Finalizar OfÃ­cio')
        self.assertContains(response, 'Baixar DOCX')
        self.assertContains(response, 'Baixar PDF')
        self.assertContains(
            response,
            reverse(
                'eventos:oficio-documento-download',
                kwargs={'pk': oficio.pk, 'tipo_documento': 'oficio', 'formato': 'docx'},
            ),
        )
        self.assertContains(
            response,
            reverse(
                'eventos:oficio-documento-download',
                kwargs={'pk': oficio.pk, 'tipo_documento': 'oficio', 'formato': 'pdf'},
            ),
        )

    def test_step3_calculo_diarias_respeita_pernoites_tres_noites(self):
        """Regra de pernoites: 3 noites fora da sede â†’ 3 x 100% (nÃ£o 2 x 100% + 1 x 30%)."""
        oficio = self._criar_oficio(ano=2026)
        cidade_destino = Cidade.objects.create(
            nome='Interior Pernoites', estado=self.estado, codigo_ibge='4113706'
        )
        oficio.viajantes.add(self.viajante_final)
        response = self.client.post(
            reverse('eventos:oficio-step3-calcular-diarias', kwargs={'pk': oficio.pk}),
            data=self._payload_step3(
                [cidade_destino],
                trechos=[
                    {
                        'saida_data': '2026-03-15',
                        'saida_hora': '14:00',
                        'chegada_data': '2026-03-15',
                        'chegada_hora': '17:30',
                    }
                ],
                retorno={
                    'saida_data': '2026-03-18',
                    'saida_hora': '08:00',
                    'chegada_data': '2026-03-18',
                    'chegada_hora': '11:30',
                },
            ),
        )
        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertTrue(payload['ok'])
        self.assertEqual(
            payload['totais']['total_diarias'],
            '3 x 100%',
            msg='3 pernoites (15â†’16, 16â†’17, 17â†’18) devem gerar 3 diÃ¡rias integrais',
        )
        self.assertEqual(payload['totais']['total_valor'], '871,65')

    def test_step4_nao_finaliza_sem_protocolo(self):
        oficio = self._criar_oficio(ano=2026)
        self._salvar_oficio_finalizavel(oficio, destino=self._criar_destino('Destino Sem Protocolo'))
        Oficio.objects.filter(pk=oficio.pk).update(protocolo='')

        response = self._post_finalizar_oficio(oficio)

        self.assertEqual(response.status_code, 200)
        oficio.refresh_from_db()
        self.assertEqual(oficio.status, Oficio.STATUS_RASCUNHO)
        self.assertContains(response, 'Informe o protocolo do ofÃ­cio.')

    def test_step4_nao_finaliza_sem_viajantes(self):
        oficio = self._criar_oficio(ano=2026)
        self._salvar_oficio_finalizavel(oficio, destino=self._criar_destino('Destino Sem Viajantes'))
        oficio.viajantes.clear()

        response = self._post_finalizar_oficio(oficio)

        self.assertEqual(response.status_code, 200)
        oficio.refresh_from_db()
        self.assertEqual(oficio.status, Oficio.STATUS_RASCUNHO)
        self.assertContains(response, 'Selecione ao menos um viajante.')

    def test_step4_nao_finaliza_sem_nome_da_instituicao_de_custeio(self):
        oficio = self._criar_oficio(ano=2026)
        self._salvar_oficio_finalizavel(oficio, destino=self._criar_destino('Destino Outra Instituicao'))
        Oficio.objects.filter(pk=oficio.pk).update(
            custeio_tipo=Oficio.CUSTEIO_OUTRA_INSTITUICAO,
            nome_instituicao_custeio='',
        )

        response = self._post_finalizar_oficio(oficio)

        self.assertEqual(response.status_code, 200)
        oficio.refresh_from_db()
        self.assertEqual(oficio.status, Oficio.STATUS_RASCUNHO)
        self.assertContains(response, 'Informe a instituiÃ§Ã£o responsÃ¡vel pelo custeio.')

    def test_step4_nao_finaliza_com_motorista_carona_sem_referencias(self):
        oficio = self._criar_oficio(ano=2026)
        self._salvar_oficio_finalizavel(oficio, destino=self._criar_destino('Destino Motorista Carona'))
        Oficio.objects.filter(pk=oficio.pk).update(
            motorista_carona=True,
            motorista_oficio='',
            motorista_oficio_numero=None,
            motorista_oficio_ano=None,
            motorista_protocolo='',
        )

        response = self._post_finalizar_oficio(oficio)

        self.assertEqual(response.status_code, 200)
        oficio.refresh_from_db()
        self.assertEqual(oficio.status, Oficio.STATUS_RASCUNHO)
        self.assertContains(response, 'Informe o nÃºmero do ofÃ­cio do motorista carona.')
        self.assertContains(response, 'Informe o protocolo do motorista carona.')

    def test_step4_nao_finaliza_com_step3_incompleto(self):
        oficio = self._criar_oficio(ano=2026)
        self._salvar_steps_1_e_2(oficio)

        response = self._post_finalizar_oficio(oficio)

        self.assertEqual(response.status_code, 200)
        oficio.refresh_from_db()
        self.assertEqual(oficio.status, Oficio.STATUS_RASCUNHO)
        self.assertContains(response, 'Preencha e salve o Step 3 antes de finalizar.')

    def test_step4_finaliza_com_dados_completos(self):
        oficio = self._criar_oficio(ano=2026)
        self._salvar_oficio_finalizavel(oficio, destino=self._criar_destino('Destino Completo'))
        before_update = oficio.updated_at

        response = self._post_finalizar_oficio(oficio)

        self.assertEqual(response.status_code, 302)
        oficio.refresh_from_db()
        self.assertEqual(oficio.status, Oficio.STATUS_FINALIZADO)
        self.assertGreater(oficio.updated_at, before_update)

    def test_step4_finalizar_com_termo_sim_gera_um_termo_por_viajante(self):
        oficio = self._criar_oficio(ano=2026)
        destino = self._criar_destino('Destino Termos Sim')
        self._salvar_oficio_finalizavel(oficio, destino=destino)
        viajante_extra = Viajante.objects.create(
            nome='Servidor Extra Termo',
            status=Viajante.STATUS_FINALIZADO,
            cargo=self.cargo,
            cpf='55544433322',
            rg='22334455',
            unidade_lotacao=self.unidade,
            telefone='41999991111',
        )
        oficio.viajantes.add(viajante_extra)

        response = self.client.post(
            reverse('eventos:oficio-step4', kwargs={'pk': oficio.pk}),
            data={'finalizar': '1', 'gerar_termo_preenchido': '1'},
        )

        self.assertEqual(response.status_code, 302)
        self.assertEqual(response.url, reverse('eventos:oficios-global'))
        oficio.refresh_from_db()
        self.assertEqual(oficio.status, Oficio.STATUS_FINALIZADO)
        termos = list(TermoAutorizacao.objects.filter(oficio=oficio).order_by('viajante__nome'))
        self.assertEqual(len(termos), 2)
        self.assertEqual({termo.viajante_id for termo in termos}, {self.viajante_final.pk, viajante_extra.pk})
        response_termos = self.client.get(reverse('eventos:documentos-termos'))
        self.assertContains(response_termos, 'Termo de autorizacao')
        self.assertContains(response_termos, self.viajante_final.nome)
        self.assertContains(response_termos, viajante_extra.nome)
        self.assertContains(response_termos, destino.nome)
        self.assertNotContains(response_termos, 'TA-000')

    def test_step4_finalizar_com_termo_sim_sincroniza_termos_existentes_sem_manter_sobra(self):
        oficio = self._criar_oficio(ano=2026)
        destino = self._criar_destino('Destino Termos Sync')
        self._salvar_oficio_finalizavel(oficio, destino=destino)
        viajante_antigo = Viajante.objects.create(
            nome='Servidor Antigo Termo',
            status=Viajante.STATUS_FINALIZADO,
            cargo=self.cargo,
            cpf='65432198700',
            rg='99887766',
            unidade_lotacao=self.unidade,
            telefone='41999992222',
        )
        termo_antigo = TermoAutorizacao.objects.create(
            evento=self.evento,
            oficio=oficio,
            modo_geracao=TermoAutorizacao.MODO_AUTOMATICO_SEM_VIATURA,
            status=TermoAutorizacao.STATUS_GERADO,
            viajante=viajante_antigo,
            destino=f'{destino.nome}/{self.estado.sigla}',
            data_evento=date(2026, 9, 1),
        )
        termo_antigo.oficios.add(oficio)

        response = self.client.post(
            reverse('eventos:oficio-step4', kwargs={'pk': oficio.pk}),
            data={'finalizar': '1', 'gerar_termo_preenchido': '1'},
        )

        self.assertEqual(response.status_code, 302)
        termos = list(TermoAutorizacao.objects.filter(oficio=oficio))
        self.assertEqual(len(termos), 1)
        self.assertEqual(termos[0].viajante_id, self.viajante_final.pk)
        termo_antigo.refresh_from_db()
        self.assertIsNone(termo_antigo.oficio_id)
        self.assertFalse(termo_antigo.oficios.filter(pk=oficio.pk).exists())

    def test_step4_finalizar_com_termo_nao_nao_gera_termos(self):
        oficio = self._criar_oficio(ano=2026)
        self._salvar_oficio_finalizavel(oficio, destino=self._criar_destino('Destino Termos Nao'))

        response = self.client.post(
            reverse('eventos:oficio-step4', kwargs={'pk': oficio.pk}),
            data={'finalizar': '1', 'gerar_termo_preenchido': '0'},
        )

        self.assertEqual(response.status_code, 302)
        self.assertEqual(response.url, reverse('eventos:oficios-global'))
        oficio.refresh_from_db()
        self.assertEqual(oficio.status, Oficio.STATUS_FINALIZADO)
        self.assertFalse(TermoAutorizacao.objects.filter(oficio=oficio).exists())

    def test_oficio_finalizado_permanece_editavel_nos_steps_1_2_e_3(self):
        oficio = self._criar_oficio(ano=2026)
        destino_inicial = self._criar_destino('Destino Inicial Finalizado')
        self._salvar_oficio_finalizavel(oficio, destino=destino_inicial)
        self.assertEqual(self._post_finalizar_oficio(oficio).status_code, 302)
        oficio.refresh_from_db()
        self.assertEqual(oficio.status, Oficio.STATUS_FINALIZADO)

        response_step1 = self.client.post(
            reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk}),
            data=self._payload_step1(oficio, motivo='Motivo editado com status finalizado'),
        )
        self.assertEqual(response_step1.status_code, 302)
        oficio.refresh_from_db()
        self.assertEqual(oficio.status, Oficio.STATUS_FINALIZADO)
        self.assertEqual(oficio.motivo, 'Motivo editado com status finalizado')

        response_step2 = self.client.post(
            reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}),
            data=self._payload_step2(placa='DEF-5678', modelo='VIATURA FINALIZADA'),
        )
        self.assertEqual(response_step2.status_code, 302)
        oficio.refresh_from_db()
        self.assertEqual(oficio.status, Oficio.STATUS_FINALIZADO)
        self.assertEqual(oficio.placa, 'DEF5678')

        destino_editado = self._criar_destino('Destino Editado Finalizado')
        response_step3 = self.client.post(
            reverse('eventos:oficio-step3', kwargs={'pk': oficio.pk}),
            data=self._payload_step3(
                [destino_editado],
                trechos=[
                    {
                        'saida_data': '2026-11-10',
                        'saida_hora': '09:00',
                        'chegada_data': '2026-11-10',
                        'chegada_hora': '13:00',
                    }
                ],
                retorno={
                    'saida_data': '2026-11-11',
                    'saida_hora': '08:00',
                    'chegada_data': '2026-11-11',
                    'chegada_hora': '17:00',
                },
            ),
        )
        self.assertEqual(response_step3.status_code, 302)
        oficio.refresh_from_db()
        self.assertEqual(oficio.status, Oficio.STATUS_FINALIZADO)
        self.assertEqual(oficio.trechos.count(), 1)


class OficioJustificativaTest(TestCase):
    """Fase 3 do ofÃ­cio: justificativa obrigatÃ³ria, tela prÃ³pria e modelos."""

    def setUp(self):
        self.user = User.objects.create_user(username='justificativa', password='justificativa')
        self.client = Client()
        self.client.login(username='justificativa', password='justificativa')
        self.estado = Estado.objects.create(nome='ParanÃ¡', sigla='PR', codigo_ibge='41')
        self.cidade = Cidade.objects.create(nome='Curitiba', estado=self.estado, codigo_ibge='4106902')
        self.evento = Evento.objects.create(
            titulo='Evento Justificativa',
            data_inicio=date(2026, 10, 1),
            data_fim=date(2026, 10, 3),
            status=Evento.STATUS_RASCUNHO,
        )
        self.cargo = Cargo.objects.create(nome='ANALISTA JUSTIFICATIVA', is_padrao=True)
        self.unidade = UnidadeLotacao.objects.create(nome='DPC JUSTIFICATIVA')
        self.viajante = Viajante.objects.create(
            nome='Viajante Justificativa',
            status=Viajante.STATUS_FINALIZADO,
            cargo=self.cargo,
            cpf='12345678901',
            telefone='41999990000',
            unidade_lotacao=self.unidade,
            rg='1234567890',
        )
        self.combustivel = CombustivelVeiculo.objects.create(nome='GASOLINA', is_padrao=True)

    def _criar_destino(self, nome='Destino Justificativa'):
        codigo = str(4108000 + Cidade.objects.count())
        return Cidade.objects.create(nome=nome, estado=self.estado, codigo_ibge=codigo)

    def _criar_oficio(self, data_criacao=None):
        oficio = Oficio.objects.create(evento=self.evento, status=Oficio.STATUS_RASCUNHO)
        if data_criacao is not None:
            Oficio.objects.filter(pk=oficio.pk).update(data_criacao=data_criacao)
            oficio.refresh_from_db()
        return oficio

    def _payload_step1(self, oficio):
        oficio.refresh_from_db()
        return {
            'oficio_numero': oficio.numero_formatado,
            'protocolo': '12.345.678-9',
            'data_criacao': oficio.data_criacao.strftime('%d/%m/%Y'),
            'modelo_motivo': '',
            'motivo': 'Motivo com justificativa',
            'assunto_tipo': Oficio.ASSUNTO_TIPO_AUTORIZACAO,
            'custeio_tipo': Oficio.CUSTEIO_UNIDADE,
            'nome_instituicao_custeio': '',
            'viajantes': [self.viajante.pk],
        }

    def _payload_step2(self):
        return {
            'placa': 'ABC-1234',
            'modelo': 'VIATURA JUSTIFICATIVA',
            'combustivel': 'GASOLINA',
            'tipo_viatura': Oficio.TIPO_VIATURA_DESCARACTERIZADA,
            'porte_transporte_armas': '1',
            'motorista_viajante': '__manual__',
            'motorista_nome': 'Motorista Justificativa',
            'motorista_oficio_numero': '5',
            'motorista_oficio_ano': str(tz.localdate().year),
            'motorista_protocolo': '12.345.678-9',
        }

    def _payload_step3(self, destino, data_saida, data_retorno):
        return {
            'roteiro_modo': Oficio.ROTEIRO_MODO_PROPRIO,
            'sede_estado': str(self.estado.pk),
            'sede_cidade': str(self.cidade.pk),
            'trecho_0_origem_estado_id': str(self.estado.pk),
            'trecho_0_origem_cidade_id': str(self.cidade.pk),
            'trecho_0_destino_estado_id': str(destino.estado_id),
            'trecho_0_destino_cidade_id': str(destino.pk),
            'trecho_0_origem_nome': f'{self.cidade.nome}/{self.estado.sigla}',
            'trecho_0_destino_nome': f'{destino.nome}/{self.estado.sigla}',
            'destino_estado_0': str(destino.estado_id),
            'destino_cidade_0': str(destino.pk),
            'trecho_0_saida_data': data_saida.strftime('%Y-%m-%d'),
            'trecho_0_saida_hora': '08:00',
            'trecho_0_chegada_data': data_saida.strftime('%Y-%m-%d'),
            'trecho_0_chegada_hora': '12:00',
            'retorno_saida_cidade': f'{destino.nome}/{self.estado.sigla}',
            'retorno_chegada_cidade': f'{self.cidade.nome}/{self.estado.sigla}',
            'retorno_saida_data': data_retorno.strftime('%Y-%m-%d'),
            'retorno_saida_hora': '09:00',
            'retorno_chegada_data': data_retorno.strftime('%Y-%m-%d'),
            'retorno_chegada_hora': '18:00',
        }

    def _salvar_oficio_finalizavel(self, oficio, data_saida, data_retorno):
        destino = self._criar_destino()
        response_step1 = self.client.post(
            reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk}),
            data=self._payload_step1(oficio),
        )
        self.assertEqual(response_step1.status_code, 302)
        response_step2 = self.client.post(
            reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}),
            data=self._payload_step2(),
        )
        self.assertEqual(response_step2.status_code, 302)
        response_step3 = self.client.post(
            reverse('eventos:oficio-step3', kwargs={'pk': oficio.pk}),
            data=self._payload_step3(destino, data_saida, data_retorno),
        )
        self.assertEqual(response_step3.status_code, 302)
        oficio.refresh_from_db()
        return destino

    def _salvar_justificativa(self, oficio, texto, modelo=None):
        justificativa, _ = Justificativa.objects.update_or_create(
            oficio=oficio,
            defaults={'modelo': modelo, 'texto': texto},
        )
        return justificativa

    def test_prazo_justificativa_usa_fallback_10_quando_configuracao_ausente(self):
        self.assertFalse(ConfiguracaoSistema.objects.exists())
        self.assertEqual(get_prazo_justificativa_dias(), 10)

    def test_step4_finaliza_sem_justificativa_quando_antecedencia_atende_prazo(self):
        ConfiguracaoSistema.objects.create(cidade_sede_padrao=self.cidade, prazo_justificativa_dias=10)
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_oficio_finalizavel(
            oficio,
            data_saida=date(2026, 10, 10),
            data_retorno=date(2026, 10, 11),
        )

        self.assertFalse(oficio_exige_justificativa(oficio))
        response = self.client.post(reverse('eventos:oficio-step4', kwargs={'pk': oficio.pk}), data={'finalizar': '1'})

        self.assertRedirects(response, reverse('eventos:oficios-global'))
        oficio.refresh_from_db()
        self.assertEqual(oficio.status, Oficio.STATUS_FINALIZADO)

    def test_step4_redireciona_para_justificativa_quando_ela_e_obrigatoria(self):
        ConfiguracaoSistema.objects.create(cidade_sede_padrao=self.cidade, prazo_justificativa_dias=10)
        oficio = self._criar_oficio(data_criacao=date(2026, 10, 5))
        self._salvar_oficio_finalizavel(
            oficio,
            data_saida=date(2026, 10, 10),
            data_retorno=date(2026, 10, 11),
        )

        self.assertTrue(oficio_exige_justificativa(oficio))
        response = self.client.post(reverse('eventos:oficio-step4', kwargs={'pk': oficio.pk}), data={'finalizar': '1'})

        self.assertEqual(response.status_code, 302)
        self.assertIn(reverse('eventos:oficio-justificativa', kwargs={'pk': oficio.pk}), response.url)
        self.assertIn('next=', response.url)
        oficio.refresh_from_db()
        self.assertEqual(oficio.status, Oficio.STATUS_RASCUNHO)

    def test_step4_finaliza_quando_justificativa_exigida_esta_preenchida(self):
        ConfiguracaoSistema.objects.create(cidade_sede_padrao=self.cidade, prazo_justificativa_dias=10)
        oficio = self._criar_oficio(data_criacao=date(2026, 10, 5))
        self._salvar_oficio_finalizavel(
            oficio,
            data_saida=date(2026, 10, 10),
            data_retorno=date(2026, 10, 11),
        )
        self._salvar_justificativa(oficio, 'Necessidade operacional urgente.')

        response = self.client.post(reverse('eventos:oficio-step4', kwargs={'pk': oficio.pk}), data={'finalizar': '1'})

        self.assertRedirects(response, reverse('eventos:oficios-global'))
        oficio.refresh_from_db()
        self.assertEqual(oficio.status, Oficio.STATUS_FINALIZADO)

    def test_step4_autosave_persiste_termo_de_autorizacao(self):
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_oficio_finalizavel(
            oficio,
            data_saida=date(2026, 10, 10),
            data_retorno=date(2026, 10, 11),
        )
        url = reverse('eventos:oficio-step4', kwargs={'pk': oficio.pk})

        response = self.client.post(
            url,
            data={'autosave': '1', 'gerar_termo_preenchido': '1'},
            HTTP_X_REQUESTED_WITH='XMLHttpRequest',
        )

        self.assertEqual(response.status_code, 200)
        oficio.refresh_from_db()
        self.assertTrue(oficio.gerar_termo_preenchido)

        reopened = self.client.get(url)

        self.assertEqual(reopened.status_code, 200)
        self.assertContains(reopened, 'id="id_gerar_termo_sim" value="1" checked')
        self.assertNotContains(reopened, 'id="id_gerar_termo_nao" value="0" checked')

    @unittest.skip('AsserÃ§Ã£o legada com encoding corrompido; coberta pelo teste limpo abaixo.')
    def test_step4_salvar_oficio_manual_persiste_escolha_do_termo(self):
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_oficio_finalizavel(
            oficio,
            data_saida=date(2026, 10, 10),
            data_retorno=date(2026, 10, 11),
        )
        url = reverse('eventos:oficio-step4', kwargs={'pk': oficio.pk})

        response = self.client.post(
            url,
            data={'salvar_oficio': '1', 'gerar_termo_preenchido': '1'},
            follow=True,
        )

        self.assertRedirects(response, url)
        oficio.refresh_from_db()
        self.assertTrue(oficio.gerar_termo_preenchido)
        self.assertContains(response, 'OfÃ­cio salvo.')

    def test_step4_salvar_oficio_manual_persiste_escolha_do_termo_sem_artefato_de_encoding(self):
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_oficio_finalizavel(
            oficio,
            data_saida=date(2026, 10, 10),
            data_retorno=date(2026, 10, 11),
        )
        url = reverse('eventos:oficio-step4', kwargs={'pk': oficio.pk})

        response = self.client.post(
            url,
            data={'salvar_oficio': '1', 'gerar_termo_preenchido': '1'},
            follow=True,
        )

        self.assertRedirects(response, url)
        oficio.refresh_from_db()
        self.assertTrue(oficio.gerar_termo_preenchido)
        self.assertContains(response, 'salvo.')

    def test_assets_do_wizard_mantem_header_sticky_e_resumo_essencial_compacto(self):
        js = (Path(settings.BASE_DIR) / 'static' / 'js' / 'oficio_wizard.js').read_text(encoding='utf-8')
        css = (Path(settings.BASE_DIR) / 'static' / 'css' / 'style.css').read_text(encoding='utf-8')

        self.assertIn("header.style.removeProperty('top');", js)
        self.assertIn("page.style.removeProperty('--oficio-sticky-header-top');", js)
        self.assertIn('bindGlanceDrawer', js)
        self.assertIn('data-oficio-glance-toggle', js)
        self.assertIn('data-glance-state', js)
        self.assertIn('setGlanceValue', js)
        self.assertIn('renderGlanceTravelers', js)
        self.assertIn("var statusElement = options.statusElement || null;", js)
        self.assertNotIn('Autosave salvo', js)
        self.assertNotIn('Autosave ativo', js)
        self.assertNotIn("window.addEventListener('scroll', refreshLayout, { passive: true })", js)
        self.assertNotIn("document.addEventListener('scroll', refreshLayout, { passive: true })", js)
        self.assertNotIn("page.style.setProperty('--oficio-sticky-header-top'", js)
        self.assertNotIn('syncQuickReportLayout', js)
        self.assertIn('.oficio-wizard-shell {', css)
        self.assertIn('.oficio-glance-panel {', css)
        self.assertIn('.oficio-glance-layout {', css)
        self.assertIn('.oficio-glance-row {', css)
        self.assertIn('.oficio-wizard-submenu-trigger {', css)
        self.assertIn('.oficio-wizard-footer-actions {', css)
        self.assertIn('.oficio-wizard-page[data-glance-state="open"] {', css)
        self.assertIn('--oficio-wizard-sticky-top: 0.9rem;', css)
        self.assertIn('top: var(--oficio-wizard-sticky-top);', css)
        self.assertIn('repeat(var(--oficio-step-count, 5), minmax(0, 1fr))', css)
        self.assertIn('transform: none !important;', css)
        self.assertIn('overflow: visible;', css)
        self.assertIn('.oficio-stepper-link.is-active {', css)
        self.assertNotIn('.oficio-quick-report', css)

    def test_steps_carregam_header_sticky_resumo_essencial_e_script_compartilhado(self):
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_oficio_finalizavel(
            oficio,
            data_saida=date(2026, 10, 10),
            data_retorno=date(2026, 10, 11),
        )
        urls = [
            reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk}),
            reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}),
            reverse('eventos:oficio-step3', kwargs={'pk': oficio.pk}),
            reverse('eventos:oficio-justificativa', kwargs={'pk': oficio.pk}),
            reverse('eventos:oficio-step4', kwargs={'pk': oficio.pk}),
        ]
        for url in urls:
            response = self.client.get(url)
            self.assertEqual(response.status_code, 200)
            self.assertContains(response, 'data-oficio-sticky-header')
            self.assertContains(response, 'data-glance-state="closed"')
            self.assertContains(response, 'data-oficio-glance')
            self.assertContains(response, 'data-oficio-glance-toggle')
            self.assertContains(response, 'data-oficio-glance-rail="1"')
            self.assertContains(response, 'id="oficio-glance-panel"')
            self.assertContains(response, 'aria-hidden="true"')
            self.assertContains(response, 'oficio-wizard-shell')
            self.assertContains(response, 'style="--oficio-step-count: 5;"')
            self.assertContains(response, 'data-step-key="justificativa"')
            self.assertContains(response, 'data-oficio-footer-actions="1"')
            self.assertContains(response, 'id="summary-oficio"')
            self.assertContains(response, 'id="summary-protocolo"')
            self.assertContains(response, 'id="summary-viajantes-meta"')
            self.assertContains(response, 'summary-viajantes-heading')
            self.assertContains(response, 'id="summary-destino"')
            self.assertContains(response, 'id="summary-data-evento"')
            self.assertNotContains(response, 'id="summary-veiculo"')
            self.assertNotContains(response, 'id="summary-motorista"')
            self.assertNotContains(response, 'RelatÃ³rio rÃ¡pido')
            self.assertNotContains(response, 'oficio-quick-report')
            self.assertNotContains(response, 'oficio-wizard-submenu-actions')
            self.assertContains(response, 'js/oficio_wizard.js')
            content = response.content.decode('utf-8')
            self.assertEqual(content.count('data-oficio-sticky-header'), 1)
            self.assertEqual(content.count('id="oficio-glance-panel"'), 1)

    def test_resumo_essencial_renderiza_dados_chave_sem_campos_legados(self):
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        destino = self._salvar_oficio_finalizavel(
            oficio,
            data_saida=date(2026, 10, 10),
            data_retorno=date(2026, 10, 11),
        )

        response = self.client.get(reverse('eventos:oficio-step3', kwargs={'pk': oficio.pk}))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'id="summary-oficio"')
        self.assertContains(response, oficio.numero_formatado)
        self.assertContains(response, 'id="summary-protocolo"')
        self.assertContains(response, '12.345.678-9')
        self.assertContains(response, 'id="summary-viajantes-meta"')
        self.assertContains(response, 'id="summary-destino"')
        self.assertContains(response, destino.nome)
        self.assertContains(response, 'id="summary-data-evento"')
        self.assertEqual(response.context['wizard_glance']['data_evento'], '01/10/2026 atÃ© 03/10/2026')
        self.assertNotContains(response, 'id="summary-veiculo"')
        self.assertNotContains(response, 'id="summary-motorista"')
        self.assertNotContains(response, 'Motivo</dt>')
        self.assertNotContains(response, 'Custeio</dt>')

    def test_protocolos_do_wizard_exibem_mascara_curta_correta(self):
        oficio = self._criar_oficio()

        response_step1 = self.client.get(reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk}))
        self.assertEqual(response_step1.status_code, 200)
        self.assertContains(response_step1, 'placeholder="12.345.678-9"')
        self.assertContains(response_step1, 'maxlength="12"')

        response_step2 = self.client.get(reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}))
        self.assertEqual(response_step2.status_code, 200)
        self.assertContains(response_step2, 'placeholder="12.345.678-9"')
        self.assertContains(response_step2, 'maxlength="12"')

    def test_salvar_justificativa_grava_texto_e_retorna_para_next(self):
        oficio = self._criar_oficio(data_criacao=date(2026, 10, 5))
        self._salvar_oficio_finalizavel(
            oficio,
            data_saida=date(2026, 10, 10),
            data_retorno=date(2026, 10, 11),
        )
        next_url = reverse('eventos:oficio-step4', kwargs={'pk': oficio.pk})
        response = self.client.post(
            reverse('eventos:oficio-justificativa', kwargs={'pk': oficio.pk}),
            data={
                'modelo_justificativa': '',
                'justificativa_texto': 'Texto final da justificativa.',
                'next': next_url,
            },
        )

        self.assertRedirects(response, next_url)
        oficio.refresh_from_db()
        self.assertEqual(oficio.justificativa.texto, 'Texto final da justificativa.')

    def test_tela_justificativa_preseleciona_modelo_padrao_quando_texto_ainda_vazio(self):
        oficio = self._criar_oficio(data_criacao=date(2026, 10, 5))
        modelo = ModeloJustificativa.objects.create(
            nome='Modelo PadrÃ£o Justificativa',
            texto='Texto padrÃ£o da justificativa.',
            padrao=True,
        )

        response = self.client.get(reverse('eventos:oficio-justificativa', kwargs={'pk': oficio.pk}))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, f'<option value="{modelo.pk}" selected>')
        self.assertContains(response, 'Texto padrÃ£o da justificativa.')

    def test_tela_justificativa_carrega_modelo_e_texto_do_modelo_novo(self):
        oficio = self._criar_oficio(data_criacao=date(2026, 10, 5))
        modelo = ModeloJustificativa.objects.create(
            nome='Modelo Existente',
            texto='Texto base do modelo.',
        )
        self._salvar_justificativa(oficio, 'Texto persistido da justificativa.', modelo=modelo)

        response = self.client.get(reverse('eventos:oficio-justificativa', kwargs={'pk': oficio.pk}))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, f'<option value="{modelo.pk}" selected>')
        self.assertContains(response, 'Texto persistido da justificativa.')
        self.assertFalse(hasattr(oficio, 'justificativa_texto'))
        self.assertFalse(hasattr(oficio, 'justificativa_modelo'))

    def test_tela_justificativa_nao_quebra_sem_trechos_validos(self):
        oficio = self._criar_oficio()

        response = self.client.get(reverse('eventos:oficio-justificativa', kwargs={'pk': oficio.pk}))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'id="summary-destino"')
        self.assertContains(response, 'id="summary-data-evento"')

    def test_tela_justificativa_reorganiza_modelo_texto_e_remove_resumo(self):
        oficio = self._criar_oficio()

        response = self.client.get(reverse('eventos:oficio-justificativa', kwargs={'pk': oficio.pk}))

        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, 'Resumo do texto')
        html = response.content.decode('utf-8')
        self.assertLess(html.index('Modelo de justificativa'), html.index('Texto da justificativa'))
        self.assertContains(response, 'id="oficio-glance-panel"')
        self.assertNotContains(response, 'RelatÃ³rio rÃ¡pido')

    def test_lista_modelos_justificativa_ordena_por_nome(self):
        ModeloJustificativa.objects.create(nome='ZZZ', texto='z')
        ModeloJustificativa.objects.create(nome='AAA', texto='a')

        response = self.client.get(reverse('eventos:modelos-justificativa-lista'))

        self.assertEqual(response.status_code, 200)
        html = response.content.decode()
        self.assertLess(html.index('AAA'), html.index('ZZZ'))

    def test_modelo_justificativa_pode_ser_editado_e_excluido(self):
        modelo = ModeloJustificativa.objects.create(nome='Modelo EditÃ¡vel', texto='Texto original')

        response_edicao = self.client.post(
            reverse('eventos:modelos-justificativa-editar', kwargs={'pk': modelo.pk}),
            data={'nome': 'Modelo Editado', 'texto': 'Texto editado', 'next': ''},
        )

        self.assertRedirects(response_edicao, reverse('eventos:modelos-justificativa-lista'))
        modelo.refresh_from_db()
        self.assertEqual(modelo.nome, 'Modelo Editado')
        self.assertEqual(modelo.texto, 'Texto editado')

        response_exclusao = self.client.post(
            reverse('eventos:modelos-justificativa-excluir', kwargs={'pk': modelo.pk}),
            data={'next': ''},
        )

        self.assertRedirects(response_exclusao, reverse('eventos:modelos-justificativa-lista'))
        self.assertFalse(ModeloJustificativa.objects.filter(pk=modelo.pk).exists())

    def test_definir_padrao_desmarca_modelo_anterior_de_justificativa(self):
        modelo_1 = ModeloJustificativa.objects.create(nome='Modelo 1', texto='1', padrao=True)
        modelo_2 = ModeloJustificativa.objects.create(nome='Modelo 2', texto='2', padrao=False)

        response = self.client.post(reverse('eventos:modelos-justificativa-definir-padrao', kwargs={'pk': modelo_2.pk}))

        self.assertRedirects(response, reverse('eventos:modelos-justificativa-lista'))
        modelo_1.refresh_from_db()
        modelo_2.refresh_from_db()
        self.assertFalse(modelo_1.padrao)
        self.assertTrue(modelo_2.padrao)


class OficioDocumentosTest(TestCase):
    """Fase 4 do ofÃ­cio: contexto, validaÃ§Ã£o e geraÃ§Ã£o documental."""

    def setUp(self):
        reset_document_backend_capabilities_cache()
        self.addCleanup(reset_document_backend_capabilities_cache)
        self.user = User.objects.create_user(username='documentos', password='documentos')
        self.client = Client()
        self.client.login(username='documentos', password='documentos')
        self.estado = Estado.objects.create(nome='ParanÃ¡', sigla='PR', codigo_ibge='41')
        self.cidade = Cidade.objects.create(nome='Curitiba', estado=self.estado, codigo_ibge='4106902')
        self.evento = Evento.objects.create(
            titulo='Evento Documental',
            data_inicio=date(2026, 10, 1),
            data_fim=date(2026, 10, 3),
            status=Evento.STATUS_RASCUNHO,
        )
        self.cargo = Cargo.objects.create(nome='ANALISTA DOCUMENTAL', is_padrao=True)
        self.unidade = UnidadeLotacao.objects.create(nome='DPC DOCUMENTAL')
        self.viajante = Viajante.objects.create(
            nome='Viajante Documental',
            status=Viajante.STATUS_FINALIZADO,
            cargo=self.cargo,
            cpf='12345678901',
            telefone='41999990000',
            unidade_lotacao=self.unidade,
            rg='1234567890',
        )
        self.assinante = Viajante.objects.create(
            nome='Autoridade Assinante',
            status=Viajante.STATUS_FINALIZADO,
            cargo=self.cargo,
            cpf='10987654321',
            telefone='41999991111',
            unidade_lotacao=self.unidade,
            rg='9988776655',
        )
        self.combustivel = CombustivelVeiculo.objects.create(nome='GASOLINA', is_padrao=True)

    def _criar_configuracao(self):
        config = ConfiguracaoSistema.objects.create(
            cidade_sede_padrao=self.cidade,
            prazo_justificativa_dias=10,
            nome_orgao='PolÃ­cia Civil do ParanÃ¡',
            sigla_orgao='PCPR',
            divisao='Diretoria de PolÃ­cia do Interior',
            unidade='Departamento de PolÃ­cia Civil',
            logradouro='Rua Exemplo',
            numero='123',
            bairro='Centro',
            cidade_endereco='Curitiba',
            uf='PR',
            cep='80000000',
            telefone='4133334444',
            email='documentos@pc.pr.gov.br',
        )
        AssinaturaConfiguracao.objects.create(
            configuracao=config,
            tipo=AssinaturaConfiguracao.TIPO_OFICIO,
            ordem=1,
            viajante=self.assinante,
            ativo=True,
        )
        AssinaturaConfiguracao.objects.create(
            configuracao=config,
            tipo=AssinaturaConfiguracao.TIPO_JUSTIFICATIVA,
            ordem=1,
            viajante=self.assinante,
            ativo=True,
        )
        AssinaturaConfiguracao.objects.create(
            configuracao=config,
            tipo=AssinaturaConfiguracao.TIPO_PLANO_TRABALHO,
            ordem=1,
            viajante=self.assinante,
            ativo=True,
        )
        AssinaturaConfiguracao.objects.create(
            configuracao=config,
            tipo=AssinaturaConfiguracao.TIPO_ORDEM_SERVICO,
            ordem=1,
            viajante=self.assinante,
            ativo=True,
        )
        AssinaturaConfiguracao.objects.create(
            configuracao=config,
            tipo=AssinaturaConfiguracao.TIPO_TERMO_AUTORIZACAO,
            ordem=1,
            viajante=self.assinante,
            ativo=True,
        )
        return config

    def _criar_destino(self, nome='Destino Documental'):
        codigo = str(4109000 + Cidade.objects.count())
        return Cidade.objects.create(nome=nome, estado=self.estado, codigo_ibge=codigo)

    def _criar_oficio(self, data_criacao=None):
        oficio = Oficio.objects.create(evento=self.evento, status=Oficio.STATUS_RASCUNHO)
        if data_criacao is not None:
            Oficio.objects.filter(pk=oficio.pk).update(data_criacao=data_criacao)
            oficio.refresh_from_db()
        return oficio

    def _payload_step1(self, oficio):
        oficio.refresh_from_db()
        return {
            'oficio_numero': oficio.numero_formatado,
            'protocolo': '12.345.678-9',
            'data_criacao': oficio.data_criacao.strftime('%d/%m/%Y'),
            'modelo_motivo': '',
            'motivo': 'Motivo documental',
            'assunto_tipo': Oficio.ASSUNTO_TIPO_AUTORIZACAO,
            'custeio_tipo': Oficio.CUSTEIO_UNIDADE,
            'nome_instituicao_custeio': '',
            'viajantes': [self.viajante.pk],
        }

    def _payload_step2(self):
        return {
            'placa': 'ABC-1234',
            'modelo': 'VIATURA DOCUMENTAL',
            'combustivel': 'GASOLINA',
            'tipo_viatura': Oficio.TIPO_VIATURA_DESCARACTERIZADA,
            'porte_transporte_armas': '1',
            'motorista_viajante': '__manual__',
            'motorista_nome': 'Motorista Documental',
            'motorista_oficio_numero': '12',
            'motorista_oficio_ano': str(tz.localdate().year),
            'motorista_protocolo': '12.345.678-9',
        }

    def _payload_step3(self, destino, data_saida, data_retorno):
        return {
            'roteiro_modo': Oficio.ROTEIRO_MODO_PROPRIO,
            'sede_estado': str(self.estado.pk),
            'sede_cidade': str(self.cidade.pk),
            'trecho_0_origem_estado_id': str(self.estado.pk),
            'trecho_0_origem_cidade_id': str(self.cidade.pk),
            'trecho_0_destino_estado_id': str(destino.estado_id),
            'trecho_0_destino_cidade_id': str(destino.pk),
            'trecho_0_origem_nome': f'{self.cidade.nome}/{self.estado.sigla}',
            'trecho_0_destino_nome': f'{destino.nome}/{self.estado.sigla}',
            'destino_estado_0': str(destino.estado_id),
            'destino_cidade_0': str(destino.pk),
            'trecho_0_saida_data': data_saida.strftime('%Y-%m-%d'),
            'trecho_0_saida_hora': '08:00',
            'trecho_0_chegada_data': data_saida.strftime('%Y-%m-%d'),
            'trecho_0_chegada_hora': '12:00',
            'retorno_saida_cidade': f'{destino.nome}/{self.estado.sigla}',
            'retorno_chegada_cidade': f'{self.cidade.nome}/{self.estado.sigla}',
            'retorno_saida_data': data_retorno.strftime('%Y-%m-%d'),
            'retorno_saida_hora': '09:00',
            'retorno_chegada_data': data_retorno.strftime('%Y-%m-%d'),
            'retorno_chegada_hora': '18:00',
        }

    def _salvar_steps_1_e_2(self, oficio):
        response_step1 = self.client.post(
            reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk}),
            data=self._payload_step1(oficio),
        )
        self.assertEqual(response_step1.status_code, 302)
        response_step2 = self.client.post(
            reverse('eventos:oficio-step2', kwargs={'pk': oficio.pk}),
            data=self._payload_step2(),
        )
        self.assertEqual(response_step2.status_code, 302)

    def _salvar_oficio_finalizavel(self, oficio, data_saida, data_retorno):
        destino = self._criar_destino()
        self._salvar_steps_1_e_2(oficio)
        response_step3 = self.client.post(
            reverse('eventos:oficio-step3', kwargs={'pk': oficio.pk}),
            data=self._payload_step3(destino, data_saida, data_retorno),
        )
        if response_step3.status_code != 302:
            errors = []
            try:
                errors = list(response_step3.context.get('validation_errors') or [])
            except Exception:
                errors = []
            self.fail(
                f'Step 3 deveria redirecionar (302), mas retornou {response_step3.status_code}. '
                f'Erros: {errors}'
            )
        oficio.refresh_from_db()
        return destino

    def _extract_docx_text(self, payload):
        if DocxDocument is None:
            self.skipTest('python-docx nÃ£o estÃ¡ instalado neste ambiente de teste.')
        document = DocxDocument(BytesIO(payload))
        texts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        texts.append(cell.text)
        return '\n'.join(texts)

    def test_contexto_documental_do_oficio(self):
        self._criar_configuracao()
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        destino = self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))

        context = build_oficio_document_context(oficio)

        self.assertEqual(context['identificacao']['numero_formatado'], oficio.numero_formatado)
        self.assertEqual(context['evento']['titulo'], self.evento.titulo)
        self.assertEqual(context['viajantes'][0]['nome'], self.viajante.nome)
        self.assertEqual(context['veiculo']['placa_formatada'], 'ABC-1234')
        self.assertEqual(context['motorista']['nome'], 'MOTORISTA DOCUMENTAL')
        self.assertEqual(context['roteiro']['destinos'], [f'{destino.nome}/{self.estado.sigla}'])
        self.assertEqual(context['institucional']['sigla_orgao'], 'PCPR')
        self.assertEqual(len(context['assinaturas']), 1)

    def test_contexto_documental_do_oficio_exclui_sede_dos_destinos_mesmo_com_retorno_em_trechos(self):
        self._criar_configuracao()
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        destino = Cidade.objects.create(nome='Colombo Documento', estado=self.estado, codigo_ibge='4105808')
        oficio.estado_sede = self.estado
        oficio.cidade_sede = self.cidade
        oficio.retorno_saida_cidade = f'{destino.nome}/{self.estado.sigla}'
        oficio.retorno_chegada_cidade = f'{self.cidade.nome}/{self.estado.sigla}'
        oficio.retorno_saida_data = date(2026, 10, 10)
        oficio.retorno_saida_hora = datetime.strptime('17:30', '%H:%M').time()
        oficio.retorno_chegada_data = date(2026, 10, 10)
        oficio.retorno_chegada_hora = datetime.strptime('18:10', '%H:%M').time()
        oficio.save()
        OficioTrecho.objects.create(
            oficio=oficio,
            ordem=0,
            origem_estado=self.estado,
            origem_cidade=self.cidade,
            destino_estado=self.estado,
            destino_cidade=destino,
            saida_data=date(2026, 10, 10),
            saida_hora=datetime.strptime('08:00', '%H:%M').time(),
            chegada_data=date(2026, 10, 10),
            chegada_hora=datetime.strptime('08:40', '%H:%M').time(),
        )
        OficioTrecho.objects.create(
            oficio=oficio,
            ordem=1,
            origem_estado=self.estado,
            origem_cidade=destino,
            destino_estado=self.estado,
            destino_cidade=self.cidade,
            saida_data=date(2026, 10, 10),
            saida_hora=datetime.strptime('17:30', '%H:%M').time(),
            chegada_data=date(2026, 10, 10),
            chegada_hora=datetime.strptime('18:10', '%H:%M').time(),
        )

        context = build_oficio_document_context(oficio)

        self.assertEqual(context['roteiro']['destinos'], [f'{destino.nome}/{self.estado.sigla}'])

    def test_template_oficio_nao_duplica_ultimo_retorno_quando_origem_explicita_vem_vazia(self):
        self._criar_configuracao()
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        destino = Cidade.objects.create(nome='Maringa Documento', estado=self.estado, codigo_ibge='4115200')

        oficio.estado_sede = self.estado
        oficio.cidade_sede = self.cidade
        oficio.retorno_saida_cidade = ''
        oficio.retorno_chegada_cidade = f'{self.cidade.nome}/{self.estado.sigla}'
        oficio.retorno_saida_data = date(2026, 10, 10)
        oficio.retorno_saida_hora = datetime.strptime('17:30', '%H:%M').time()
        oficio.retorno_chegada_data = date(2026, 10, 10)
        oficio.retorno_chegada_hora = datetime.strptime('18:10', '%H:%M').time()
        oficio.save()

        OficioTrecho.objects.create(
            oficio=oficio,
            ordem=0,
            origem_estado=self.estado,
            origem_cidade=self.cidade,
            destino_estado=self.estado,
            destino_cidade=destino,
            saida_data=date(2026, 10, 10),
            saida_hora=datetime.strptime('08:00', '%H:%M').time(),
            chegada_data=date(2026, 10, 10),
            chegada_hora=datetime.strptime('08:40', '%H:%M').time(),
        )
        OficioTrecho.objects.create(
            oficio=oficio,
            ordem=1,
            origem_estado=self.estado,
            origem_cidade=destino,
            destino_estado=self.estado,
            destino_cidade=self.cidade,
            saida_data=date(2026, 10, 10),
            saida_hora=datetime.strptime('17:30', '%H:%M').time(),
            chegada_data=date(2026, 10, 10),
            chegada_hora=datetime.strptime('18:10', '%H:%M').time(),
        )

        mapping = build_oficio_template_context(oficio)

        self.assertEqual(mapping['col_volta_saida'], f'SaÃ­da {destino.nome}/{self.estado.sigla}: 10/10/2026 17:30')
        self.assertEqual(mapping['col_volta_chegada'], f'Chegada {self.cidade.nome}/{self.estado.sigla}: 10/10/2026 18:10')

    def test_contexto_documental_da_justificativa(self):
        self._criar_configuracao()
        oficio = self._criar_oficio(data_criacao=date(2026, 10, 5))
        self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))
        Justificativa.objects.update_or_create(
            oficio=oficio,
            defaults={'texto': 'Justificativa documental.'},
        )
        oficio.refresh_from_db()

        context = build_justificativa_document_context(oficio)

        self.assertTrue(context['justificativa']['exigida'])
        self.assertEqual(context['justificativa']['dias_antecedencia'], 5)
        self.assertEqual(context['conteudo']['justificativa_texto'], 'Justificativa documental.')
        self.assertEqual(len(context['assinaturas']), 1)

    def test_contexto_documental_do_termo_autorizacao(self):
        self._criar_configuracao()
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        destino = self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))

        context = build_termo_autorizacao_document_context(oficio)

        self.assertEqual(context['termo']['destinos_texto'], f'{destino.nome}/{self.estado.sigla}')
        self.assertIn('10/10/2026', context['termo']['periodo_viagem'])
        self.assertEqual(context['assinaturas'], [])

    def test_contexto_documental_do_plano_trabalho(self):
        self._criar_configuracao()
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))

        context = build_plano_trabalho_document_context(oficio)

        self.assertEqual(context['plano_trabalho']['objetivo'], 'Motivo documental')
        self.assertTrue(context['plano_trabalho']['roteiro_resumo'])
        self.assertIn('UNIDADE', context['plano_trabalho']['custeio_resumo'])

    def test_contexto_documental_da_ordem_servico(self):
        self._criar_configuracao()
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        destino = self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))

        context = build_ordem_servico_document_context(oficio)

        self.assertEqual(context['ordem_servico']['finalidade'], 'Motivo documental')
        self.assertIn(destino.nome, context['ordem_servico']['destinos_texto'])
        self.assertEqual(context['ordem_servico']['participantes_texto'], self.viajante.nome)

    def test_mapping_oficio_formata_placeholders_humanos_em_title_case(self):
        config = self._criar_configuracao()
        ConfiguracaoSistema.objects.filter(pk=config.pk).update(unidade='ASSESSORIA DE COMUNICAÃ‡ÃƒO SOCIAL')
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))
        oficio.refresh_from_db()

        mapping = build_oficio_template_context(oficio)

        self.assertEqual(mapping['nome_chefia'], 'Autoridade Assinante')
        self.assertEqual(mapping['cargo_chefia'], 'Analista Documental')
        self.assertEqual(mapping['divisao'], 'DIRETORIA DE POLÃCIA DO INTERIOR')
        self.assertEqual(mapping['unidade'], 'Assessoria de ComunicaÃ§Ã£o Social')
        self.assertEqual(mapping['unidade_cabecalho'], 'ASSESSORIA DE COMUNICAÃ‡ÃƒO SOCIAL')
        self.assertEqual(mapping['unidade_rodape'], 'Assessoria de ComunicaÃ§Ã£o Social')
        self.assertEqual(mapping['assunto_oficio'], '(AutorizaÃ§Ã£o)')
        self.assertEqual(mapping['viatura'], 'Viatura Documental')
        self.assertEqual(mapping['combustivel'], 'Gasolina')
        self.assertIn('Viajante Documental', mapping['col_servidor'])
        self.assertIn('Analista Documental', mapping['col_cargo'])
        self.assertEqual(mapping['placa'], mapping['placa'].upper())

    def test_validacao_documental_do_oficio_apto(self):
        self._criar_configuracao()
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))

        validation = validate_oficio_for_document_generation(oficio, DocumentoOficioTipo.OFICIO)

        self.assertTrue(validation['ok'])
        self.assertEqual(validation['errors'], [])

    def test_validacao_documental_bloqueia_oficio_com_pendencia(self):
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_steps_1_e_2(oficio)

        validation = validate_oficio_for_document_generation(oficio, DocumentoOficioTipo.OFICIO)

        self.assertFalse(validation['ok'])
        self.assertTrue(any('Step 3' in error for error in validation['errors']))

    def test_validacao_documental_do_termo_autorizacao_apto(self):
        self._criar_configuracao()
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))

        validation = validate_oficio_for_document_generation(oficio, DocumentoOficioTipo.TERMO_AUTORIZACAO)

        self.assertTrue(validation['ok'])
        self.assertEqual(validation['errors'], [])

    def test_validacao_documental_do_termo_nao_exige_assinatura_de_chefia(self):
        self._criar_configuracao()
        AssinaturaConfiguracao.objects.filter(
            tipo=AssinaturaConfiguracao.TIPO_TERMO_AUTORIZACAO,
        ).delete()
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))

        validation = validate_oficio_for_document_generation(oficio, DocumentoOficioTipo.TERMO_AUTORIZACAO)

        self.assertTrue(validation['ok'])
        self.assertEqual(validation['errors'], [])

    def test_validacao_documental_do_plano_trabalho_apto(self):
        self._criar_configuracao()
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))

        validation = validate_oficio_for_document_generation(oficio, DocumentoOficioTipo.PLANO_TRABALHO)

        self.assertTrue(validation['ok'])
        self.assertEqual(validation['errors'], [])

    def test_validacao_documental_da_ordem_servico_apta(self):
        self._criar_configuracao()
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))

        validation = validate_oficio_for_document_generation(oficio, DocumentoOficioTipo.ORDEM_SERVICO)

        self.assertTrue(validation['ok'])
        self.assertEqual(validation['errors'], [])

    def test_download_do_oficio_e_bloqueado_sem_justificativa_obrigatoria(self):
        self._criar_configuracao()
        oficio = self._criar_oficio(data_criacao=date(2026, 10, 5))
        self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))

        response = self.client.get(
            reverse(
                'eventos:oficio-documento-download',
                kwargs={'pk': oficio.pk, 'tipo_documento': 'oficio', 'formato': 'docx'},
            ),
            follow=True,
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Preencha a justificativa')
        oficio.refresh_from_db()
        self.assertFalse(Justificativa.objects.filter(oficio=oficio).exists())

    def test_download_docx_do_oficio_quando_apto(self):
        self._criar_configuracao()
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))

        response = self.client.get(
            reverse(
                'eventos:oficio-documento-download',
                kwargs={'pk': oficio.pk, 'tipo_documento': 'oficio', 'formato': 'docx'},
            )
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response['Content-Type'],
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        self.assertIn(build_document_filename(oficio, DocumentoOficioTipo.OFICIO, 'docx'), response['Content-Disposition'])
        text = self._extract_docx_text(response.content)
        self.assertIn('Senhor Delegado', text)
        self.assertIn('Motivo documental', text)
        self.assertIn(self.viajante.nome, text)

    def test_download_docx_da_justificativa_quando_apta(self):
        self._criar_configuracao()
        oficio = self._criar_oficio(data_criacao=date(2026, 10, 5))
        self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))
        Justificativa.objects.update_or_create(
            oficio=oficio,
            defaults={'texto': 'Justificativa pronta para DOCX.'},
        )
        oficio.refresh_from_db()

        response = self.client.get(
            reverse(
                'eventos:oficio-documento-download',
                kwargs={'pk': oficio.pk, 'tipo_documento': 'justificativa', 'formato': 'docx'},
            )
        )

        self.assertEqual(response.status_code, 200)
        self.assertIn(build_document_filename(oficio, DocumentoOficioTipo.JUSTIFICATIVA, 'docx'), response['Content-Disposition'])
        text = self._extract_docx_text(response.content)
        self.assertIn('JUSTIFICATIVA', text)
        self.assertIn('Justificativa pronta para DOCX.', text)

    def test_download_docx_do_termo_autorizacao_quando_apto(self):
        self._criar_configuracao()
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))
        initial_response = self.client.get(
            reverse(
                'eventos:oficio-documento-download',
                kwargs={'pk': oficio.pk, 'tipo_documento': 'termo-autorizacao', 'formato': 'docx'},
            )
        )
        termo = TermoAutorizacao.objects.get(oficio=oficio)
        download_url = reverse(
            'eventos:documentos-termos-download',
            kwargs={'pk': termo.pk, 'formato': 'docx'},
        )
        self.assertRedirects(initial_response, download_url, fetch_redirect_response=False)

        response = self.client.get(download_url)

        self.assertEqual(response.status_code, 200)
        self.assertIn(f'termo_autorizacao_{termo.pk}_', response['Content-Disposition'])
        text = self._extract_docx_text(response.content)
        self.assertIn('TERMO DE AUTORIZA', text)
        self.assertIn(self.viajante.nome, text)

    def test_download_docx_do_plano_trabalho_quando_apto(self):
        self._criar_configuracao()
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))

        response = self.client.get(
            reverse(
                'eventos:oficio-documento-download',
                kwargs={'pk': oficio.pk, 'tipo_documento': 'plano-trabalho', 'formato': 'docx'},
            )
        )

        self.assertEqual(response.status_code, 200)
        self.assertIn(
            build_document_filename(oficio, DocumentoOficioTipo.PLANO_TRABALHO, 'docx'),
            response['Content-Disposition'],
        )
        text = self._extract_docx_text(response.content)
        self.assertIn('PLANO DE TRABALHO', text)
        self.assertIn('Objetivo / finalidade', text)
        self.assertIn('Motivo documental', text)

    def test_download_docx_da_ordem_servico_quando_apta(self):
        self._criar_configuracao()
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))

        response = self.client.get(
            reverse(
                'eventos:oficio-documento-download',
                kwargs={'pk': oficio.pk, 'tipo_documento': 'ordem-servico', 'formato': 'docx'},
            )
        )

        self.assertEqual(response.status_code, 200)
        self.assertIn(
            build_document_filename(oficio, DocumentoOficioTipo.ORDEM_SERVICO, 'docx'),
            response['Content-Disposition'],
        )
        text = self._extract_docx_text(response.content)
        self.assertIn('ORDEM DE SERVI', text)
        self.assertIn('O deslocamento do servidor', text)
        self.assertIn(self.viajante.nome, text)

    def test_templates_docx_versionados_estao_disponiveis(self):
        from eventos.services.documentos.renderer import get_document_template_path

        self.assertTrue(get_document_template_path(DocumentoOficioTipo.OFICIO).exists())
        self.assertTrue(get_document_template_path(DocumentoOficioTipo.JUSTIFICATIVA).exists())
        self.assertTrue(get_document_template_path(DocumentoOficioTipo.TERMO_AUTORIZACAO).exists())
        self.assertTrue(get_document_template_path(DocumentoOficioTipo.ORDEM_SERVICO).exists())

    def test_download_pdf_do_oficio_quando_apto(self):
        self._criar_configuracao()
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))

        with patch(
            'eventos.services.documentos.renderer.convert_docx_bytes_to_pdf_bytes',
            return_value=b'%PDF-1.4 oficio',
        ):
            response = self.client.get(
                reverse(
                    'eventos:oficio-documento-download',
                    kwargs={'pk': oficio.pk, 'tipo_documento': 'oficio', 'formato': 'pdf'},
                )
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response['Content-Type'], 'application/pdf')
        self.assertIn(build_document_filename(oficio, DocumentoOficioTipo.OFICIO, 'pdf'), response['Content-Disposition'])
        self.assertTrue(response.content.startswith(b'%PDF-1.4'))

    def test_download_pdf_da_justificativa_quando_apta(self):
        self._criar_configuracao()
        oficio = self._criar_oficio(data_criacao=date(2026, 10, 5))
        self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))
        Justificativa.objects.update_or_create(
            oficio=oficio,
            defaults={'texto': 'Justificativa pronta para PDF.'},
        )
        oficio.refresh_from_db()

        with patch(
            'eventos.services.documentos.renderer.convert_docx_bytes_to_pdf_bytes',
            return_value=b'%PDF-1.4 justificativa',
        ):
            response = self.client.get(
                reverse(
                    'eventos:oficio-documento-download',
                    kwargs={'pk': oficio.pk, 'tipo_documento': 'justificativa', 'formato': 'pdf'},
                )
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response['Content-Type'], 'application/pdf')
        self.assertIn(build_document_filename(oficio, DocumentoOficioTipo.JUSTIFICATIVA, 'pdf'), response['Content-Disposition'])
        self.assertTrue(response.content.startswith(b'%PDF-1.4'))

    def test_download_pdf_do_termo_autorizacao_quando_apto(self):
        self._criar_configuracao()
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))

        with patch(
            'eventos.views_global.convert_docx_bytes_to_pdf_bytes',
            return_value=b'%PDF-1.4 termo',
        ):
            initial_response = self.client.get(
                reverse(
                    'eventos:oficio-documento-download',
                    kwargs={'pk': oficio.pk, 'tipo_documento': 'termo-autorizacao', 'formato': 'pdf'},
                )
            )
            termo = TermoAutorizacao.objects.get(oficio=oficio)
            download_url = reverse(
                'eventos:documentos-termos-download',
                kwargs={'pk': termo.pk, 'formato': 'pdf'},
            )
            self.assertRedirects(initial_response, download_url, fetch_redirect_response=False)
            response = self.client.get(download_url)

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response['Content-Type'], 'application/pdf')
        self.assertIn(f'termo_autorizacao_{termo.pk}_', response['Content-Disposition'])
        self.assertTrue(response.content.startswith(b'%PDF-1.4'))

    def test_download_pdf_do_plano_trabalho_quando_apto(self):
        self._criar_configuracao()
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))

        with patch(
            'eventos.services.documentos.renderer.convert_docx_bytes_to_pdf_bytes',
            return_value=b'%PDF-1.4 plano',
        ):
            response = self.client.get(
                reverse(
                    'eventos:oficio-documento-download',
                    kwargs={'pk': oficio.pk, 'tipo_documento': 'plano-trabalho', 'formato': 'pdf'},
                )
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response['Content-Type'], 'application/pdf')
        self.assertIn(
            build_document_filename(oficio, DocumentoOficioTipo.PLANO_TRABALHO, 'pdf'),
            response['Content-Disposition'],
        )
        self.assertTrue(response.content.startswith(b'%PDF-1.4'))

    def test_download_pdf_da_ordem_servico_quando_apta(self):
        self._criar_configuracao()
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))

        with patch(
            'eventos.services.documentos.renderer.convert_docx_bytes_to_pdf_bytes',
            return_value=b'%PDF-1.4 ordem',
        ):
            response = self.client.get(
                reverse(
                    'eventos:oficio-documento-download',
                    kwargs={'pk': oficio.pk, 'tipo_documento': 'ordem-servico', 'formato': 'pdf'},
                )
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response['Content-Type'], 'application/pdf')
        self.assertIn(
            build_document_filename(oficio, DocumentoOficioTipo.ORDEM_SERVICO, 'pdf'),
            response['Content-Disposition'],
        )
        self.assertTrue(response.content.startswith(b'%PDF-1.4'))

    def test_download_do_termo_autorizacao_redireciona_para_registro_salvo_mesmo_com_oficio_incompleto(self):
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_steps_1_e_2(oficio)

        response = self.client.get(
            reverse(
                'eventos:oficio-documento-download',
                kwargs={'pk': oficio.pk, 'tipo_documento': 'termo-autorizacao', 'formato': 'docx'},
            )
        )
        termo = TermoAutorizacao.objects.get(oficio=oficio)
        self.assertRedirects(
            response,
            reverse('eventos:documentos-termos-download', kwargs={'pk': termo.pk, 'formato': 'docx'}),
            fetch_redirect_response=False,
        )

    def test_download_do_plano_trabalho_e_bloqueado_com_oficio_incompleto(self):
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_steps_1_e_2(oficio)

        response = self.client.get(
            reverse(
                'eventos:oficio-documento-download',
                kwargs={'pk': oficio.pk, 'tipo_documento': 'plano-trabalho', 'formato': 'docx'},
            ),
            follow=True,
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Preencha e salve o Step 3')

    def test_download_da_ordem_servico_e_bloqueado_com_oficio_incompleto(self):
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_steps_1_e_2(oficio)

        response = self.client.get(
            reverse(
                'eventos:oficio-documento-download',
                kwargs={'pk': oficio.pk, 'tipo_documento': 'ordem-servico', 'formato': 'docx'},
            ),
            follow=True,
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Preencha e salve o Step 3')

    def test_tela_de_documentos_mostra_status_corretos(self):
        self._criar_configuracao()
        oficio = self._criar_oficio(data_criacao=date(2026, 10, 5))
        self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))

        response = self.client.get(reverse('eventos:oficio-documentos', kwargs={'pk': oficio.pk}), follow=True)

        self.assertEqual(response.status_code, 200)
        self.assertRedirects(response, reverse('eventos:oficio-step4', kwargs={'pk': oficio.pk}))
        self.assertContains(response, 'Os downloads agora acontecem diretamente no resumo do oficio.')
        self.assertContains(response, 'Resumo do of')
        self.assertNotContains(response, 'Documentos do of')

    def test_tela_de_documentos_mostra_acoes_docx_e_pdf_quando_apto(self):
        self._criar_configuracao()
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))
        Justificativa.objects.update_or_create(
            oficio=oficio,
            defaults={'texto': 'Justificativa liberada.'},
        )
        oficio.refresh_from_db()

        response = self.client.get(reverse('eventos:oficio-documentos', kwargs={'pk': oficio.pk}), follow=True)

        self.assertEqual(response.status_code, 200)
        self.assertRedirects(response, reverse('eventos:oficio-step4', kwargs={'pk': oficio.pk}))
        self.assertContains(response, 'Os downloads agora acontecem diretamente no resumo do oficio.')
        self.assertContains(response, 'Resumo do of')

    def test_download_docx_indisponivel_redireciona_para_resumo_do_oficio(self):
        self._criar_configuracao()
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))

        with patch(
            'eventos.views.get_document_generation_status',
            return_value={
                'status': 'unavailable',
                'errors': ['Backend DOCX nao instalado neste ambiente atual.'],
            },
        ):
            response = self.client.get(
                reverse(
                    'eventos:oficio-documento-download',
                    kwargs={'pk': oficio.pk, 'tipo_documento': 'oficio', 'formato': 'docx'},
                ),
                follow=True,
            )

        self.assertEqual(response.status_code, 200)
        self.assertRedirects(response, reverse('eventos:oficio-step4', kwargs={'pk': oficio.pk}))
        self.assertContains(response, 'Backend DOCX nao instalado neste ambiente atual.')
        self.assertNotContains(response, 'Documentos do of')

    def test_download_pdf_indisponivel_redireciona_para_resumo_do_oficio(self):
        self._criar_configuracao()
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))

        with patch(
            'eventos.views.get_document_generation_status',
            return_value={
                'status': 'unavailable',
                'errors': ['Backend PDF indisponivel neste ambiente atual.'],
            },
        ):
            response = self.client.get(
                reverse(
                    'eventos:oficio-documento-download',
                    kwargs={'pk': oficio.pk, 'tipo_documento': 'oficio', 'formato': 'pdf'},
                ),
                follow=True,
            )

        self.assertEqual(response.status_code, 200)
        self.assertRedirects(response, reverse('eventos:oficio-step4', kwargs={'pk': oficio.pk}))
        self.assertContains(response, 'Backend PDF indisponivel neste ambiente atual.')
        self.assertNotContains(response, 'Documentos do of')

    def test_leitura_de_configuracoes_e_assinaturas_nao_quebra_sem_dados(self):
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))

        context = build_oficio_document_context(oficio)

        self.assertEqual(context['institucional']['orgao'], '')
        self.assertEqual(get_assinaturas_documento(DocumentoOficioTipo.OFICIO), [])
        self.assertEqual(context['assinaturas'], [])

    def test_import_das_urls_nao_quebra_sem_python_docx(self):
        import eventos.services.documentos as documentos_package
        import eventos.urls as eventos_urls
        import eventos.views as eventos_views

        real_import = builtins.__import__

        def guarded_import(name, globals=None, locals=None, fromlist=(), level=0):
            if name == 'docx' or name.startswith('docx.'):
                raise ImportError('docx ausente para teste')
            return real_import(name, globals, locals, fromlist, level)

        with patch('builtins.__import__', side_effect=guarded_import):
            importlib.reload(documentos_package)
            importlib.reload(eventos_views)
            importlib.reload(eventos_urls)

        importlib.reload(documentos_package)
        importlib.reload(eventos_views)
        importlib.reload(eventos_urls)

    def test_capabilities_separam_docx_e_pdf_quando_docx2pdf_falta(self):
        def module_status_side_effect(module_name, *, label, install_hint=''):
            if module_name == 'docx2pdf':
                return {
                    'available': False,
                    'module': None,
                    'reason': 'docx2pdf nÃ£o estÃ¡ instalado neste ambiente. Instale docx2pdf.',
                    'exception': ImportError('docx2pdf ausente'),
                }
            return {
                'available': True,
                'module': object(),
                'reason': '',
                'exception': None,
            }

        with patch(
            'eventos.services.documentos.backends._module_status',
            side_effect=module_status_side_effect,
        ), patch(
            'eventos.services.documentos.backends._check_word_com_availability',
            return_value={'available': True, 'reason': ''},
        ):
            reset_document_backend_capabilities_cache()
            capabilities = get_document_backend_capabilities()
            docx_status = get_docx_backend_availability()
            pdf_status = get_pdf_backend_availability()

        self.assertTrue(capabilities['docx_available'])
        self.assertFalse(capabilities['pdf_available'])
        self.assertTrue(docx_status['available'])
        self.assertFalse(pdf_status['available'])
        self.assertIn('docx2pdf', pdf_status['message'])

    def test_capabilities_reportam_pywin32_indisponivel_para_pdf(self):
        def module_status_side_effect(module_name, *, label, install_hint=''):
            if module_name == 'win32com.client':
                return {
                    'available': False,
                    'module': None,
                    'reason': 'pywin32 / win32com.client nÃ£o estÃ¡ disponÃ­vel neste ambiente. Instale pywin32.',
                    'exception': ImportError('win32com ausente'),
                }
            return {
                'available': True,
                'module': object(),
                'reason': '',
                'exception': None,
            }

        with patch(
            'eventos.services.documentos.backends._module_status',
            side_effect=module_status_side_effect,
        ):
            reset_document_backend_capabilities_cache()
            pdf_status = get_pdf_backend_availability()

        self.assertFalse(pdf_status['available'])
        self.assertIn('pywin32', pdf_status['message'])

    def test_capabilities_reportam_word_com_indisponivel_para_pdf(self):
        def module_status_side_effect(module_name, *, label, install_hint=''):
            return {
                'available': True,
                'module': object(),
                'reason': '',
                'exception': None,
            }

        with patch(
            'eventos.services.documentos.backends._module_status',
            side_effect=module_status_side_effect,
        ), patch(
            'eventos.services.documentos.backends._check_word_com_availability',
            return_value={
                'available': False,
                'reason': 'Microsoft Word / COM nÃ£o estÃ¡ disponÃ­vel para conversÃ£o PDF neste ambiente Windows (RuntimeError).',
            },
        ):
            reset_document_backend_capabilities_cache()
            pdf_status = get_pdf_backend_availability()

        self.assertFalse(pdf_status['available'])
        self.assertIn('Microsoft Word / COM', pdf_status['message'])

    def test_check_word_com_falha_quando_documents_nao_esta_acessivel(self):
        from eventos.services.documentos.backends import _check_word_com_availability

        class BrokenWordApplication:
            DisplayAlerts = 1

            def __init__(self):
                self.visible = True
                self.quit_called = False

            @property
            def Visible(self):
                return self.visible

            @Visible.setter
            def Visible(self, value):
                self.visible = value

            def Quit(self):
                self.quit_called = True

        broken_word = BrokenWordApplication()
        win32_client_stub = MagicMock()
        win32_client_stub.DispatchEx.return_value = broken_word
        real_import_module = importlib.import_module

        def import_module_side_effect(module_name):
            if module_name == 'win32com.client':
                return win32_client_stub
            return real_import_module(module_name)

        with patch(
            'eventos.services.documentos.backends.importlib.import_module',
            side_effect=import_module_side_effect,
        ):
            _check_word_com_availability.cache_clear()
            status = _check_word_com_availability()
            _check_word_com_availability.cache_clear()

        self.assertFalse(status['available'])
        self.assertIn('Microsoft Word / COM', status['reason'])
        self.assertTrue(broken_word.quit_called)

    def test_download_docx_continua_funcionando_mesmo_quando_pdf_esta_indisponivel(self):
        self._criar_configuracao()
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))

        def backend_side_effect(formato):
            if getattr(formato, 'value', str(formato)) == 'pdf':
                return {'available': False, 'message': 'Backend PDF indisponÃ­vel neste ambiente atual.', 'reasons': []}
            return {'available': True, 'message': '', 'reasons': []}

        with patch(
            'eventos.services.documentos.validators.get_document_backend_availability',
            side_effect=backend_side_effect,
        ):
            response = self.client.get(
                reverse(
                    'eventos:oficio-documento-download',
                    kwargs={'pk': oficio.pk, 'tipo_documento': 'oficio', 'formato': 'docx'},
                )
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response['Content-Type'],
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )

    @unittest.skip('Fluxo antigo de tela de documentos removido; cobertura migrada para downloads diretos e lista global.')
    def test_tela_documentos_mostra_status_separados_por_formato(self):
        self._criar_configuracao()
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))

        def backend_side_effect(formato):
            if getattr(formato, 'value', str(formato)) == 'pdf':
                return {
                    'available': False,
                    'message': 'docx2pdf nÃ£o estÃ¡ instalado neste ambiente. Instale docx2pdf.',
                    'reasons': ['docx2pdf nÃ£o estÃ¡ instalado neste ambiente. Instale docx2pdf.'],
                }
            return {'available': True, 'message': '', 'reasons': []}

        with patch(
            'eventos.views.get_pdf_backend_availability',
            return_value={
                'available': False,
                'message': 'docx2pdf nÃ£o estÃ¡ instalado neste ambiente. Instale docx2pdf.',
                'reasons': ['docx2pdf nÃ£o estÃ¡ instalado neste ambiente. Instale docx2pdf.'],
            },
        ), patch(
            'eventos.services.documentos.validators.get_document_backend_availability',
            side_effect=backend_side_effect,
        ):
            response = self.client.get(reverse('eventos:oficio-documentos', kwargs={'pk': oficio.pk}), follow=True)

        self.assertEqual(response.status_code, 200)
        oficio_item = next(item for item in response.context['documentos'] if item['slug'] == 'oficio')
        self.assertEqual(oficio_item['docx']['status'], 'available')
        self.assertEqual(oficio_item['pdf']['status'], 'unavailable')
        self.assertContains(response, 'DOCX: Dispon')
        self.assertContains(response, 'PDF: Indispon')
        self.assertContains(response, 'docx2pdf nÃ£o estÃ¡ instalado')

    @unittest.skip('Fluxo antigo de tela de documentos removido; cobertura migrada para downloads diretos e lista global.')
    def test_tela_documentos_abre_com_backend_docx_indisponivel(self):
        self._criar_configuracao()
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))

        with patch(
            'eventos.views.get_docx_backend_availability',
            return_value={'available': False, 'message': 'Backend DOCX n?o instalado neste ambiente atual.'},
        ), patch(
            'eventos.services.documentos.validators.get_document_backend_availability',
            return_value={'available': False, 'message': 'Backend DOCX n?o instalado neste ambiente atual.'},
        ):
            response = self.client.get(reverse('eventos:oficio-documentos', kwargs={'pk': oficio.pk}))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Backend DOCX')
        self.assertContains(response, 'DOCX indispon')

    @unittest.skip('Fluxo antigo de tela de documentos removido; cobertura migrada para downloads diretos e lista global.')
    def test_download_redireciona_com_mensagem_quando_backend_docx_indisponivel(self):
        self._criar_configuracao()
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))

        with patch(
            'eventos.services.documentos.validators.get_document_backend_availability',
            return_value={'available': False, 'message': 'Backend DOCX n?o instalado neste ambiente atual.'},
        ):
            response = self.client.get(
                reverse(
                    'eventos:oficio-documento-download',
                    kwargs={'pk': oficio.pk, 'tipo_documento': 'oficio', 'formato': 'docx'},
                ),
                follow=True,
            )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Backend DOCX')
        self.assertContains(response, 'Documentos do of')

    @unittest.skip('Fluxo antigo de tela de documentos removido; cobertura migrada para downloads diretos e lista global.')
    def test_tela_documentos_abre_com_backend_pdf_indisponivel(self):
        self._criar_configuracao()
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))

        def backend_side_effect(formato):
            if getattr(formato, 'value', str(formato)) == 'pdf':
                return {'available': False, 'message': 'Backend PDF indispon?vel neste ambiente atual.'}
            return {'available': True, 'message': ''}

        with patch(
            'eventos.views.get_pdf_backend_availability',
            return_value={'available': False, 'message': 'Backend PDF indispon?vel neste ambiente atual.'},
        ), patch(
            'eventos.services.documentos.validators.get_document_backend_availability',
            side_effect=backend_side_effect,
        ):
            response = self.client.get(reverse('eventos:oficio-documentos', kwargs={'pk': oficio.pk}))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Backend PDF indispon')
        self.assertContains(response, 'PDF indispon')

    @unittest.skip('Fluxo antigo de tela de documentos removido; cobertura migrada para downloads diretos e lista global.')
    def test_download_pdf_redireciona_com_mensagem_quando_backend_pdf_indisponivel(self):
        self._criar_configuracao()
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        self._salvar_oficio_finalizavel(oficio, date(2026, 10, 10), date(2026, 10, 11))

        def backend_side_effect(formato):
            if getattr(formato, 'value', str(formato)) == 'pdf':
                return {'available': False, 'message': 'Backend PDF indispon?vel neste ambiente atual.'}
            return {'available': True, 'message': ''}

        with patch(
            'eventos.services.documentos.validators.get_document_backend_availability',
            side_effect=backend_side_effect,
        ):
            response = self.client.get(
                reverse(
                    'eventos:oficio-documento-download',
                    kwargs={'pk': oficio.pk, 'tipo_documento': 'oficio', 'formato': 'pdf'},
                ),
                follow=True,
            )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Backend PDF indispon')
        self.assertContains(response, 'Documentos do of')

    @unittest.skip('Fluxo antigo de tela de documentos removido; cobertura migrada para downloads diretos e lista global.')
    def test_tela_documentos_abre_sem_dependencia_do_check_legado_de_schema(self):
        oficio = self._criar_oficio(data_criacao=date(2026, 9, 20))
        response = self.client.get(reverse('eventos:oficio-documentos', kwargs={'pk': oficio.pk}))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Documentos do of')


class OficioGeracaoDocumentoCorrecoesTest(TestCase):
    """Testes para as correÃ§Ãµes de destino, assunto e protocolo no gerador de OfÃ­cios."""

    def setUp(self):
        self.estado_pr = Estado.objects.create(nome='ParanÃ¡', sigla='PR', codigo_ibge='41')
        self.estado_sp = Estado.objects.create(nome='SÃ£o Paulo', sigla='SP', codigo_ibge='35')
        self.cidade_ctba = Cidade.objects.create(nome='Curitiba', estado=self.estado_pr, codigo_ibge='4106902')
        self.cidade_sp = Cidade.objects.create(nome='SÃ£o Paulo', estado=self.estado_sp, codigo_ibge='3550308')
        self.cidade_interior_pr = Cidade.objects.create(nome='Londrina', estado=self.estado_pr, codigo_ibge='4113700')

    def _criar_oficio_simples(self, assunto_tipo=None):
        oficio = Oficio.objects.create(status=Oficio.STATUS_RASCUNHO)
        if assunto_tipo:
            Oficio.objects.filter(pk=oficio.pk).update(assunto_tipo=assunto_tipo)
            oficio.refresh_from_db()
        return oficio

    def _add_trecho(self, oficio, destino_cidade):
        OficioTrecho.objects.create(
            oficio=oficio,
            ordem=oficio.trechos.count(),
            destino_estado=destino_cidade.estado,
            destino_cidade=destino_cidade,
        )

    def test_destino_cabecalho_e_sesp_quando_destino_fora_do_parana(self):
        from eventos.services.documentos.oficio import _get_destino_cabecalho_oficio, DESTINO_FORA_PARANA
        oficio = self._criar_oficio_simples()
        self._add_trecho(oficio, self.cidade_sp)

        resultado = _get_destino_cabecalho_oficio(oficio)

        self.assertEqual(resultado, DESTINO_FORA_PARANA)
        self.assertEqual(resultado, 'SESP')

    def test_destino_cabecalho_e_gabinete_quando_destino_dentro_do_parana(self):
        from eventos.services.documentos.oficio import _get_destino_cabecalho_oficio, DESTINO_DENTRO_PARANA
        oficio = self._criar_oficio_simples()
        self._add_trecho(oficio, self.cidade_interior_pr)

        resultado = _get_destino_cabecalho_oficio(oficio)

        self.assertEqual(resultado, DESTINO_DENTRO_PARANA)
        self.assertEqual(resultado, 'Gabinete do Delegado Geral Adjunto')

    def test_multiplos_destinos_com_um_fora_do_pr_resulta_em_sesp(self):
        from eventos.services.documentos.oficio import _get_destino_cabecalho_oficio, DESTINO_FORA_PARANA
        oficio = self._criar_oficio_simples()
        self._add_trecho(oficio, self.cidade_ctba)
        self._add_trecho(oficio, self.cidade_sp)

        resultado = _get_destino_cabecalho_oficio(oficio)

        self.assertEqual(resultado, DESTINO_FORA_PARANA)

    def test_multiplos_destinos_todos_no_pr_resulta_em_gabinete(self):
        from eventos.services.documentos.oficio import _get_destino_cabecalho_oficio, DESTINO_DENTRO_PARANA
        oficio = self._criar_oficio_simples()
        self._add_trecho(oficio, self.cidade_ctba)
        self._add_trecho(oficio, self.cidade_interior_pr)

        resultado = _get_destino_cabecalho_oficio(oficio)

        self.assertEqual(resultado, DESTINO_DENTRO_PARANA)

    def test_oficio_sem_trechos_resulta_em_gabinete(self):
        from eventos.services.documentos.oficio import _get_destino_cabecalho_oficio, DESTINO_DENTRO_PARANA
        oficio = self._criar_oficio_simples()

        resultado = _get_destino_cabecalho_oficio(oficio)

        self.assertEqual(resultado, DESTINO_DENTRO_PARANA)

    def test_assunto_padrao_e_autorizacao(self):
        from eventos.services.documentos.oficio import _get_assunto_for_oficio, ASSUNTO_AUTORIZACAO
        oficio = self._criar_oficio_simples()

        resultado = _get_assunto_for_oficio(oficio)

        self.assertEqual(resultado, ASSUNTO_AUTORIZACAO)
        self.assertEqual(resultado, 'SolicitaÃ§Ã£o de autorizaÃ§Ã£o e concessÃ£o de diÃ¡rias.')

    def test_assunto_muda_para_convalidacao_quando_tipo_convalidacao(self):
        from eventos.services.documentos.oficio import _get_assunto_for_oficio, ASSUNTO_CONVALIDACAO
        oficio = self._criar_oficio_simples(assunto_tipo=Oficio.ASSUNTO_TIPO_CONVALIDACAO)

        resultado = _get_assunto_for_oficio(oficio)

        self.assertEqual(resultado, ASSUNTO_CONVALIDACAO)
        self.assertEqual(resultado, 'SolicitaÃ§Ã£o de convalidaÃ§Ã£o e concessÃ£o de diÃ¡rias.')

    def test_col_solicitacao_fica_em_branco(self):
        from eventos.services.documentos.oficio import _build_col_solicitacao
        context_fake = {'viajantes': [{'nome': 'Fulano'}, {'nome': 'Beltrano'}]}

        resultado = _build_col_solicitacao(context_fake)

        self.assertEqual(resultado, '')

    def test_assunto_tipo_default_e_autorizacao(self):
        oficio = Oficio.objects.create(status=Oficio.STATUS_RASCUNHO)
        oficio.refresh_from_db()

        self.assertEqual(oficio.assunto_tipo, Oficio.ASSUNTO_TIPO_AUTORIZACAO)

    def test_choices_assunto_tipo_corretos(self):
        self.assertIn(Oficio.ASSUNTO_TIPO_AUTORIZACAO, dict(Oficio.ASSUNTO_TIPO_CHOICES))
        self.assertIn(Oficio.ASSUNTO_TIPO_CONVALIDACAO, dict(Oficio.ASSUNTO_TIPO_CHOICES))


class OficioStep1AjustesFinosTest(TestCase):
    """Ajustes finos do Step 1 e gerenciadores auxiliares."""

    def setUp(self):
        self.user = User.objects.create_user(username='ajustes-finos', password='ajustes-finos')
        self.client = Client()
        self.client.login(username='ajustes-finos', password='ajustes-finos')
        self.estado = Estado.objects.create(nome='ParanÃ¡', sigla='PR', codigo_ibge='41')
        self.cidade = Cidade.objects.create(nome='Curitiba', estado=self.estado, codigo_ibge='4106902')
        self.evento = Evento.objects.create(
            titulo='Evento Ajustes Finos',
            data_inicio=date(2026, 3, 1),
            data_fim=date(2026, 3, 2),
            status=Evento.STATUS_RASCUNHO,
        )
        self.cargo = Cargo.objects.create(nome='ANALISTA AJUSTE', is_padrao=True)
        self.unidade = UnidadeLotacao.objects.create(nome='DPC AJUSTE')

    def _criar_oficio(self):
        return Oficio.objects.create(evento=self.evento, status=Oficio.STATUS_RASCUNHO)

    def test_lista_oficios_exibe_protocolo_mascarado_sem_texto_auxiliar(self):
        oficio = self._criar_oficio()
        Oficio.objects.filter(pk=oficio.pk).update(protocolo='123456789')
        response = self.client.get(reverse('eventos:guiado-etapa-5', kwargs={'evento_id': self.evento.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, '12.345.678-9')
        self.assertNotContains(response, 'Formato: XX.XXX.XXX-X')

    def test_form_modelo_motivo_nao_exibe_codigo_ordem_ativo(self):
        response = self.client.get(reverse('eventos:modelos-motivo-cadastrar'))
        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, 'id_codigo')
        self.assertNotContains(response, 'id_ordem')
        self.assertNotContains(response, 'id_ativo')

    def test_lista_modelos_ordem_alfabetica(self):
        ModeloMotivoViagem.objects.create(codigo='zzz', nome='ZZZ', texto='z')
        ModeloMotivoViagem.objects.create(codigo='aaa', nome='AAA', texto='a')
        response = self.client.get(reverse('eventos:modelos-motivo-lista'))
        self.assertEqual(response.status_code, 200)
        html = response.content.decode()
        self.assertLess(html.index('AAA'), html.index('ZZZ'))

    def test_salvar_modelo_redireciona_para_lista(self):
        oficio = self._criar_oficio()
        url = f"{reverse('eventos:modelos-motivo-cadastrar')}?volta_step1={oficio.pk}"
        response = self.client.post(url, data={'nome': 'Modelo Novo', 'texto': 'Texto novo'})
        self.assertRedirects(
            response,
            f"{reverse('eventos:modelos-motivo-lista')}?volta_step1={oficio.pk}",
        )

    def test_botao_definir_padrao_e_excluir_aparecem_na_lista(self):
        modelo = ModeloMotivoViagem.objects.create(codigo='modelo_btn', nome='Modelo BotÃ£o', texto='Texto')
        response = self.client.get(reverse('eventos:modelos-motivo-lista'))
        self.assertEqual(response.status_code, 200)
        self.assertContains(
            response,
            reverse('eventos:modelos-motivo-definir-padrao', kwargs={'pk': modelo.pk}),
        )
        self.assertContains(
            response,
            reverse('eventos:modelos-motivo-excluir', kwargs={'pk': modelo.pk}),
        )

    def test_definir_novo_padrao_desmarca_anterior(self):
        m1 = ModeloMotivoViagem.objects.create(codigo='padrao_1', nome='PadrÃ£o 1', texto='1', padrao=True)
        m2 = ModeloMotivoViagem.objects.create(codigo='padrao_2', nome='PadrÃ£o 2', texto='2', padrao=False)
        response = self.client.post(reverse('eventos:modelos-motivo-definir-padrao', kwargs={'pk': m2.pk}))
        self.assertRedirects(response, reverse('eventos:modelos-motivo-lista'))
        m1.refresh_from_db()
        m2.refresh_from_db()
        self.assertFalse(m1.padrao)
        self.assertTrue(m2.padrao)

    def test_excluir_modelo_funciona(self):
        modelo = ModeloMotivoViagem.objects.create(codigo='modelo_excluir', nome='Excluir', texto='texto')
        response = self.client.post(reverse('eventos:modelos-motivo-excluir', kwargs={'pk': modelo.pk}))
        self.assertRedirects(response, reverse('eventos:modelos-motivo-lista'))
        self.assertFalse(ModeloMotivoViagem.objects.filter(pk=modelo.pk).exists())

    def test_step1_preseleciona_modelo_padrao(self):
        modelo = ModeloMotivoViagem.objects.create(
            codigo='modelo_padrao_step1',
            nome='Modelo PadrÃ£o Step1',
            texto='Texto padrÃ£o Step1',
            padrao=True,
        )
        oficio = self._criar_oficio()
        response = self.client.get(reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, f'<option value="{modelo.pk}" selected>')
        self.assertContains(response, 'Texto padrÃ£o Step1')

    def test_cadastrar_novo_viajante_com_next_retorna_para_step1(self):
        oficio = self._criar_oficio()
        step1_url = reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk})
        cadastrar_url = f"{reverse('cadastros:viajante-cadastrar')}?next={quote(step1_url)}"
        response = self.client.post(
            cadastrar_url,
            data={
                'nome': 'Viajante Retorno',
                'cargo': str(self.cargo.pk),
                'rg': '',
                'sem_rg': 'on',
                'cpf': '529.982.247-25',
                'telefone': '(41) 99999-8888',
                'unidade_lotacao': str(self.unidade.pk),
                'next': step1_url,
            },
        )
        self.assertRedirects(response, step1_url)

    def test_fluxo_step1_continua_funcional_apos_retorno_do_viajante(self):
        oficio = self._criar_oficio()
        step1_url = reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk})
        cadastrar_url = f"{reverse('cadastros:viajante-cadastrar')}?next={quote(step1_url)}"
        self.client.post(
            cadastrar_url,
            data={
                'nome': 'Viajante Fluxo',
                'cargo': str(self.cargo.pk),
                'rg': '',
                'sem_rg': 'on',
                'cpf': '390.533.447-05',
                'telefone': '(41) 99999-7777',
                'unidade_lotacao': str(self.unidade.pk),
                'next': step1_url,
            },
        )
        viajante = Viajante.objects.get(nome='VIAJANTE FLUXO')
        response_get = self.client.get(step1_url)
        self.assertEqual(response_get.status_code, 200)
        api_response = self.client.get(
            reverse('eventos:oficio-step1-viajantes-api'),
            {'q': 'FLUXO'},
        )
        self.assertEqual(api_response.status_code, 200)
        self.assertEqual(api_response.json()['results'][0]['id'], viajante.pk)
        oficio.refresh_from_db()
        response_post = self.client.post(
            step1_url,
            data={
                'oficio_numero': oficio.numero_formatado,
                'protocolo': '12.345.678-9',
                'data_criacao': oficio.data_criacao.strftime('%d/%m/%Y'),
                'modelo_motivo': '',
                'motivo': 'Fluxo apÃ³s retorno',
                'assunto_tipo': Oficio.ASSUNTO_TIPO_AUTORIZACAO,
                'custeio_tipo': Oficio.CUSTEIO_UNIDADE,
                'nome_instituicao_custeio': '',
                'viajantes': [viajante.pk],
            },
        )
        self.assertEqual(response_post.status_code, 302)
        oficio.refresh_from_db()
        self.assertEqual(list(oficio.viajantes.values_list('pk', flat=True)), [viajante.pk])


class OficioStep1ProtocolRegressionTest(TestCase):
    """RegressÃµes crÃ­ticas: GET sem save indevido e protocolo canÃ´nico/visual desacoplados."""

    def setUp(self):
        self.user = User.objects.create_user(username='prot-reg', password='prot-reg')
        self.client = Client()
        self.client.login(username='prot-reg', password='prot-reg')
        self.estado = Estado.objects.create(nome='ParanÃ¡', sigla='PR', codigo_ibge='41')
        self.cidade = Cidade.objects.create(nome='Curitiba', estado=self.estado, codigo_ibge='4106902')
        self.evento = Evento.objects.create(
            titulo='Evento RegressÃ£o Protocolo',
            data_inicio=date(2026, 2, 1),
            data_fim=date(2026, 2, 2),
            status=Evento.STATUS_RASCUNHO,
        )
        self.cargo = Cargo.objects.create(nome='ANALISTA PROTO', is_padrao=True)
        self.unidade = UnidadeLotacao.objects.create(nome='DPC PROTO')
        self.viajante = Viajante.objects.create(
            nome='Viajante Protocolo',
            status=Viajante.STATUS_FINALIZADO,
            cargo=self.cargo,
            cpf='12312312312',
            telefone='41988887777',
            unidade_lotacao=self.unidade,
            rg='1234567890',
        )

    def _criar_oficio(self):
        return Oficio.objects.create(evento=self.evento, status=Oficio.STATUS_RASCUNHO)

    def _payload(self, oficio, protocolo):
        oficio.refresh_from_db()
        return {
            'oficio_numero': oficio.numero_formatado,
            'protocolo': protocolo,
            'data_criacao': oficio.data_criacao.strftime('%d/%m/%Y'),
            'motivo': 'Motivo protocolo',
            'assunto_tipo': Oficio.ASSUNTO_TIPO_AUTORIZACAO,
            'custeio_tipo': Oficio.CUSTEIO_UNIDADE,
            'nome_instituicao_custeio': '',
            'viajantes': [self.viajante.pk],
        }

    def test_get_step1_nao_chama_save_indevido(self):
        oficio = self._criar_oficio()
        url = reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk})
        with patch('eventos.views.Oficio.save') as save_mock:
            response = self.client.get(url)
        self.assertEqual(response.status_code, 200)
        save_mock.assert_not_called()

    def test_get_step1_abre_com_protocolo_salvo_sem_mascara(self):
        oficio = self._criar_oficio()
        Oficio.objects.filter(pk=oficio.pk).update(protocolo='123456789')
        response = self.client.get(reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, '12.345.678-9')

    def test_get_step1_abre_com_protocolo_legado_mascarado(self):
        oficio = self._criar_oficio()
        Oficio.objects.filter(pk=oficio.pk).update(protocolo='12.345.678-9')
        response = self.client.get(reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, '12.345.678-9')

    def test_post_protocolo_mascarado_salva_canonico(self):
        oficio = self._criar_oficio()
        response = self.client.post(
            reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk}),
            data=self._payload(oficio, '12.345.678-9'),
        )
        self.assertEqual(response.status_code, 302)
        oficio.refresh_from_db()
        self.assertEqual(oficio.protocolo, '123456789')

    def test_post_protocolo_sem_mascara_salva_canonico(self):
        oficio = self._criar_oficio()
        response = self.client.post(
            reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk}),
            data=self._payload(oficio, '123456789'),
        )
        self.assertEqual(response.status_code, 302)
        oficio.refresh_from_db()
        self.assertEqual(oficio.protocolo, '123456789')

    def test_autosave_step1_sem_header_xhr_salva_com_sendbeacon(self):
        oficio = self._criar_oficio()
        response = self.client.post(
            reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk}),
            data={
                **self._payload(oficio, '12.345.678-9'),
                'autosave': '1',
            },
        )
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()['ok'], True)
        oficio.refresh_from_db()
        self.assertEqual(oficio.protocolo, '123456789')
        self.assertEqual(oficio.motivo, 'Motivo protocolo')
        self.assertEqual(list(oficio.viajantes.values_list('pk', flat=True)), [self.viajante.pk])

    def test_reabrir_step1_reexibe_protocolo_mascarado(self):
        oficio = self._criar_oficio()
        self.client.post(
            reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk}),
            data=self._payload(oficio, '123456789'),
        )
        response = self.client.get(reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, '12.345.678-9')

    def test_glance_context_reexibe_protocolo_mascarado(self):
        oficio = self._criar_oficio()
        Oficio.objects.filter(pk=oficio.pk).update(protocolo='123456789')
        response = self.client.get(reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk}))
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.context['wizard_glance']['protocolo'], '12.345.678-9')

    def test_get_edicao_nao_muda_numero_nem_data(self):
        oficio = self._criar_oficio()
        oficio.refresh_from_db()
        numero_antes = oficio.numero
        ano_antes = oficio.ano
        data_antes = oficio.data_criacao
        response = self.client.get(reverse('eventos:oficio-step1', kwargs={'pk': oficio.pk}))
        self.assertEqual(response.status_code, 200)
        oficio.refresh_from_db()
        self.assertEqual(oficio.numero, numero_antes)
        self.assertEqual(oficio.ano, ano_antes)
        self.assertEqual(oficio.data_criacao, data_antes)


class DocumentosHubTest(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username='u_documentos', password='p_documentos')
        self.client = Client()
        self.client.login(username='u_documentos', password='p_documentos')

    def test_hub_nao_exibe_fluxo_generico_de_documentos(self):
        response = self.client.get(reverse('eventos:documentos-hub'))
        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, 'Novo documento' + ' avulso')
        self.assertNotContains(response, 'modalNovoDocumento' + 'Avulso')
        self.assertNotContains(response, 'Documentos criados no fluxo' + ' avulso')
        self.assertContains(response, 'Planos de trabalho')
        self.assertContains(response, 'Ordens de servico')
        self.assertContains(response, 'Justificativas')
        self.assertContains(response, 'Termos')


class EventoEtapa5TermosPadraoFluxoTest(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username='u_etapa5_termos', password='p_etapa5_termos')
        self.client = Client()
        self.client.login(username='u_etapa5_termos', password='p_etapa5_termos')
        self.estado = Estado.objects.create(nome='Parana', sigla='PR', codigo_ibge='41')
        self.cidade = Cidade.objects.create(nome='Curitiba', estado=self.estado, codigo_ibge='4106902')
        self.cargo = Cargo.objects.create(nome='Analista de testes', is_padrao=True)
        self.unidade = UnidadeLotacao.objects.create(nome='Unidade de testes')

        self.evento = Evento.objects.create(
            titulo='Evento com termos no fluxo',
            data_inicio=date(2026, 4, 10),
            data_fim=date(2026, 4, 12),
            cidade_principal=self.cidade,
            cidade_base=self.cidade,
            estado_principal=self.estado,
            status=Evento.STATUS_RASCUNHO,
        )
        self.evento_outro = Evento.objects.create(
            titulo='Outro evento com termos',
            data_inicio=date(2026, 5, 1),
            data_fim=date(2026, 5, 2),
            cidade_principal=self.cidade,
            cidade_base=self.cidade,
            estado_principal=self.estado,
            status=Evento.STATUS_RASCUNHO,
        )

        self.viajante_evento = Viajante.objects.create(
            nome='Viajante Etapa 5',
            cargo=self.cargo,
            unidade_lotacao=self.unidade,
            cpf='52998224725',
        )
        self.viajante_oficio = Viajante.objects.create(
            nome='Viajante Oficio Etapa 5',
            cargo=self.cargo,
            unidade_lotacao=self.unidade,
            cpf='39053344705',
        )
        self.viajante_outro = Viajante.objects.create(
            nome='Viajante Outro Evento',
            cargo=self.cargo,
            unidade_lotacao=self.unidade,
            cpf='16899535009',
        )

        self.oficio_evento = Oficio.objects.create(
            evento=self.evento,
            status=Oficio.STATUS_FINALIZADO,
            data_criacao=date(2026, 4, 10),
        )
        self.oficio_outro = Oficio.objects.create(
            evento=self.evento_outro,
            status=Oficio.STATUS_FINALIZADO,
            data_criacao=date(2026, 5, 1),
        )

        self.termo_evento = TermoAutorizacao.objects.create(
            evento=self.evento,
            oficio=self.oficio_evento,
            viajante=self.viajante_evento,
            destino='Curitiba/PR',
            data_evento=date(2026, 4, 10),
            data_evento_fim=date(2026, 4, 12),
            status=TermoAutorizacao.STATUS_GERADO,
        )
        self.termo_evento.oficios.add(self.oficio_evento)

        self.termo_por_oficio = TermoAutorizacao.objects.create(
            oficio=self.oficio_evento,
            viajante=self.viajante_oficio,
            destino='Londrina/PR',
            data_evento=date(2026, 4, 11),
            data_evento_fim=date(2026, 4, 11),
            status=TermoAutorizacao.STATUS_RASCUNHO,
        )
        self.termo_por_oficio.oficios.add(self.oficio_evento)

        self.termo_outro = TermoAutorizacao.objects.create(
            evento=self.evento_outro,
            oficio=self.oficio_outro,
            viajante=self.viajante_outro,
            destino='Maringa/PR',
            data_evento=date(2026, 5, 1),
            data_evento_fim=date(2026, 5, 2),
            status=TermoAutorizacao.STATUS_GERADO,
        )
        self.termo_outro.oficios.add(self.oficio_outro)

    def test_etapa_5_reaproveita_shell_moderno_e_filtra_termos_do_evento(self):
        response = self.client.get(reverse('eventos:guiado-etapa-5', kwargs={'evento_id': self.evento.pk}))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            [item.pk for item in response.context['object_list']],
            [self.termo_por_oficio.pk, self.termo_evento.pk],
        )
        self.assertContains(response, 'guided-flow-header-card', html=False)
        self.assertContains(response, 'data-view-storage-key="central-viagens.guiado.etapa5.termos.view-mode"', html=False)
        self.assertContains(response, 'data-list-view-toggle="rich"', html=False)
        self.assertContains(response, 'data-list-view-toggle="basic"', html=False)
        self.assertContains(response, 'data-scroll-top', html=False)
        self.assertContains(response, self.termo_evento.numero_formatado)
        self.assertContains(response, self.termo_por_oficio.numero_formatado)
        self.assertContains(response, self.viajante_evento.nome)
        self.assertContains(response, self.viajante_oficio.nome)
        self.assertNotContains(response, self.termo_outro.numero_formatado)
        self.assertNotContains(response, self.viajante_outro.nome)

        etapa_ativa = next(step for step in response.context['wizard_steps'] if step['key'] == 'termos')
        self.assertTrue(etapa_ativa['active'])
        self.assertEqual(etapa_ativa['number'], 5)

    def test_etapa_5_novo_termo_aponta_para_cadastro_real_contextualizado(self):
        response = self.client.get(reverse('eventos:guiado-etapa-5', kwargs={'evento_id': self.evento.pk}))

        self.assertEqual(response.status_code, 200)
        expected_return_to = quote(reverse('eventos:guiado-etapa-5', kwargs={'evento_id': self.evento.pk}))
        self.assertIn('context_source=evento', response.context['termo_novo_url'])
        self.assertIn(f'preselected_event_id={self.evento.pk}', response.context['termo_novo_url'])
        self.assertIn(f'return_to={expected_return_to}', response.context['termo_novo_url'])
        self.assertContains(response, f'preselected_event_id={self.evento.pk}')
        self.assertContains(response, 'Novo termo')



